# SPDX-License-Identifier: AGPL-3.0-or-later
"""redirecall.ingest — extracted from main.py (see main.py for architecture).

Split out mechanically for maintainability; cross-module references are
module-qualified so runtime rebinding and test monkeypatching stay live.
"""
import logging
import asyncio
import csv
import hashlib
import time
from datetime import datetime, timezone
from pathlib import Path
import redis
try:
    import fitz          # PyMuPDF — PDF text extraction
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
try:
    from docx import Document as _DocxDocument   # python-docx — DOCX text extraction
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
try:
    from openpyxl import load_workbook as _load_workbook  # openpyxl — XLSX text extraction
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
from . import config, rag, redis_store, state, textutil

log = logging.getLogger(__name__)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT INGESTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def ingest_text(
    instance: str,
    text: str,
    source: str,
    rc: redis.Redis | None = None,
) -> int:
    """
    Chunk raw text and store all chunks in the RAG index.
    Returns the number of chunks ingested.

    IDs are reserved atomically with next_chunk_id so concurrent ingests
    don't collide even without a distributed lock.
    """
    cfg = state._config.get("rag", {})
    chunks = textutil.chunk_text(
        text,
        cfg.get("chunk_size", 180),
        cfg.get("chunk_overlap", 32),
    )
    if not chunks:
        return 0

    rc = rc or redis_store.r()
    hash_set_key = f"rag:{instance}:chunk_hashes"

    # Deduplicate: compute all hashes, then check + add in a single pipeline.
    # This turns N round-trips into one network operation.
    # Scoped per source. A global content hash stores an identical chunk once,
    # so deleting the document that happened to be ingested first also removed
    # content the other document still needed — and released the shared hash, so
    # re-ingesting the survivor was then skipped as a duplicate. Scoping trades a
    # little storage for documents that can be deleted independently.
    chunk_hashes = [
        hashlib.sha256(f"{source}\x00{' '.join(c.lower().split())}".encode()).hexdigest()
        for c in chunks
    ]
    pipe = rc.pipeline(transaction=False)
    for h in chunk_hashes:
        pipe.sadd(hash_set_key, h)
    results = pipe.execute()          # one round-trip; returns [1 or 0, ...]
    new_chunks: list[str] = [c for c, added in zip(chunks, results) if added]

    if not new_chunks:
        return 0

    # Reserve N IDs at once — single O(1) Redis call
    start_id = rag.next_chunk_id(instance, len(new_chunks), rc)
    now = int(time.time())
    # `pos` is the ordinal within THIS document (not the global chunk_id), so
    # neighbouring chunks of the same source can be located later.
    records = [
        {"id": start_id + i, "text": c, "source": source, "ingested_at": now, "pos": i}
        for i, c in enumerate(new_chunks)
    ]
    rag.add_chunks(instance, records, rc)
    return len(records)


def _prepare_chunks(
    instance: str,
    text: str,
    source: str,
    rc: redis.Redis,
    force_reindex: bool = False,
) -> list[dict]:
    """
    Chunk, deduplicate, and reserve Redis IDs — but do NOT embed yet.
    Returns a list of {id, text, source} records ready to be embedded in a batch.
    Designed to be called via asyncio.to_thread(); it is CPU-light (no model calls).
    When force_reindex=True the hash-dedup filter is bypassed so previously seen
    chunks are re-embedded rather than silently dropped.
    """
    cfg = state._config.get("rag", {})
    chunks = textutil.chunk_text(
        text,
        cfg.get("chunk_size", 180),
        cfg.get("chunk_overlap", 32),
    )
    if not chunks:
        return []

    hash_set_key = f"rag:{instance}:chunk_hashes"
    # Scoped per source. A global content hash stores an identical chunk once,
    # so deleting the document that happened to be ingested first also removed
    # content the other document still needed — and released the shared hash, so
    # re-ingesting the survivor was then skipped as a duplicate. Scoping trades a
    # little storage for documents that can be deleted independently.
    chunk_hashes = [
        hashlib.sha256(f"{source}\x00{' '.join(c.lower().split())}".encode()).hexdigest()
        for c in chunks
    ]
    pipe = rc.pipeline(transaction=False)
    for h in chunk_hashes:
        pipe.sadd(hash_set_key, h)
    results = pipe.execute()

    if force_reindex:
        # Allow all chunks through; hashes were already updated above
        new_chunks = chunks
    else:
        new_chunks = [c for c, added in zip(chunks, results) if added]

    if not new_chunks:
        return []

    start_id = rag.next_chunk_id(instance, len(new_chunks), rc)
    now = int(time.time())
    return [{"id": start_id + i, "text": c, "source": source, "ingested_at": now, "pos": i}
            for i, c in enumerate(new_chunks)]


# Legacy binary .doc/.xls are deliberately NOT here: python-docx and openpyxl read
# only OOXML, so those files fail with an opaque "File is not a zip file".
# extract_file_text raises a clear message for them instead.
_CHAT_FILE_ACCEPT = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx"}
_CHAT_FILE_MAX_BYTES = 10 * 1024 * 1024   # 10 MB
_CHAT_FILE_MAX_CHARS = 150_000             # ~100 k tokens — keeps context manageable


def extract_file_text(filename: str, data: bytes) -> str:
    """Convert an uploaded document to plain text.

    Supported formats: .txt, .md, .csv, .pdf, .doc/.docx, .xls/.xlsx
    Raises ValueError for unsupported or unparseable files.
    """
    import io as _io

    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        return data.decode("utf-8", errors="ignore")

    if suffix == ".csv":
        text_io = _io.StringIO(data.decode("utf-8", errors="ignore"))
        rows = [" | ".join(row) for row in csv.reader(text_io)]
        return "\n".join(rows)

    if suffix == ".pdf":
        if not HAS_PYMUPDF:
            raise ValueError("PDF support requires PyMuPDF (pip install pymupdf)")
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)

    if suffix in (".doc", ".xls"):
        raise ValueError(
            f"Legacy binary {suffix} is not supported — re-save it as "
            f"{'.docx' if suffix == '.doc' else '.xlsx'} and upload again")

    if suffix == ".docx":
        if not HAS_PYTHON_DOCX:
            raise ValueError("DOCX support requires python-docx (pip install python-docx)")
        doc = _DocxDocument(_io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells if cell.text.strip()))
        return "\n".join(parts)

    if suffix == ".xlsx":
        if not HAS_OPENPYXL:
            raise ValueError("Excel support requires openpyxl (pip install openpyxl)")
        wb = _load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {suffix}. Accepted: {', '.join(sorted(_CHAT_FILE_ACCEPT))}")


async def ingest_file(
    instance: str,
    path: Path,
    source: str,
    rc: redis.Redis | None = None,
) -> dict:
    """
    Ingest a single file into a RAG instance.

    Formats are whatever ``extract_file_text`` understands — .txt, .md, .csv,
    .pdf, .doc/.docx, .xls/.xlsx — so the knowledge base accepts the same set as
    a chat attachment. (These were previously divergent: the docs advertised
    md/docx/xlsx while this path rejected them as "Unsupported type".)

    Returns a log entry dict with status information.
    """
    suffix = path.suffix.lower()

    # Parsing (a big PDF is seconds of CPU) and the embed+Redis write in ingest_text
    # are both synchronous — run them off the event loop so a large upload never freezes
    # concurrent chat sessions or the WS receive loop.
    def _extract():
        return extract_file_text(path.name, path.read_bytes())

    try:
        if suffix not in _CHAT_FILE_ACCEPT:
            hint = (f"Legacy binary {suffix} is not supported — re-save it as "
                    f"{'.docx' if suffix == '.doc' else '.xlsx'} and upload again"
                    if suffix in (".doc", ".xls") else f"Unsupported type: {suffix}")
            return {"source": source, "status": "skipped", "error": hint}

        text = await asyncio.to_thread(_extract)

        if not text.strip():
            # A scanned PDF extracts to nothing. Reporting "ok, 0 chunks" made this
            # look like a successful ingest; say plainly that nothing was indexed.
            entry = {
                "ts": datetime.now(timezone.utc).isoformat(),
                "instance": instance, "source": source, "chunks": 0,
                "status": "error",
                "error": "No extractable text (an image-only/scanned document needs OCR)",
            }
            config.append_log(entry)
            return entry

        n = await asyncio.to_thread(ingest_text, instance, text, source, rc)
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "instance": instance,
            "source": source,
            "chunks": n,
            "status": "ok",
        }
        config.append_log(entry)
        return entry

    except Exception as e:
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "instance": instance,
            "source": source,
            "chunks": 0,
            "status": "error",
            "error": str(e),
        }
        config.append_log(entry)
        return entry

