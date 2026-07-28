# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 Sebastian Cyris
#
# RediRecall is free software: you can redistribute it and/or modify it under
# the terms of the GNU Affero General Public License as published by the Free
# Software Foundation, either version 3 of the License, or (at your option) any
# later version. It is distributed WITHOUT ANY WARRANTY; see the LICENSE file,
# or <https://www.gnu.org/licenses/>, for details.
#
# Source: https://github.com/SFCyris/RediRecall
"""
RediRecall — FastAPI backend (RAG + semantic cache chat)
========================================================
Architecture overview
---------------------
* Redis is used for three things:
    1. RAG vector store  — chunks stored as HASH keys, searched via FT.SEARCH KNN
    2. Semantic cache    — query→response pairs with a vector index for fast lookup
    3. Instance metadata — lightweight JSON blobs (rag_meta:* keys)

* LLM providers supported: Ollama (local), Anthropic Claude (API), OpenAI (API)
  All three expose the same async-generator interface: stream → (token, done).

* Multiple Redis endpoints are supported so different RAG instances can live on
  different Redis servers (e.g. a shared enterprise cluster vs. a local dev box).

* Performance notes:
    - Embeddings are batched per-document at ingest time (one model call per file).
    - next_chunk_id() uses an atomic INCRBY counter instead of a KEYS scan.
    - cache_lookup() uses FT.SEARCH KNN — O(log N) instead of O(N) linear scan.
    - HNSW index uses M=16, EF_CONSTRUCTION=200 for a good quality/speed tradeoff.
    - Pipelines are used for all bulk writes.
"""

import asyncio
import base64
import copy
import csv
import hashlib
import io
import ipaddress
import json
import logging
import os
import re
import shutil
import socket
import sys
import threading
import time
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser

import httpx
import numpy as np
try:
    from google import genai as _google_genai
    from google.genai import types as _genai_types
    _GENAI_AVAILABLE = True
except ImportError:
    _GENAI_AVAILABLE = False

try:
    import anthropic as _anthropic
    _ANTHROPIC_AVAILABLE = True
except ImportError:
    _ANTHROPIC_AVAILABLE = False

try:
    from openai import AsyncOpenAI as _AsyncOpenAI
    _OPENAI_SDK_AVAILABLE = True
except ImportError:
    _OPENAI_SDK_AVAILABLE = False
import redis
from redisvl.index import SearchIndex
from redisvl.query import VectorQuery
from redisvl.schema import IndexSchema
from redisvl.extensions.cache.llm import SemanticCache
# HFTextVectorizer is imported lazily inside _make_cache_vectorizer() to keep
# startup fast (it triggers a ~1.7 s sentence-transformers/torch import chain).
from fastapi import FastAPI, File, Form, HTTPException, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
# sentence_transformers is imported lazily inside get_embed_model() to avoid
# blocking Python startup with the 5-15 s PyTorch/transformers import time.

# ── Optional heavy imports ────────────────────────────────────────────────────
# These enhance functionality but are not required to run the app.

try:
    import fitz          # PyMuPDF — PDF text extraction
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import trafilatura   # Best-in-class web content extraction
    HAS_TRAFILATURA = True
except ImportError:
    HAS_TRAFILATURA = False

try:
    from bs4 import BeautifulSoup   # HTML parsing fallback + link extraction
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    from crawl4ai import AsyncWebCrawler as _C4AIWebCrawler   # parallel JS-capable crawler
    from crawl4ai import BrowserConfig  as _C4AIBrowserConfig
    from crawl4ai import CrawlerRunConfig as _C4AIRunConfig
    HAS_CRAWL4AI = True
except ImportError:
    HAS_CRAWL4AI = False

import platform as _platform

# Performance flags passed to every Chromium instance.
# --no-sandbox is required in Linux containers/CI (no user namespace isolation).
# The rendering and networking flags cut per-tab resource use significantly.
_BROWSER_EXTRA_ARGS: list[str] = [
    "--disable-dev-shm-usage",           # avoid /dev/shm exhaustion in containers
    "--disable-gpu",                     # no GPU needed for headless text extraction
    "--disable-background-networking",   # suppress background pings (Safe Browsing etc.)
    "--disable-sync",                    # no Chrome account sync
    "--no-first-run",                    # skip first-run wizard
    "--disable-extensions",              # no extensions
    "--blink-settings=imagesEnabled=false",  # block image rendering at the engine level
]
if _platform.system() == "Linux":
    _BROWSER_EXTRA_ARGS += ["--no-sandbox", "--disable-setuid-sandbox"]

try:
    from sentence_transformers import CrossEncoder
    HAS_CROSSENCODER = True
except ImportError:
    HAS_CROSSENCODER = False

try:
    from docx import Document as _DocxDocument   # python-docx — DOCX text extraction
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    from openpyxl import load_workbook as _load_workbook  # openpyxl — XLSX text extraction
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CONSTANTS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _default_data_dir() -> Path:
    """Platform-appropriate local data directory (always on local disk).

    Priority:
      1. REDIRECALL_DATA_DIR env var (used by the Docker image and run.sh)
      2. Platform default:
           macOS:  ~/Library/Application Support/RediRecall
           Linux:  $XDG_DATA_HOME/redirecall  or  ~/.local/share/redirecall
           other:  ~/.redirecall
    """
    env = os.environ.get("REDIRECALL_DATA_DIR", "").strip()
    if env:
        return Path(env).expanduser()
    if sys.platform == "darwin":
        return Path.home() / "Library" / "Application Support" / "RediRecall"
    if sys.platform.startswith("linux"):
        xdg = os.environ.get("XDG_DATA_HOME", "").strip()
        base = Path(xdg) if xdg else Path.home() / ".local" / "share"
        return base / "redirecall"
    return Path.home() / ".redirecall"


DATA_DIR = _default_data_dir()
DATA_DIR.mkdir(parents=True, exist_ok=True)

CONFIG_PATH   = DATA_DIR / "config.json"
LOGS_PATH     = DATA_DIR / "ingestion_logs.json"
FEEDBACK_PATH = DATA_DIR / "feedback.json"
UPLOAD_DIR    = DATA_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def _migrate_legacy_state() -> None:
    """Relocate state files that older versions kept next to the code.

    Copies (never deletes) any config/log/feedback file found beside this
    module into DATA_DIR when the new location does not yet hold one, so an
    in-place upgrade keeps the user's existing settings and history.
    """
    # Older versions kept state beside main.py (project root). Since main.py now
    # lives in the redirecall/ package, check the package dir, its parent (the
    # repo root), and the current working directory.
    here = Path(__file__).resolve().parent
    legacy_roots = [here, here.parent, Path.cwd()]
    for name, target in (
        ("config.json",         CONFIG_PATH),
        ("ingestion_logs.json", LOGS_PATH),
        ("feedback.json",       FEEDBACK_PATH),
    ):
        if target.exists():
            continue
        for legacy_root in legacy_roots:
            legacy = legacy_root / name
            try:
                if legacy.exists():
                    shutil.copy2(legacy, target)
                    log.info(f"Migrated legacy state file {legacy} -> {target}")
                    break
            except Exception as e:  # never let a migration hiccup block startup
                log.warning(f"Could not migrate legacy {legacy}: {e}")


_migrate_legacy_state()


def safe_upload_dest(filename: str) -> tuple[Path, str]:
    """Resolve a client-supplied upload filename to a safe path inside UPLOAD_DIR.

    Strips every path component (``../``, absolute paths, Windows separators) to
    the bare basename, rejects empty/dot names, and verifies the resolved path
    stays within UPLOAD_DIR — defeating path-traversal via a crafted filename.
    Returns ``(dest_path, safe_name)``; raises HTTPException(400) on a bad name.
    """
    raw = (filename or "").replace("\\", "/")
    safe_name = Path(raw).name
    if not safe_name or safe_name in (".", ".."):
        raise HTTPException(status_code=400, detail="Invalid filename")
    dest = (UPLOAD_DIR / safe_name).resolve()
    if not str(dest).startswith(str(UPLOAD_DIR.resolve()) + os.sep):
        raise HTTPException(status_code=400, detail="Invalid filename")
    return dest, safe_name

# Redis key prefix used by the redisvl SemanticCache (must match SemanticCache name).
CACHE_PREFIX = "semcache:"

# Shared SemanticCache instance — lazily created, reset when Redis config changes.
_semantic_cache: "SemanticCache | None" = None
# Building the cache's HF vectorizer takes ~3 s (model load). The background warm
# (_bg_init) does it off the request path, but Uvicorn serves immediately, so a chat in
# the first few seconds after a restart would otherwise pay that build inline. This flag
# lets cache_lookup/store treat the not-yet-warm cache as a miss until the warm finishes.
_semantic_cache_ready = False

# Per-endpoint Search module availability cache.
# Key = endpoint name ("default" for primary).  Value = True/False/None (None = unchecked).
_search_available: dict[str, bool | None] = {}

# ── Identity / AGPL §13 source offer ─────────────────────────────────────────
# Where a network user can obtain the Corresponding Source. Override only if you
# publish YOUR modified source somewhere else — AGPL-3.0 §13 requires the offer to
# point at the source of the version actually running.
from redirecall import __version__
SOURCE_URL = os.environ.get("REDIRECALL_SOURCE_URL", "https://github.com/SFCyris/RediRecall")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DEFAULT CONFIGURATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# This dict is merged with config.json on startup so new keys are always
# available even on an existing install.  Nested dicts are merged shallowly.

# Default global system instruction. Prepended to every chat turn (before any
# selected template). Editable in Settings -> Templates -> Base Instruction.
# Tuned to this app's SVG renderer + DOMPurify sanitiser (see the marked `code`
# handler and the svg-profile sanitise in index.html) so charts actually render.
DEFAULT_BASE_INSTRUCTION = """You are RediRecall's assistant. Answer clearly and concisely: prefer short paragraphs and bullet lists, put tabular data in Markdown tables, and write mathematical formulas as LaTeX (inline $...$ or a ```latex fenced block).

Draw a chart ONLY when the user asks for one (graph/plot/chart/diagram/visualise) or when a picture genuinely makes numeric data easier to grasp — otherwise just answer in text. You may put normal prose, tables, and LaTeX around a chart, but every chart goes in its OWN fenced block.

=== VISUAL BLOCKS — pick the RIGHT fence; the app renders each one ===
Always prefer one of these over hand-drawing an SVG: you supply a short description and the app does the drawing exactly.
Every fenced block below is BLOCK-LEVEL: put it on its own lines with a blank line before and after, and with REAL newlines inside (never the literal characters backslash-n). NEVER place a fence inside a Markdown table cell, list item, or blockquote — a fence in a table cell renders as raw text, not a figure. To show structures/figures for several items, put the plain formula or SMILES string as text in the table, then render each item as its own labelled block AFTER the table (e.g. a short heading followed by a ```molecule block).
- ```plot — a function graph. Body: `y = x^2 + 3x - 2` (one function per line; optional `x = -5 .. 5`). The app has NO predefined constants beyond a bare number you write yourself — physical/mathematical constants (speed of light, gravitational constant, Boltzmann constant, …) are NOT built in. If a formula uses a named constant instead of writing its number directly, you MUST declare it with its own `param:` line, BEFORE any other line that uses it — never leave a symbol undeclared and never assume the app knows what it stands for. A fixed (non-adjustable) constant is just a param with equal bounds, e.g. for `y = sqrt(1-(v/c)^2)` declare `param: c = 299792458 .. 299792458 (299792458)` first, then `param: v = 0 .. 0.99c (0.01)` — a param's own bounds can reference an EARLIER param (implicit multiplication: `0.99c` = 0.99×c). An undeclared symbol, or a bound that references a constant declared later or not at all, makes evaluation fail and the whole plot error out. Declaring a param with a REAL range instead renders a live slider that re-plots instantly with no further request: `y = a*sin(b*x)` then `param: a = 0.5 .. 3 (1)` and `param: b = 1 .. 5 (2)`. Prefer this over answering the same question again for a different value.
- ```chart — data chart (bar/line/pie/doughnut/scatter/radar). Body: Chart.js JSON, e.g. {"type":"bar","data":{"labels":["Q1","Q2"],"datasets":[{"label":"Sales","data":[120,150]}]}}
- ```mermaid — flowchart, sequence, class, state, ER or Gantt diagram. Body: mermaid syntax, e.g. `graph TD` then `A[Start] --> B{Choice}`.
- ```dot — a graph best drawn by automatic layout (dependencies, call graphs). Body: Graphviz DOT, e.g. `digraph G { A -> B }`.
- ```geometry — an exact geometric construction. Body: JSON {"boundingbox":[xmin,ymax,xmax,ymin],"axis":true,"elements":[…]}. Each element is {"type":…,"args":[…],"attrs":{…}}. Coordinates must be NUMBERS. A text element takes its content as the LAST arg: {"type":"text","args":[1.5,-1.5,"a²=9"]}. Colours use fillColor/strokeColor/strokeWidth (plain fill/stroke are also accepted). Allowed types: point, line, segment, circle, ellipse, polygon, text, angle, arc, sector, midpoint, perpendicular, parallel, tangent, intersection, arrow, vector. Data only — never a function or code string (use ```plot for curves).
- ```map — a map. Body: JSON {"center":[lat,lng],"zoom":11,"markers":[{"lat":..,"lng":..,"label":".."}]} (optional "geojson").
- ```plot3d — a 3-D surface/scatter. For a formula, let the app compute it: {"zfunction":"x*y","x":[-5,5],"y":[-5,5],"layout":{"title":"…"}}. For explicit data use Plotly JSON {"data":[{"type":"surface","z":[[1,2],[3,4]]}]} — z must be NUMBERS; never write an expression such as [[x*y]] inside JSON.
- ```molecule — a chemical structure. Body: one SMILES string, e.g. CC(=O)Oc1ccccc1C(=O)O
- ```molecule3d — the same structure, rotatable in 3D. Body: standard XYZ format — atom count, a comment line, then one `Element x y z` line per atom (plain numbers, Å). No bonds needed; they're inferred from distance.
- ```calc — arithmetic, unit conversion, dates, matrices. Body: one expression per line, e.g. `5 km/h to m/s` or `(1250 * 1.19) / 3`. Never do multi-step arithmetic in prose — emit it here and the app computes it exactly.
- ```stats — descriptive statistics and linear regression. Body: `data: 4, 8, 15, 16, 23, 42` (optionally `x:` and `y:` lines for regression). The app computes mean/median/sd/quartiles/correlation — do not compute them yourself.
- ```solve — symbolic algebra. Body: one per line — `derivative: x^3 + 2x` , `simplify: (x^2-1)/(x-1)` , `solve: x^2 - 5x + 6 = 0` , `evaluate: ...`.
- ```table — a sortable, filterable data table with computed totals and CSV export. Body: JSON {"columns":["Item","Qty","Price"],"rows":[["A",2,9.99]],"total":["Price"]} — never add up a column yourself; list `total` and the app sums it.
- ```diff — compare two texts. Body: `--- before` / lines / `--- after` / lines. The app computes the real diff.
- ```regex — test a pattern. Body: `pattern: \d{4}-\d{2}` then `test:` lines. The app runs it and shows real matches.
- ```truth — truth table for a boolean expression, e.g. `(A and B) or not C`. The app enumerates every row.
- ```abc — music notation (see below). ```svg — only for a custom diagram none of the above can express.
Use ordinary ```language fences for code (they are syntax-highlighted). One block per figure; put explanation as prose outside the block.

=== GRAPHING A FUNCTION (y = f(x)) — ALWAYS use this, never hand-draw the curve ===
To graph a mathematical function, output a fenced block whose language tag is exactly `plot` (lowercase) containing ONLY the formula — the app evaluates it and draws the exact curve, so you must NOT compute points or draw the curve yourself in SVG.
```plot
y = x^2 + 3x - 2
```
Several curves and an explicit domain (default domain is -10..10 if you omit it):
```plot
f(x) = sin(x)
g(x) = cos(x)
x = -6.28 .. 6.28
```
Syntax: standard math — ^ for powers, * optional (3x is fine), functions sin cos tan asin acos atan sqrt exp log log10 abs floor ceil, constants pi and e. One function per line; set the range with a line like `x = a .. b`. This is far more reliable than drawing an SVG polyline by hand — use it for ANY formula.

=== CUSTOM SVG — only when no lane above fits ===
Every chart, diagram, plot, map and construction has a dedicated lane above; use it. Hand-drawn SVG is a last resort for a bespoke illustration, because coordinates you compute by hand are frequently wrong.
If you must: one ```svg block, root `<svg width="640" height="360" viewBox="0 0 640 360" xmlns="http://www.w3.org/2000/svg">`, a white background `<rect>`, and explicit non-white colours on every shape (the card is white in both themes). Allowed: svg g path line polyline polygon rect circle ellipse text tspan defs linearGradient radialGradient stop clipPath marker use title desc animateTransform animateMotion, styled with presentation attributes. The app genuinely animates: nest an `<animateTransform>` (type="translate"/"rotate"/"scale"/"skewX") or `<animateMotion>` (a `path` to follow) inside the shape it moves, e.g. `<animateTransform attributeName="transform" type="translate" values="0,0; 100,0; 0,0" dur="2s" repeatCount="indefinite"/>` — use this whenever asked to animate, move, or bounce something. Stripped, so never used: <style>, <script>, <foreignObject>, plain `<animate>` (silently dropped — use animateTransform/animateMotion instead), on* handlers, external URLs. <text> does not wrap and does not render LaTeX or Markdown — keep labels short, use Unicode superscripts (a², x²), and keep everything inside 0..640 x 0..360.

=== MUSIC / SHEET NOTATION ===
For music or sheet-music notation, do NOT draw notes as SVG. Output ABC notation in a fenced block whose language tag is exactly `abc` (lowercase) — the app renders it to a proper score. Emit valid ABC: an information header, then the tune body. Minimal template:
```abc
X:1
T:Tune title
M:4/4
L:1/4
K:C
C D E F | G A B c | c B A G | F E D C |]
```
Rules: X: (tune number), K: (key) are required; M: (metre) and L: (default note length) are strongly recommended, and every header line comes BEFORE the notes. Notes are letters A-G (c-b are the octave above; ',' lowers and ' raises an octave); digits set duration relative to L: (C2 = two units), z is a rest, | is a barline. Keep it to a few bars unless asked for more, use one ```abc block per piece, and put any explanation as prose outside the block."""

DEFAULT_CONFIG: dict = {
    # ── Primary Redis connection (the default endpoint) ───────────────────────
    "redis": {
        "host": "localhost", "port": 6379, "db": 0,
        "password": "", "ssl": False,
    },
    # ── Additional named Redis endpoints (list of connection configs) ─────────
    # Each item: {name, host, port, db, password, ssl}
    # RAG instances store which endpoint they live on in their metadata.
    "redis_endpoints": [],

    # ── Ollama (local LLM server) ─────────────────────────────────────────────
    "ollama": {"host": "http://localhost", "port": 11434, "model": ""},

    # ── Anthropic Claude API ──────────────────────────────────────────────────
    # api_key is never written to disk if it came from the ANTHROPIC_API_KEY env var.
    "claude": {
        "api_key": "",
        "model": "claude-sonnet-4-6",
        "base_url": "https://api.anthropic.com",
    },

    # ── OpenAI API ────────────────────────────────────────────────────────────
    # api_key is never written to disk if it came from the OPENAI_API_KEY env var.
    "openai": {
        "api_key": "",
        "model": "gpt-4o",
        "base_url": "https://api.openai.com",
    },

    # ── Qwen (Alibaba — OpenAI-compatible, free tier available) ──────────────
    # Get a free API key at qwen.ai · Free tier includes qwen-plus and qwen-turbo.
    "qwen": {
        "api_key": "",
        "model": "qwen-plus",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
    },

    # ── Mistral (OpenAI-compatible, EU-hosted, free "Experiment" tier) ────────
    # Get a free API key at console.mistral.ai · Set MISTRAL_API_KEY env var.
    "mistral": {
        "api_key": "",
        "model": "mistral-small-latest",
        "base_url": "https://api.mistral.ai/v1",
    },

    # ── Groq (OpenAI-compatible, generous free tier) ──────────────────────────
    # Get a free API key at console.groq.com · Fast inference, no credit card.
    "groq": {
        "api_key": "",
        "model": "llama-3.3-70b-versatile",
        "base_url": "https://api.groq.com/openai",
    },

    # ── Google Gemini (native google-genai SDK, free tier available) ─────────
    # Get a free API key at aistudio.google.com · Set GEMINI_API_KEY env var.
    "gemini": {
        "api_key": "",
        "model": "gemini-3-flash-preview",
    },

    # ── Active provider: "ollama" | "claude" | "openai" | "qwen" | "mistral" | "groq" | "gemini"
    "provider": "ollama",

    # ── Embedding model (SentenceTransformer) ─────────────────────────────────
    "embedding": {"model": "all-MiniLM-L6-v2", "max_image_dim": 1024},

    # ── RAG retrieval settings ────────────────────────────────────────────────
    "rag": {
        "chunk_size": 512,
        "chunk_overlap": 64,
        "top_k": 5,
        "similarity_threshold": 0.75,
        "hybrid_search": True,
    },

    # ── Semantic cache settings ───────────────────────────────────────────────
    "cache": {"enabled": True, "similarity_threshold": 0.92, "ttl": 3600},

    # ── UI preferences ────────────────────────────────────────────────────────
    "ui": {"theme": "auto", "show_rag_matches": False},

    "web_sources": [],
    # Global system instruction prepended to every chat turn (before any selected
    # template). Edited in Settings -> Templates -> Base Instruction.
    "base_instruction": DEFAULT_BASE_INSTRUCTION,
    "prompt_templates": [
        {"name": "Default",      "system": ""},
        {"name": "Redis Expert", "system": "You are a Redis expert. Answer concisely with examples."},
        {"name": "ELI5",         "system": "Explain everything like I'm 5 years old, using simple analogies."},
    ],
    "security": {"password": "", "enabled": False},
    "active_rag": "default",
    "sessions":   {"persist": True, "ttl": 86400},
    "reranker":   {"enabled": False, "model": "cross-encoder/ms-marco-MiniLM-L-6-v2", "top_n": 5},
    "hyde":       {"enabled": False},
    "recrawl":    {"enabled": False, "interval_minutes": 60},
    "scheduled_sources": [],
    # ── Web crawler settings ──────────────────────────────────────────────────
    "crawl": {
        "concurrency":    10,    # max parallel httpx fetches
        "js_render":      False, # force Playwright for all pages (requires crawl4ai)
        "js_concurrency": 3,     # max simultaneous Playwright browser tabs
        "smart_mode":     True,  # httpx-first; fall back to Playwright only when needed
        "min_words":      100,   # word threshold below which smart mode triggers JS fallback
    },
}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FASTAPI APP
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

app = FastAPI(title="RediRecall")

# CORS — same-origin by default. The bundled UI is served from this app, so it
# needs no cross-origin grant and none is given (a wildcard would let any site
# script the API from a user's browser). Set REDIRECALL_CORS_ORIGINS to a
# comma-separated list of trusted origins only if you host the UI separately.
_cors_origins = [
    o.strip() for o in os.environ.get("REDIRECALL_CORS_ORIGINS", "").split(",")
    if o.strip()
]
if _cors_origins:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=_cors_origins,
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

# Security headers on every response — defense in depth alongside the frontend's
# DOMPurify sanitisation. 'unsafe-inline' is unavoidable (the single-file UI uses
# an inline <script>, inline styles, and static inline handlers), but the policy
# still blocks external script injection, plugins, framing, and cross-origin
# data exfiltration. cdnjs is allowed for marked / KaTeX / DOMPurify.
_CSP = (
    "default-src 'self'; "
    # cdnjs: marked/DOMPurify/KaTeX/abcjs/math.js + the lazy-loaded renderers
    # (mermaid, Chart.js, highlight.js, Viz/Graphviz, JSXGraph, Leaflet, Plotly).
    # jsDelivr: smiles-drawer (molecules) — not published on cdnjs.
    "script-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com https://cdn.jsdelivr.net; "
    "style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
    # https: is needed for map tiles (e.g. OpenStreetMap) in the ```map lane.
    "img-src 'self' data: blob: https:; "
    "font-src 'self' https://cdnjs.cloudflare.com data:; "
    "connect-src 'self'; "
    "worker-src 'self' blob:; "
    "object-src 'none'; "
    "base-uri 'self'; "
    "frame-ancestors 'self'; "
    "form-action 'self'"
)


@app.middleware("http")
async def _security_headers(request, call_next):
    resp = await call_next(request)
    resp.headers.setdefault("Content-Security-Policy", _CSP)
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    resp.headers.setdefault("Referrer-Policy", "no-referrer")
    return resp

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# GLOBAL STATE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_config: dict = {}

# Redis client pool: "default" key → primary client; other keys → named endpoints
_redis_clients: dict[str, redis.Redis] = {}

# Runtime reachability of each configured Redis endpoint.
# Keys are endpoint names ("default", "RAG", …); value is True/False.
# Unknown endpoints are treated as reachable (optimistic default).
_endpoint_health: dict[str, bool] = {}

_embed_model: Optional[Any] = None  # SentenceTransformer, lazily loaded
_embed_model_name: str = ""

# In-memory session store: session_id → list of {role, content} dicts
_sessions: dict[str, list] = {}

_feedback: list = []
_ingestion_logs: list = []
_crawl_tasks: dict[str, asyncio.Task] = {}
# Active crawl state — survives browser refresh/reconnect.
# Key: url.  Value: {instance, pages_done, chunks, errors, blocked, start_ts, done}
_active_crawls: dict[str, dict] = {}

# Per-instance RAG query statistics (in-memory, reset on restart).
# Keys: instance name.  Values: counters used to derive hit-rate and avg score.
# Structure: {name: {queries, hits, chunks_total, score_sum}}
_rag_stats: dict[str, dict] = {}
_reranker: Optional["CrossEncoder"] = None
_reranker_model_name: str = ""
_recrawl_task: Optional[asyncio.Task] = None


def _record_rag_stats(
    instance: str,
    results: list[dict],
    raw_results: list[dict] | None = None,
) -> None:
    """Update per-instance RAG stats after every search_rag() call.

    ``results``     — chunks that passed the similarity threshold (used for hit counting).
    ``raw_results`` — all KNN results before threshold filtering; used to track the best
                      raw score so we can detect threshold misconfiguration (e.g. good
                      matches getting filtered out because the threshold is too strict).
    """
    s = _rag_stats.setdefault(instance, {
        "queries": 0, "hits": 0, "chunks_total": 0, "score_sum": 0.0,
        "raw_score_sum": 0.0,  # sum of best raw scores (pre-threshold) across all queries
    })
    s["queries"] += 1
    # Best raw score — top-1 from the unfiltered KNN results
    best_raw = raw_results[0]["score"] if raw_results else 0.0
    s["raw_score_sum"] += best_raw
    if results:
        s["hits"]         += 1
        s["chunks_total"] += len(results)
        s["score_sum"]    += results[0]["score"]   # top-1 cosine similarity (post-filter)
# Active streaming task per session — used to cancel mid-stream when the client
# sends {"type":"abort"}.  Keyed by session id.
_chat_tasks: dict[str, asyncio.Task] = {}

# API keys sourced from environment variables — never persisted to disk
_env_key: str = ""          # ANTHROPIC_API_KEY
_openai_env_key: str = ""   # OPENAI_API_KEY
_qwen_env_key: str = ""     # DASHSCOPE_API_KEY
_mistral_env_key: str = ""  # MISTRAL_API_KEY
_groq_env_key: str = ""     # GROQ_API_KEY
_gemini_env_key: str = ""   # GEMINI_API_KEY

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CONFIG HELPERS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def load_config() -> dict:
    """Load config.json and deep-merge with defaults so new keys always exist."""
    if CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH) as f:
                cfg = json.load(f)
            merged = {**DEFAULT_CONFIG, **cfg}
            # Shallow-merge each nested dict so we pick up new sub-keys
            for k, v in DEFAULT_CONFIG.items():
                if isinstance(v, dict):
                    merged[k] = {**v, **cfg.get(k, {})}
            return merged
        except Exception:
            pass
    return dict(DEFAULT_CONFIG)


def save_config(cfg: dict):
    """
    Persist config to disk.
    API keys that were loaded from environment variables are always stripped
    before writing so they never end up on disk.
    """
    to_save = copy.deepcopy(cfg)
    if _env_key and to_save.get("claude", {}).get("api_key") == _env_key:
        to_save.setdefault("claude", {})["api_key"] = ""
    if _openai_env_key and to_save.get("openai", {}).get("api_key") == _openai_env_key:
        to_save.setdefault("openai", {})["api_key"] = ""
    if _qwen_env_key and to_save.get("qwen", {}).get("api_key") == _qwen_env_key:
        to_save.setdefault("qwen", {})["api_key"] = ""
    if _mistral_env_key and to_save.get("mistral", {}).get("api_key") == _mistral_env_key:
        to_save.setdefault("mistral", {})["api_key"] = ""
    if _groq_env_key and to_save.get("groq", {}).get("api_key") == _groq_env_key:
        to_save.setdefault("groq", {})["api_key"] = ""
    if _gemini_env_key and to_save.get("gemini", {}).get("api_key") == _gemini_env_key:
        to_save.setdefault("gemini", {})["api_key"] = ""
    with open(CONFIG_PATH, "w") as f:
        json.dump(to_save, f, indent=2)


# ── Provider SDK client reuse ────────────────────────────────────────────────
# A fresh SDK client per request means a fresh connection pool and a fresh TLS
# handshake on every turn (~30 ms per cloud provider, measured) — and a first-turn
# conversation pays it three times (answer + title + HyDE). Cache by credentials so
# a key change still creates a new client.
_provider_clients: dict = {}

def _cached_client(kind: str, api_key: str, base_url: str):
    """Return a pooled SDK client for (kind, key, base_url), creating it on first use."""
    ck = (kind, api_key, base_url)
    c = _provider_clients.get(ck)
    if c is None:
        c = (_anthropic.AsyncAnthropic(api_key=api_key, base_url=base_url)
             if kind == "anthropic" else
             _AsyncOpenAI(api_key=api_key, base_url=base_url))
        _provider_clients[ck] = c
    return c

def invalidate_provider_clients():
    """Drop cached clients (called on config save, since keys may have changed)."""
    _provider_clients.clear()


def compose_system_prompt(client_system: str | None) -> str:
    """
    Build the effective system prompt for a chat turn.

    The global base_instruction (Settings -> Templates -> Base Instruction) is
    always prepended; a selected template's system prompt is added on top of it
    (templates are additive to the base). If both are empty we fall back to a
    plain default so the model still gets a system message.
    """
    base   = (_config.get("base_instruction") or "").strip()
    client = (client_system or "").strip()
    parts  = [p for p in (base, client) if p]
    return "\n\n".join(parts) if parts else "You are a helpful assistant."


# ── Secret redaction ──────────────────────────────────────────────────────────
# Config sent to the browser must never carry provider keys or passwords. Each
# stored secret is swapped for a sentinel on the way out and restored from the
# stored value on the way back in, so the UI can round-trip settings without
# ever seeing a secret — and a blank field still clears a key (blank != sentinel).

_SECRET_SENTINEL = "__REDIRECALL_SECRET_KEPT__"
_PROVIDER_SECRET_KEYS = ("claude", "openai", "qwen", "mistral", "groq", "gemini")


def _redact_secrets(cfg: dict) -> dict:
    """Deep copy of cfg with every set secret replaced by _SECRET_SENTINEL."""
    red = copy.deepcopy(cfg)
    for p in _PROVIDER_SECRET_KEYS:
        if isinstance(red.get(p), dict) and red[p].get("api_key"):
            red[p]["api_key"] = _SECRET_SENTINEL
    if isinstance(red.get("redis"), dict) and red["redis"].get("password"):
        red["redis"]["password"] = _SECRET_SENTINEL
    if isinstance(red.get("security"), dict) and red["security"].get("password"):
        red["security"]["password"] = _SECRET_SENTINEL
    for ep in red.get("redis_endpoints", []) or []:
        if isinstance(ep, dict) and ep.get("password"):
            ep["password"] = _SECRET_SENTINEL
    return red


def _unredact_secrets(new_cfg: dict, old_cfg: dict) -> None:
    """In place: swap any sentinel secret in new_cfg back to the stored value."""
    for p in _PROVIDER_SECRET_KEYS:
        if isinstance(new_cfg.get(p), dict) and new_cfg[p].get("api_key") == _SECRET_SENTINEL:
            new_cfg[p]["api_key"] = (old_cfg.get(p) or {}).get("api_key", "")
    if isinstance(new_cfg.get("redis"), dict) and new_cfg["redis"].get("password") == _SECRET_SENTINEL:
        new_cfg["redis"]["password"] = (old_cfg.get("redis") or {}).get("password", "")
    if isinstance(new_cfg.get("security"), dict) and new_cfg["security"].get("password") == _SECRET_SENTINEL:
        new_cfg["security"]["password"] = (old_cfg.get("security") or {}).get("password", "")
    old_eps = {e.get("name"): e for e in (old_cfg.get("redis_endpoints") or []) if isinstance(e, dict)}
    for ep in new_cfg.get("redis_endpoints", []) or []:
        if isinstance(ep, dict) and ep.get("password") == _SECRET_SENTINEL:
            ep["password"] = (old_eps.get(ep.get("name")) or {}).get("password", "")


def load_logs():
    """Load persisted ingestion log from disk into memory."""
    global _ingestion_logs
    if LOGS_PATH.exists():
        with open(LOGS_PATH) as f:
            _ingestion_logs = json.load(f)


def append_log(entry: dict):
    """Append an ingestion event and keep the last 500 entries on disk."""
    _ingestion_logs.append(entry)
    with open(LOGS_PATH, "w") as f:
        json.dump(_ingestion_logs[-500:], f)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# REDIS HELPERS — multi-endpoint aware
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _build_redis_client(cfg: dict) -> redis.Redis:
    """Build a redis.Redis client from a connection-config dict."""
    return redis.Redis(
        host=cfg.get("host") or "localhost",
        port=int(cfg.get("port") or 6379),    # or-default handles None (JSON null) safely
        db=int(cfg.get("db") or 0),
        password=cfg.get("password") or None,
        ssl=bool(cfg.get("ssl") or False),
        decode_responses=False,   # we handle bytes manually for embedding binary data
        socket_connect_timeout=5,  # fail fast if endpoint is unreachable
        socket_timeout=10,         # prevent blocking forever on slow endpoints
    )


def get_redis(cfg: dict | None = None) -> redis.Redis:
    """
    Return (and cache) the primary Redis client.
    Pass cfg to temporarily use a different connection config (e.g. for test).
    """
    c = cfg or _config.get("redis", {})
    client = _build_redis_client(c)
    _redis_clients["default"] = client
    return client


def r_for(endpoint_name: str = "default") -> redis.Redis:
    """
    Return the Redis client for a named endpoint.
    "default" maps to the primary redis config.
    Other names are looked up in config["redis_endpoints"].
    Clients are cached so we reuse connections.
    """
    if endpoint_name in _redis_clients:
        return _redis_clients[endpoint_name]

    if endpoint_name == "default":
        return get_redis()

    # Find the named endpoint in config
    for ep in _config.get("redis_endpoints", []):
        if ep.get("name") == endpoint_name:
            client = _build_redis_client(ep)
            _redis_clients[endpoint_name] = client
            return client

    # Fallback to default if not found
    log.warning(f"Redis endpoint '{endpoint_name}' not found — falling back to default")
    return r()


def r() -> redis.Redis:
    """Shorthand: return the default Redis client, creating it if needed."""
    if "default" not in _redis_clients:
        return get_redis()
    return _redis_clients["default"]


def invalidate_redis_clients():
    """
    Clear the client cache so they're rebuilt on next use.
    Also resets the SemanticCache so it reconnects with the new client.
    Called after config changes that affect Redis connections.
    """
    global _semantic_cache
    _redis_clients.clear()
    _semantic_cache = None


def _probe_endpoint(ep_name: str) -> bool:
    """
    Return True if the named Redis endpoint responds to PING within the
    socket_connect_timeout.  Updates _endpoint_health in place.
    Never raises.
    """
    global _endpoint_health
    try:
        rc = r_for(ep_name)
        rc.ping()
        _endpoint_health[ep_name] = True
        return True
    except Exception:
        _endpoint_health[ep_name] = False
        log.warning(f"Redis endpoint '{ep_name}' is unreachable — RAG instances on it will be marked offline")
        return False


def refresh_endpoint_health() -> dict[str, bool]:
    """
    Probe all configured Redis endpoints (default + extras) and return the
    updated health dict.  Called at startup and periodically in background.
    """
    _probe_endpoint("default")
    for ep in _config.get("redis_endpoints", []):
        ep_name = ep.get("name", "")
        if ep_name and ep_name != "default":
            _probe_endpoint(ep_name)
    return dict(_endpoint_health)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# EMBEDDING HELPERS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def get_embed_model(name: str | None = None) -> Any:
    """
    Lazy-load the SentenceTransformer model.
    The sentence_transformers import is deferred here so Python startup is
    not blocked by the 5-15 s PyTorch initialisation time.
    Reloads if the model name has changed (e.g. user switched in settings).
    """
    global _embed_model, _embed_model_name
    name = name or _config.get("embedding", {}).get("model", "all-MiniLM-L6-v2")
    if _embed_model is None or _embed_model_name != name:
        log.info(f"Loading embedding model: {name}")
        from sentence_transformers import SentenceTransformer  # noqa: PLC0415
        _embed_model = SentenceTransformer(name)
        _embed_model_name = name
    return _embed_model


def embed(text: str) -> np.ndarray:
    """Embed a single string. Returns a normalised float32 vector."""
    return get_embed_model().encode(text, normalize_embeddings=True)


def embed_batch(texts: list[str]) -> np.ndarray:
    """
    Embed a list of strings in one model call.
    Much faster than calling embed() in a loop because the model can use
    batched GPU/CPU operations.  Returns shape (N, dim) float32 array.
    """
    return get_embed_model().encode(
        texts,
        normalize_embeddings=True,
        batch_size=32,          # tune based on available RAM/VRAM
        show_progress_bar=False,
    )


def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    """Cosine similarity of two pre-normalised vectors (just a dot product)."""
    return float(np.dot(a, b))

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CROSS-ENCODER RERANKER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def get_reranker() -> "CrossEncoder | None":
    """Lazy-load the cross-encoder reranker model if enabled in config."""
    global _reranker, _reranker_model_name
    if not _config.get("reranker", {}).get("enabled", False):
        return None
    if not HAS_CROSSENCODER:
        log.warning("cross-encoder reranking requested but sentence-transformers CrossEncoder not available")
        return None
    model = _config.get("reranker", {}).get("model", "cross-encoder/ms-marco-MiniLM-L-6-v2")
    if _reranker is None or _reranker_model_name != model:
        log.info(f"Loading cross-encoder reranker: {model}")
        try:
            _reranker = CrossEncoder(model)
            _reranker_model_name = model
        except Exception as e:
            log.warning(f"Cross-encoder load failed: {e}")
            return None
    return _reranker


def rerank_chunks(query: str, chunks: list[dict], top_n: int = 5) -> list[dict]:
    """
    Re-rank retrieved chunks using a cross-encoder model.

    Cross-encoders jointly encode (query, passage) pairs and produce a
    relevance score that is much more accurate than bi-encoder cosine
    similarity.  This step runs after the fast HNSW+BM25 retrieval and
    re-orders the candidates so the most relevant chunks bubble to the top.

    Falls back to the original order if the reranker is unavailable.
    """
    reranker = get_reranker()
    if not reranker or not chunks:
        return chunks
    try:
        pairs  = [(query, c["text"]) for c in chunks]
        scores = reranker.predict(pairs)
        ranked = sorted(zip(scores, chunks), key=lambda x: x[0], reverse=True)
        result = [c for _, c in ranked[:top_n]]
        # Annotate each chunk with its reranker score for debugging/analytics
        for (score, _), chunk in zip(sorted(zip(scores, chunks), key=lambda x: x[0], reverse=True), result):
            chunk["reranker_score"] = round(float(score), 4)
        return result
    except Exception as e:
        log.warning(f"Reranking failed, using original order: {e}")
        return chunks

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SESSION PERSISTENCE (Redis-backed)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_SESSION_PREFIX = "session:"


def _session_key(sid: str) -> str:
    return f"{_SESSION_PREFIX}{sid}"


def load_session(sid: str) -> list:
    """Load a session's message list from Redis. Returns [] if not found or persistence disabled."""
    if not _config.get("sessions", {}).get("persist", True):
        return []
    try:
        raw = r().get(_session_key(sid))
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    return []


def save_session(sid: str, messages: list):
    """Persist a session to Redis with the configured TTL."""
    if not _config.get("sessions", {}).get("persist", True):
        return
    try:
        ttl = int(_config.get("sessions", {}).get("ttl", 86400))
        r().setex(_session_key(sid), ttl, json.dumps(messages))
    except Exception as e:
        log.warning(f"Session save failed for {sid}: {e}")


def delete_session_from_redis(sid: str):
    """Remove a session from Redis."""
    try:
        r().delete(_session_key(sid))
    except Exception:
        pass


def list_sessions_from_redis() -> list[dict]:
    """Return session summaries from Redis, excluding sessions already in _sessions."""
    result = []
    try:
        client = r()
        # Collect keys first, then fetch all values in a single pipeline.
        keys = []
        sids = []
        for k in client.scan_iter(f"{_SESSION_PREFIX}*", count=200):
            sid = k.decode().removeprefix(_SESSION_PREFIX)
            if sid in _sessions:
                continue   # already handled by the in-memory dict
            keys.append(k)
            sids.append(sid)

        if not keys:
            return result

        pipe = client.pipeline(transaction=False)
        for k in keys:
            pipe.get(k)
        values = pipe.execute()   # one round-trip for all GETs

        for sid, raw in zip(sids, values):
            if not raw:
                continue
            msgs = json.loads(raw)
            preview = ""
            for m in reversed(msgs):
                if m.get("role") == "assistant":
                    preview = m.get("content", "")[:60]
                    break
            result.append({"id": sid, "messages": len(msgs), "preview": preview})
    except Exception:
        pass
    return result

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# RAG INDEX HELPERS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def rag_prefix(instance: str) -> str:
    """Key namespace for a RAG instance, e.g. 'rag:default'."""
    return f"rag:{instance}"


def _get_rag_index(instance: str, rc: redis.Redis) -> SearchIndex:
    """
    Build a redisvl SearchIndex for a RAG instance without creating it in Redis.
    The index schema is declared once here; use .create() / .delete() / .query()
    on the returned object to interact with Redis.

    HNSW parameters:
        M=16            — neighbours per node (higher = better recall, more RAM)
        EF_CONSTRUCTION=200 — build-time beam width (higher = better quality)
    """
    dim    = get_embed_model().get_sentence_embedding_dimension()
    prefix = rag_prefix(instance)
    schema = IndexSchema.from_dict({
        "index": {
            "name":         f"{prefix}:idx",
            "prefix":       f"{prefix}:chunk:",
            "storage_type": "hash",
        },
        "fields": [
            {"name": "text",      "type": "text"},
            {"name": "source",    "type": "text"},
            {"name": "chunk_id",  "type": "numeric"},
            {"name": "embedding", "type": "vector",  "attrs": {
                "algorithm":       "hnsw",
                "datatype":        "float32",
                "dims":            dim,
                "distance_metric": "cosine",
                "m":               16,
                "ef_construction": 200,
            }},
        ],
    })
    return SearchIndex(schema, redis_client=rc)


_index_ensured: set[str] = set()

def ensure_rag_index(instance: str, rc: redis.Redis | None = None):
    """
    Create the RediSearch HNSW vector index for a RAG instance if it doesn't exist.
    Tracks confirmed instances in-process to avoid redundant Redis round-trips and
    suppress the repeated 'Non-Sortable non-Indexable' notices from Redis.
    """
    if instance in _index_ensured:
        return
    rc = rc or r()
    try:
        _get_rag_index(instance, rc).create(overwrite=False)
        log.info(f"RAG index ensured for '{instance}'")
    except Exception as e:
        msg = str(e)
        # "Index already exists" is not an error — the index is present, which is what we want.
        if "already exists" not in msg.lower():
            log.warning(f"RAG index creation for '{instance}': {e}")
    _index_ensured.add(instance)


def next_chunk_id(instance: str, reserve: int = 1, rc: redis.Redis | None = None) -> int:
    """
    Atomically reserve `reserve` sequential chunk IDs using a Redis INCRBY counter.

    This is O(1) — a single Redis round-trip — whereas the old approach of
    scanning all chunk keys was O(N) and very slow for large instances.

    Returns the first ID in the reserved range, so caller can use
    [start, start+1, ..., start+reserve-1].
    """
    rc = rc or r()
    counter_key = f"rag:{instance}:chunk_counter"
    new_val = int(rc.incrby(counter_key, reserve))
    return new_val - reserve   # return the start of the reserved range


def add_chunks(instance: str, chunks: list[dict], rc: redis.Redis | None = None):
    """
    Store a batch of text chunks in Redis with their embeddings.

    Optimisations:
    - All texts are embedded in ONE model call (batch inference).
    - All Redis writes are pipelined (one round-trip for the whole batch).
    """
    if not chunks:
        return

    rc = rc or r()
    ensure_rag_index(instance, rc)
    prefix = rag_prefix(instance)

    # Batch embed — one model call for the entire chunk list
    texts = [ch["text"] for ch in chunks]
    embeddings = embed_batch(texts)   # shape: (N, dim)

    pipe = rc.pipeline(transaction=False)
    for ch, emb in zip(chunks, embeddings):
        key = f"{prefix}:chunk:{ch['id']}"
        pipe.hset(key, mapping={
            "text":      ch["text"].encode(),
            "source":    ch.get("source", "").encode(),
            "chunk_id":  str(ch["id"]),
            "embedding": emb.astype(np.float32).tobytes(),
        })
    pipe.execute()


def _decode(v) -> str:
    """Decode a bytes value to str; pass strings through unchanged."""
    return v.decode() if isinstance(v, bytes) else (v or "")


def search_rag(
    instance: str,
    query: str,
    top_k: int = 5,
    threshold: float = 0.0,
    rc: redis.Redis | None = None,
    hybrid: bool = True,
    query_vec: "np.ndarray | None" = None,
    source_filter: str = "",
) -> list[dict]:
    """
    Search a RAG instance and return the top-K most relevant chunks.

    When ``hybrid=True`` (default) the search combines two strategies via
    Reciprocal Rank Fusion (RRF):

      1. **Vector KNN** — semantic similarity via redisvl VectorQuery (HNSW cosine).
         Catches paraphrases and related concepts.
      2. **BM25 full-text** — exact/near-exact keyword matching.
         Catches precise terms that may score below the cosine threshold.

    RRF formula: each result gets 1/(K+rank) for every list it appears in.
    K=60 is the standard constant that prevents high ranks from dominating.
    Results are then re-ranked by combined RRF score.  The final cosine score
    (from the KNN result, or 0 if the chunk only appeared in BM25) is used
    for threshold filtering and analytics.

    When ``hybrid=False`` only the vector search is performed.
    """
    rc     = rc or r()
    prefix = rag_prefix(instance)
    idx_name = f"{prefix}:idx"
    # Use a pre-computed query vector (e.g. from HyDE) if provided, otherwise embed the query.
    q_emb = query_vec if query_vec is not None else embed(query).astype(np.float32)

    try:
        # ── 1. Vector KNN search via redisvl VectorQuery ──────────────────────
        # Fetch top_k*2 so that after RRF merging we still have enough candidates.
        fetch_k = top_k * 2 if hybrid else top_k
        vq = VectorQuery(
            vector=q_emb.tolist(),
            vector_field_name="embedding",
            return_fields=["text", "source"],
            num_results=fetch_k,
        )
        idx = _get_rag_index(instance, rc)
        raw_vec = idx.query(vq)   # list[dict]: id, text, source, vector_distance

        # vector_distance is cosine DISTANCE (0=identical); convert to similarity
        vec_rows: list[dict] = []
        for row in raw_vec:
            vec_rows.append({
                "_key":       row.get("id", ""),
                "text":       _decode(row.get("text", "")),
                "source":     _decode(row.get("source", "")),
                "_vec_score": round(1.0 - float(row.get("vector_distance", 1.0)), 4),
            })

        # ── 2. BM25 full-text search (hybrid mode only) ───────────────────────
        bm25_rows: list[dict] = []
        if hybrid:
            kws = _keywords_for_bm25(query)
            if kws:
                text_q = " | ".join(kws)
                try:
                    txt_res = rc.execute_command(
                        "FT.SEARCH", idx_name,
                        f"@text:({text_q})",
                        "RETURN", "2", "text", "source",
                        "LIMIT", "0", str(fetch_k),
                        "DIALECT", "2",
                    )
                    # Parse raw FT.SEARCH response (key, [field, val, ...], ...)
                    items = txt_res[1:]
                    for i in range(0, len(items), 2):
                        key = _decode(items[i])
                        fields = items[i + 1]
                        d: dict = {"_key": key, "_vec_score": 0.0}
                        for j in range(0, len(fields), 2):
                            d[_decode(fields[j])] = _decode(fields[j + 1])
                        bm25_rows.append(d)
                except Exception:
                    pass   # text index unavailable; degrade to vector-only

        # ── 3. RRF merge and re-rank ─────────────────────────────────────────
        K_RRF  = 60
        scores: dict[str, dict] = {}
        for rank, row in enumerate(vec_rows):
            key = row["_key"]
            scores.setdefault(key, {"row": row, "rrf": 0.0})
            scores[key]["rrf"] += 1.0 / (K_RRF + rank + 1)

        for rank, row in enumerate(bm25_rows):
            key = row["_key"]
            if key in scores:
                scores[key]["rrf"] += 1.0 / (K_RRF + rank + 1)
            else:
                scores[key] = {"row": row, "rrf": 1.0 / (K_RRF + rank + 1)}

        ranked = sorted(scores.values(), key=lambda x: x["rrf"], reverse=True)[:top_k]

        raw_results = [
            {
                "text":   item["row"].get("text", ""),
                "source": item["row"].get("source", ""),
                "score":  float(item["row"].get("_vec_score", 0.0)),
            }
            for item in ranked
        ]

        # Threshold is applied to the cosine vector score
        results = [c for c in raw_results if c["score"] >= threshold]
        # Optional source filter (substring match on source field)
        if source_filter:
            results = [c for c in results if source_filter.lower() in c.get("source", "").lower()]
        _record_rag_stats(instance, results, raw_results)
        return results

    except Exception as e:
        # Auto-recover: if the FT index was dropped (e.g. Redis restart without
        # RDB persistence), recreate it and retry the query once — no recursion.
        if "no such index" in str(e).lower():
            log.warning(f"RAG index missing for '{instance}', recreating and retrying…")
            _index_ensured.discard(instance)
            ensure_rag_index(instance, rc)
            try:
                idx = _get_rag_index(instance, rc)
                fetch_k = top_k * 2 if hybrid else top_k
                q_emb2 = query_vec if query_vec is not None else embed(query).astype(np.float32)
                vq2 = VectorQuery(
                    vector=q_emb2.tolist(),
                    vector_field_name="embedding",
                    return_fields=["text", "source"],
                    num_results=fetch_k,
                )
                raw2 = idx.query(vq2)
                results2 = [
                    {
                        "text":   _decode(row.get("text", "")),
                        "source": _decode(row.get("source", "")),
                        "score":  round(1.0 - float(row.get("vector_distance", 1.0)), 4),
                    }
                    for row in raw2
                ]
                results2 = [c for c in results2 if c["score"] >= threshold]
                if source_filter:
                    results2 = [c for c in results2 if source_filter.lower() in c.get("source", "").lower()]
                _record_rag_stats(instance, results2, results2)
                return results2
            except Exception as e2:
                log.warning(f"RAG search still failing for '{instance}' after index recreate: {e2}")
        else:
            log.warning(f"RAG search skipped for '{instance}': {e}")
        _record_rag_stats(instance, [], [])
        return []


async def search_rag_parallel(
    instances: list[str],
    query: str,
    top_k: int = 5,
    threshold: float = 0.0,
    hybrid: bool = True,
    query_vec: "np.ndarray | None" = None,
) -> list[dict]:
    """
    Search multiple RAG instances concurrently and return a merged, score-sorted list.

    Each instance may live on a different Redis server (resolved via rc_for_instance).
    Only enabled instances are queried.  Results from all instances are merged and
    re-ranked by similarity score so the top_k best chunks are returned regardless
    of which instance they came from.

    The `instance` key is added to each chunk so the caller knows the origin.
    """
    loop = asyncio.get_event_loop()

    # Filter to enabled + reachable instances only
    enabled: list[str] = []
    for inst in instances:
        try:
            meta, ep = await _rag_meta_cached_async(inst)   # cached; miss resolves off-loop
            enabled_flag = True
            if meta:
                enabled_flag = meta.get("enabled", True)
                # Skip if the owning endpoint is known to be offline
                if not _endpoint_health.get(ep, True):
                    continue
            if enabled_flag:
                enabled.append(inst)
        except Exception:
            pass  # skip unreachable instances gracefully

    if not enabled:
        return []

    async def _search_one(inst: str) -> list[dict]:
        """Run a single synchronous search in a thread-pool so searches are parallel."""
        rc = rc_for_instance(inst)
        results = await loop.run_in_executor(
            None, search_rag, inst, query, top_k, threshold, rc, hybrid, query_vec
        )
        # Tag each chunk with its origin instance
        for c in results:
            c["instance"] = inst
        return results

    # Fan out to all enabled instances simultaneously
    per_instance = await asyncio.gather(*[_search_one(i) for i in enabled])

    # Merge, sort by descending score, keep global top_k
    merged: list[dict] = []
    for results in per_instance:
        merged.extend(results)
    merged.sort(key=lambda c: c["score"], reverse=True)
    return merged[:top_k]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SEMANTIC CACHE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _make_cache_vectorizer() -> Any:
    """
    Build an HFTextVectorizer that uses the same sentence-transformer model
    as the RAG embedder.  redisvl uses this to embed queries and responses
    before storing/looking up cache entries.

    The short model name (e.g. "all-MiniLM-L6-v2") is expanded to its full
    HuggingFace path so redisvl can resolve it correctly.
    """
    from redisvl.utils.vectorize import HFTextVectorizer  # noqa: PLC0415
    model = _config.get("embedding", {}).get("model", "all-MiniLM-L6-v2")
    if "/" not in model:
        model = f"sentence-transformers/{model}"
    return HFTextVectorizer(model=model)


def _get_semantic_cache() -> "SemanticCache | None":
    """
    Lazily create and return the shared SemanticCache.

    Uses redisvl's SemanticCache which stores prompt→response pairs as
    vector embeddings and performs approximate-nearest-neighbour lookup on
    every incoming query.  A query is a cache hit when its cosine distance
    to a stored prompt is ≤ distance_threshold (i.e. similarity ≥ threshold).

    Returns None if the cache is disabled in config or if the Redis Search
    module is not available (logged once, never repeated).
    """
    global _semantic_cache
    if not _config.get("cache", {}).get("enabled", True):
        return None
    if _semantic_cache is not None:
        return _semantic_cache
    try:
        similarity_threshold = _config.get("cache", {}).get("similarity_threshold", 0.92)
        ttl  = _config.get("cache", {}).get("ttl", 3600)
        # SemanticCache uses cosine DISTANCE not SIMILARITY — convert
        _semantic_cache = SemanticCache(
            name=CACHE_PREFIX.rstrip(":"),   # "semcache"
            vectorizer=_make_cache_vectorizer(),
            distance_threshold=round(1.0 - similarity_threshold, 4),
            ttl=ttl,
            redis_client=r(),
        )
        log.info("SemanticCache initialised (redisvl)")
    except Exception as e:
        if "unknown command" in str(e).lower():
            log.warning(
                "Semantic cache disabled — Redis Search module not available. "
                "Use Redis Stack or Redis Enterprise with the Search module enabled."
            )
        else:
            log.warning(f"SemanticCache init failed: {e}")
        _semantic_cache = None
    return _semantic_cache


_VISUAL_INTENT_RE = re.compile(
    r"\b(charts?|graphs?|graphing|plots?|plotting|diagrams?|histograms?|scatter|"
    r"bar\s*(?:graph|chart)|pie\s*chart|line\s*(?:graph|chart)|"
    r"visuali[sz]e|visuali[sz]ations?)\b"
    r"|\by\s*=\s*\S"        # "y = <expr>" — a function to plot (e.g. "render function y=3x+sin(x)*x")
    r"|\bf\s*\(\s*x\s*\)",  # f(x)
    re.I,
)


def wants_visual(query: str) -> bool:
    """
    True when the query asks for a chart/graph/plot.

    The semantic cache keys only on query text, so without this a chart request
    that is similar to an earlier text answer would return that text (no chart),
    and two similar chart requests would return the first one's SVG. We skip the
    cache (lookup and store) for visual-intent queries so every chart is fresh.
    """
    return bool(_VISUAL_INTENT_RE.search(query or ""))


def cache_lookup(query: str, threshold: float = 0.92) -> dict | None:
    """
    Look up the nearest cached response via redisvl SemanticCache.
    Returns {"response": str, "score": float} or None.
    """
    # During the first few seconds after a restart the background warm may not have built
    # the vectorizer yet; treat that as a cache miss rather than paying the ~3 s build here.
    if _semantic_cache is None and not _semantic_cache_ready:
        return None
    cache = _get_semantic_cache()
    if cache is None:
        return None
    try:
        hits = cache.check(prompt=query, num_results=1)
        if hits:
            h = hits[0]
            dist  = float(h.get("vector_distance", h.get("score", 1.0)))
            score = round(1.0 - dist, 4)
            if score >= threshold:
                try:
                    meta = h.get("metadata") or {}
                    cached_chunks = json.loads(meta.get("chunks_json", "[]"))
                except Exception:
                    cached_chunks = []
                return {"response": h.get("response", ""), "score": score, "entry_id": h.get("entry_id", ""), "chunks": cached_chunks}
    except Exception as e:
        log.error(f"Cache lookup error: {e}")
    return None


def cache_store(query: str, response: str, chunks: list | None = None):
    """Store a query→response pair in the SemanticCache with TTL.
    Chunks are stored as JSON metadata so they can be re-displayed on cache hits.
    """
    # Skip storing during the warm window rather than triggering the ~3 s inline build.
    if _semantic_cache is None and not _semantic_cache_ready:
        return
    cache = _get_semantic_cache()
    if cache is None:
        return
    try:
        metadata = {"chunks_json": json.dumps(chunks)} if chunks else None
        cache.store(prompt=query, response=response, metadata=metadata)
    except Exception as e:
        log.error(f"Cache store error: {e}")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# TEXT CHUNKING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Sentence boundary: period / ! / ? followed by whitespace (lookbehind so
# the punctuation stays attached to the preceding sentence).
_SENT_END = re.compile(r'(?<=[.!?])\s+')

# English stop-words excluded from BM25 keyword extraction
_STOPWORDS = frozenset({
    "a","an","the","is","it","in","on","at","to","of","for","and","or",
    "but","not","with","from","by","as","be","are","was","were","been",
    "has","have","had","do","does","did","will","would","can","could",
    "should","may","might","this","that","these","those","i","you","he",
    "she","we","they","my","your","his","her","our","their","what","which",
    "who","how","when","where","why","all","any","each","if","then","than",
    "so","no","yes","also","just","about","into","up","out","more","some",
    "use","using","like","get","set","see","make","know","want","need",
})


def _keywords_for_bm25(query: str) -> list[str]:
    """Extract significant words for BM25/text search — remove stop-words and short tokens."""
    return [w for w in re.findall(r'\w+', query.lower())
            if len(w) > 2 and w not in _STOPWORDS]


def chunk_text(text: str, size: int = 512, overlap: int = 64) -> list[str]:
    """
    Split text into overlapping chunks that respect sentence boundaries.

    Unlike naive word-count splitting, this first splits on sentence endings
    (./?/!) so no sentence is ever cut in the middle.  Sentences are then
    grouped into windows of approximately `size` words.  Overlap is achieved
    by carrying the last N words worth of sentences forward into the next chunk.

    size    — target words per chunk (approximate)
    overlap — words of sentence-level context shared between adjacent chunks
    """
    sentences = [s.strip() for s in _SENT_END.split(text) if s.strip()]
    if not sentences:
        # Fallback: no sentence punctuation found — split on whitespace
        words = text.split()
        result, i = [], 0
        while i < len(words):
            chunk = " ".join(words[i : i + size])
            if chunk.strip():
                result.append(chunk)
            i += max(1, size - overlap)
        return result

    chunks: list[str] = []
    i = 0
    while i < len(sentences):
        window: list[str] = []
        wc = 0
        j = i
        # Accumulate sentences until we reach the word target
        while j < len(sentences):
            s_words = len(sentences[j].split())
            if wc > 0 and wc + s_words > size:
                break
            window.append(sentences[j])
            wc += s_words
            j += 1
        # Safety: if a single sentence exceeds `size`, include it anyway
        if not window:
            window.append(sentences[i])
            j = i + 1

        chunks.append(" ".join(window))

        if overlap == 0 or j >= len(sentences):
            i = j
        else:
            # Walk backwards from j to find enough sentences to cover `overlap` words
            carried_wc = 0
            new_i = j
            for k in range(j - 1, i, -1):
                carried_wc += len(sentences[k].split())
                if carried_wc >= overlap:
                    new_i = k
                    break
            i = max(new_i, i + 1)   # always advance at least one sentence

    return chunks

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT INGESTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def ingest_text(
    instance: str,
    text: str,
    source: str,
    rc: redis.Redis | None = None,
) -> int:
    """
    Chunk raw text and store all chunks in the RAG index.
    Returns the number of chunks ingested.

    IDs are reserved atomically with next_chunk_id so concurrent ingests
    don't collide even without a distributed lock.
    """
    cfg = _config.get("rag", {})
    chunks = chunk_text(
        text,
        cfg.get("chunk_size", 512),
        cfg.get("chunk_overlap", 64),
    )
    if not chunks:
        return 0

    rc = rc or r()
    hash_set_key = f"rag:{instance}:chunk_hashes"

    # Deduplicate: compute all hashes, then check + add in a single pipeline.
    # This turns N round-trips into one network operation.
    chunk_hashes = [
        hashlib.sha256(" ".join(c.lower().split()).encode()).hexdigest()
        for c in chunks
    ]
    pipe = rc.pipeline(transaction=False)
    for h in chunk_hashes:
        pipe.sadd(hash_set_key, h)
    results = pipe.execute()          # one round-trip; returns [1 or 0, ...]
    new_chunks: list[str] = [c for c, added in zip(chunks, results) if added]

    if not new_chunks:
        return 0

    # Reserve N IDs at once — single O(1) Redis call
    start_id = next_chunk_id(instance, len(new_chunks), rc)
    records = [
        {"id": start_id + i, "text": c, "source": source}
        for i, c in enumerate(new_chunks)
    ]
    add_chunks(instance, records, rc)
    return len(records)


def _prepare_chunks(
    instance: str,
    text: str,
    source: str,
    rc: redis.Redis,
    force_reindex: bool = False,
) -> list[dict]:
    """
    Chunk, deduplicate, and reserve Redis IDs — but do NOT embed yet.
    Returns a list of {id, text, source} records ready to be embedded in a batch.
    Designed to be called via asyncio.to_thread(); it is CPU-light (no model calls).
    When force_reindex=True the hash-dedup filter is bypassed so previously seen
    chunks are re-embedded rather than silently dropped.
    """
    cfg = _config.get("rag", {})
    chunks = chunk_text(
        text,
        cfg.get("chunk_size", 512),
        cfg.get("chunk_overlap", 64),
    )
    if not chunks:
        return []

    hash_set_key = f"rag:{instance}:chunk_hashes"
    chunk_hashes = [
        hashlib.sha256(" ".join(c.lower().split()).encode()).hexdigest()
        for c in chunks
    ]
    pipe = rc.pipeline(transaction=False)
    for h in chunk_hashes:
        pipe.sadd(hash_set_key, h)
    results = pipe.execute()

    if force_reindex:
        # Allow all chunks through; hashes were already updated above
        new_chunks = chunks
    else:
        new_chunks = [c for c, added in zip(chunks, results) if added]

    if not new_chunks:
        return []

    start_id = next_chunk_id(instance, len(new_chunks), rc)
    return [{"id": start_id + i, "text": c, "source": source} for i, c in enumerate(new_chunks)]


_CHAT_FILE_ACCEPT = {".txt", ".md", ".csv", ".pdf", ".doc", ".docx", ".xls", ".xlsx"}
_CHAT_FILE_MAX_BYTES = 10 * 1024 * 1024   # 10 MB
_CHAT_FILE_MAX_CHARS = 150_000             # ~100 k tokens — keeps context manageable


def extract_file_text(filename: str, data: bytes) -> str:
    """Convert an uploaded document to plain text.

    Supported formats: .txt, .md, .csv, .pdf, .doc/.docx, .xls/.xlsx
    Raises ValueError for unsupported or unparseable files.
    """
    import io as _io

    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        return data.decode("utf-8", errors="ignore")

    if suffix == ".csv":
        text_io = _io.StringIO(data.decode("utf-8", errors="ignore"))
        rows = [" | ".join(row) for row in csv.reader(text_io)]
        return "\n".join(rows)

    if suffix == ".pdf":
        if not HAS_PYMUPDF:
            raise ValueError("PDF support requires PyMuPDF (pip install pymupdf)")
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)

    if suffix in (".doc", ".docx"):
        if not HAS_PYTHON_DOCX:
            raise ValueError("DOCX support requires python-docx (pip install python-docx)")
        doc = _DocxDocument(_io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells if cell.text.strip()))
        return "\n".join(parts)

    if suffix in (".xls", ".xlsx"):
        if not HAS_OPENPYXL:
            raise ValueError("Excel support requires openpyxl (pip install openpyxl)")
        wb = _load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {suffix}. Accepted: {', '.join(sorted(_CHAT_FILE_ACCEPT))}")


async def ingest_file(
    instance: str,
    path: Path,
    source: str,
    rc: redis.Redis | None = None,
) -> dict:
    """
    Ingest a single file (TXT, CSV, or PDF) into a RAG instance.
    Returns a log entry dict with status information.
    """
    text = ""
    suffix = path.suffix.lower()

    # File parsing (a big PDF is seconds of CPU) and the embed+Redis write in ingest_text
    # are both synchronous — run them off the event loop so a large upload never freezes
    # concurrent chat sessions or the WS receive loop.
    def _extract_txt():
        return path.read_text(errors="ignore")

    def _extract_csv():
        rows = []
        with open(path, newline="", errors="ignore") as f:
            for row in csv.reader(f):
                rows.append(" | ".join(row))
        return "\n".join(rows)

    def _extract_pdf():
        doc = fitz.open(str(path))
        return "\n".join(p.get_text() for p in doc)

    try:
        if suffix == ".txt":
            text = await asyncio.to_thread(_extract_txt)

        elif suffix == ".csv":
            text = await asyncio.to_thread(_extract_csv)

        elif suffix == ".pdf":
            if HAS_PYMUPDF:
                text = await asyncio.to_thread(_extract_pdf)
            else:
                return {"source": source, "status": "error", "error": "PyMuPDF not installed (pip install pymupdf)"}

        else:
            return {"source": source, "status": "skipped", "error": f"Unsupported type: {suffix}"}

        n = await asyncio.to_thread(ingest_text, instance, text, source, rc)
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "instance": instance,
            "source": source,
            "chunks": n,
            "status": "ok",
        }
        append_log(entry)
        return entry

    except Exception as e:
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "instance": instance,
            "source": source,
            "chunks": 0,
            "status": "error",
            "error": str(e),
        }
        append_log(entry)
        return entry

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# WEB CRAWLING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Shared httpx client — connection pooling + keep-alive reused across all crawls.
# Created lazily on first use; never closed (lives for the server process lifetime).
_shared_http_client: httpx.AsyncClient | None = None


async def _get_http_client() -> httpx.AsyncClient:
    """Return (or lazily create) the shared pooled httpx AsyncClient.

    Redirects are followed manually in fetch_url (not by httpx) so the SSRF
    guard can re-validate every hop — a page on a public host could otherwise
    302 the crawler onto an internal address.
    """
    global _shared_http_client
    if _shared_http_client is None or _shared_http_client.is_closed:
        _shared_http_client = httpx.AsyncClient(
            follow_redirects=False,
            timeout=httpx.Timeout(connect=10.0, read=30.0, write=10.0, pool=5.0),
            limits=httpx.Limits(max_connections=100, max_keepalive_connections=30),
            headers={"User-Agent": "RediRecallBot/1.0"},
        )
    return _shared_http_client


def _ip_is_public(ip_str: str) -> bool:
    """True only for globally-routable unicast addresses."""
    try:
        ip = ipaddress.ip_address(ip_str)
    except ValueError:
        return False
    return not (
        ip.is_private or ip.is_loopback or ip.is_link_local
        or ip.is_multicast or ip.is_reserved or ip.is_unspecified
    )


def assert_public_url(url: str) -> None:
    """SSRF guard: raise ValueError unless ``url`` is a public http(s) target.

    Rejects non-http(s) schemes and resolves the host, requiring *every* A/AAAA
    record to be a public address — so a hostname that resolves to a private,
    loopback, link-local, or reserved IP (e.g. 127.0.0.1, 169.254.169.254,
    10.x, cloud metadata endpoints) is blocked before any connection is made.
    """
    p = urlparse(url)
    if p.scheme not in ("http", "https"):
        raise ValueError(f"blocked non-http(s) URL scheme: {p.scheme or '(none)'}")
    host = p.hostname
    if not host:
        raise ValueError("blocked URL with no host")
    port = p.port or (443 if p.scheme == "https" else 80)
    try:
        infos = socket.getaddrinfo(host, port, proto=socket.IPPROTO_TCP)
    except socket.gaierror as e:
        raise ValueError(f"could not resolve host '{host}': {e}")
    addrs = {info[4][0] for info in infos}
    if not addrs:
        raise ValueError(f"host '{host}' did not resolve to any address")
    for addr in addrs:
        if not _ip_is_public(addr):
            raise ValueError(f"blocked private/reserved address {addr} for host '{host}'")


async def fetch_url(url: str) -> str:
    """Fetch a URL via the shared pooled client, guarding against SSRF.

    Every request — the initial URL and each redirect hop — is validated by
    assert_public_url() before the connection is made. 4xx responses return an
    empty string (page missing/forbidden — skip silently); 5xx responses raise
    so they surface as errors in the crawl log.
    """
    client = await _get_http_client()
    current = _strip_fragment(url)
    for _ in range(10):  # bounded redirect chain
        assert_public_url(current)          # re-checked on every hop
        r = await client.get(current)
        if r.is_redirect:
            location = r.headers.get("location")
            if not location:
                break
            current = _strip_fragment(urljoin(current, location))
            continue
        if 400 <= r.status_code < 500:
            return ""
        r.raise_for_status()
        return r.text
    raise ValueError(f"too many redirects fetching {url}")


def _assert_c4ai_result_public(r, requested_url: str) -> None:
    """Re-check a crawl4ai result's final (post-redirect) URL against the SSRF guard.

    The headless browser follows redirects itself with no per-hop validation, so
    a public seed that redirects to an internal host would otherwise be fetched
    and its content indexed. If the landing URL is internal this raises, and the
    caller drops the page (nothing is indexed or shown to the user).
    """
    final = getattr(r, "redirected_url", None) or getattr(r, "url", None) or ""
    if final and final != requested_url:
        assert_public_url(final)


def is_llms_txt(url: str, content: str) -> bool:
    """Return True if the URL points to an llms.txt / llms-full.txt manifest."""
    return url.endswith("llms.txt") or url.endswith("llms-full.txt")


def parse_llms_txt(content: str, base_url: str) -> list[dict]:
    """
    Parse an llms.txt manifest and return a list of {url, description} dicts.

    The llms.txt spec (https://llmstxt.org/) uses Markdown-style links:
        - [Page Title](https://example.com/page): optional description

    We also handle bare absolute URLs (one per line) as a fallback.
    Comment lines (#) and blockquotes (>) are ignored.
    Duplicate URLs are deduplicated.
    """
    links = []
    seen: set[str] = set()

    for line in content.splitlines():
        line = line.strip()
        if not line or line.startswith("#") or line.startswith(">"):
            continue

        # Extract all [text](url) patterns on this line
        for m in re.finditer(r"\[([^\]]*)\]\(([^)]+)\)", line):
            desc, href = m.group(1), m.group(2)
            if href.startswith("http") or href.startswith("/"):
                full = _strip_fragment(urljoin(base_url, href))
                if full not in seen:
                    seen.add(full)
                    links.append({"url": full, "description": desc})

        # Bare absolute URL on its own line (alternative llms.txt format)
        bare = re.match(r"^(https?://\S+)$", line)
        if bare:
            full = _strip_fragment(bare.group(1))
            if full not in seen:
                seen.add(full)
                links.append({"url": full, "description": full})

    return links


def extract_text(html: str, url: str) -> str:
    """
    Extract clean readable text from a fetched page.

    Plain-text and Markdown files (URL ends in .md/.txt, or content has no
    HTML tags) are returned directly — no stripping needed and HTML parsers
    would mangle them.  For real HTML pages we try trafilatura first (best
    quality), then BeautifulSoup, then fall back to the raw content.
    """
    url_lower = url.split("?")[0].lower()
    if url_lower.endswith(".md") or url_lower.endswith(".txt"):
        return html.strip()

    stripped = html.strip()
    if stripped and not stripped.startswith("<"):
        return stripped

    if HAS_TRAFILATURA:
        t = trafilatura.extract(html, include_links=False, include_images=False)
        if t:
            return t
    if HAS_BS4:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    return html


def _strip_fragment(url: str) -> str:
    """Remove the #anchor from a URL — same page regardless of anchor."""
    parsed = urlparse(url)
    return parsed._replace(fragment="").geturl()


# File extensions that are never useful to crawl as text content.
# Binary downloads, media, archives, executables, and data files.
_SKIP_EXTENSIONS: frozenset[str] = frozenset({
    # Archives / compressed
    ".zip", ".tar", ".gz", ".tgz", ".bz2", ".xz", ".zst", ".rar", ".7z",
    # Executables / installers
    ".exe", ".msi", ".dmg", ".pkg", ".deb", ".rpm", ".apk", ".appimage",
    # Images
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".ico", ".bmp", ".tiff",
    # Audio / video
    ".mp3", ".mp4", ".wav", ".ogg", ".flac", ".avi", ".mov", ".mkv", ".webm",
    # Documents (handled separately via file upload, not crawl)
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".odt",
    # Data / binary
    ".bin", ".dat", ".iso", ".img", ".whl", ".jar",
    # Fonts
    ".woff", ".woff2", ".ttf", ".eot",
})


def _is_crawlable_url(url: str) -> bool:
    """Return False if the URL points to a known binary/non-text file type."""
    path = urlparse(url).path.lower().split("?")[0]
    _, ext = os.path.splitext(path)
    return ext not in _SKIP_EXTENSIONS


def _extract_html_links(html: str, base_url: str) -> list[str]:
    """Return all absolute HTTP(S) links found in an HTML page, excluding binary file types."""
    if not HAS_BS4:
        return []
    try:
        soup = BeautifulSoup(html, "html.parser")
        links: list[str] = []
        for a in soup.find_all("a", href=True):
            href = _strip_fragment(urljoin(base_url, a["href"]))
            p = urlparse(href)
            if p.scheme in ("http", "https") and p.netloc and _is_crawlable_url(href):
                links.append(href)
        return links
    except Exception:
        return []


_robots_cache: dict[str, RobotFileParser] = {}

def can_crawl(url: str) -> bool:
    """
    Check robots.txt to see if RediRecallBot is permitted to crawl this URL.
    Caches the parsed robots.txt per netloc so the same domain is only
    fetched once per crawl session (not once per page).
    Returns True on any error (conservative: allow if unsure).
    """
    try:
        parsed = urlparse(url)
        netloc = parsed.netloc
        if netloc not in _robots_cache:
            robots_url = f"{parsed.scheme}://{netloc}/robots.txt"
            # SSRF guard: robots.txt is fetched before the page, so validate the
            # host here too — otherwise this is a blind GET against an internal
            # address for any URL that reached the crawler unvalidated.
            assert_public_url(robots_url)
            rp = RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            _robots_cache[netloc] = rp
        return _robots_cache[netloc].can_fetch("RediRecallBot", url)
    except Exception:
        return True


async def crawl_url(
    instance: str,
    url: str,
    depth: int = 0,
    visited: set | None = None,       # unused — kept for caller compatibility
    progress_cb=None,
    respect_robots: bool = True,
    local_only: bool = True,
    path_prefix_only: bool = False,
    max_pages: int = 0,
    _root_netloc: str = "",            # unused — kept for caller compatibility
    _counter: dict | None = None,      # unused — kept for caller compatibility
    rc: redis.Redis | None = None,
    concurrency: int = 10,
    js_render: bool = False,
    js_concurrency: int = 3,
    smart_mode: bool = True,
    min_words: int = 100,
    force_reindex: bool = False,
):
    """
    Parallel BFS web crawler with optional JavaScript rendering.

    Architecture
    ------------
    • BFS queue (asyncio.Queue) drives link discovery — breadth-first so
      shallow pages are indexed first and progress is visible immediately.
    • asyncio.Semaphore(concurrency) limits how many pages are in-flight at
      once without spawning an unbounded number of tasks.
    • A single shared httpx.AsyncClient is reused across all requests in the
      crawl for connection keep-alive and HTTP/2 multiplexing (fast path).
    • With js_render=True (requires crawl4ai + playwright) a single browser
      context is opened once and reused across all workers — far cheaper than
      launching a new browser tab per page.

    Speed improvement vs old sequential recursive crawler
    -----------------------------------------------------
    Old: each page was fetch→extract→embed→store before the next URL was even
         requested.  100 pages × 250 ms avg latency = ~25 s minimum wall time.
    New: up to `concurrency` (default 10) pages fetched concurrently.
         100 pages ÷ 10 workers × 250 ms = ~2.5 s minimum (≈10× faster).
    """
    seed_url    = _strip_fragment(url)
    _seed_parsed = urlparse(seed_url)
    root_netloc  = _seed_parsed.netloc
    # Derive path prefix: use the directory portion of the seed URL's path.
    # e.g. http://my.org/list/3/       → prefix "/list/3/"
    #      http://my.org/list/3/index  → prefix "/list/3/"
    _seed_path = _seed_parsed.path
    if not _seed_path.endswith("/"):
        _seed_path = _seed_path.rsplit("/", 1)[0] + "/"
    root_path_prefix = _seed_path   # used when path_prefix_only=True

    # Clear per-crawl robots.txt cache so stale results don't carry over
    _robots_cache.clear()

    # ── Shared mutable state (all workers in the same crawl share these) ───
    visited_urls: set[str]   = set()
    visited_lock             = asyncio.Lock()
    counter                  = {"count": 0}
    sem                      = asyncio.Semaphore(max(1, concurrency))
    queue: asyncio.Queue     = asyncio.Queue()

    visited_urls.add(seed_url)
    queue.put_nowait((seed_url, depth))

    # ── URL-level deduplication across crawl sessions ──────────────────────
    # Load the full indexed-URL set into memory once at crawl start.
    # All per-page checks are then O(1) Python set lookups — zero Redis
    # round-trips and zero thread-pool overhead during the crawl.
    _crawl_rc         = rc or rc_for_instance(instance)
    _indexed_urls_key = f"rag:{instance}:indexed_urls"
    if force_reindex:
        # Wipe the URL skip-list so every page is fetched and re-embedded fresh
        try:
            await asyncio.to_thread(_crawl_rc.delete, _indexed_urls_key)
        except Exception:
            pass
        _indexed_urls_mem: set[str] = set()
    else:
        try:
            # Loading the whole indexed-URL set is one sync command; keep it off the loop
            # so a large skip-list doesn't stall chat at crawl start.
            members = await asyncio.to_thread(_crawl_rc.smembers, _indexed_urls_key)
            _indexed_urls_mem = {u.decode() if isinstance(u, bytes) else u for u in members}
        except Exception:
            _indexed_urls_mem = set()

    # Collect newly indexed URLs in memory; flush to Redis in a single pipeline
    # after the crawl finishes rather than one SADD per page.
    _newly_indexed: list[str] = []
    _newly_indexed_lock = asyncio.Lock()

    async def _mark_url_indexed(u: str) -> None:
        _indexed_urls_mem.add(u)
        async with _newly_indexed_lock:
            _newly_indexed.append(u)
            # Flush to Redis every 50 URLs so indexed_urls survives cancellation
            if len(_newly_indexed) % 50 == 0:
                to_flush = list(_newly_indexed)
                try:
                    pipe = _crawl_rc.pipeline(transaction=False)
                    for fu in to_flush:
                        pipe.sadd(_indexed_urls_key, fu)
                    await asyncio.to_thread(pipe.execute)   # only the round-trip blocks
                except Exception:
                    pass

    # ── Inner: fetch, extract, ingest one page and enqueue its children ────
    async def process_page(page_url: str, page_depth: int, fetch_fn):
        # ── Cheap pre-checks BEFORE acquiring the semaphore ────────────────
        # These are all O(1) in-memory lookups — no I/O, no threads needed.
        if max_pages > 0 and counter["count"] >= max_pages:
            return
        if not _is_crawlable_url(page_url):
            if progress_cb:
                await progress_cb(page_url, "skipped", 0, "binary file type", counter["count"])
            return
        if not is_llms_txt(page_url, "") and page_url in _indexed_urls_mem:
            if progress_cb:
                await progress_cb(page_url, "skipped", 0, "already indexed", counter["count"])
            return
        # robots.txt: if domain already cached, check inline; otherwise fetch in thread
        if respect_robots:
            parsed_url = urlparse(page_url)
            if parsed_url.netloc in _robots_cache:
                # Cache hit — synchronous dict lookup, no thread needed
                if not _robots_cache[parsed_url.netloc].can_fetch("RediRecallBot", page_url):
                    if progress_cb:
                        await progress_cb(page_url, "blocked", 0, "", counter["count"])
                    return
            else:
                # Cache miss — fetch robots.txt in a thread (blocks until done)
                if not await asyncio.to_thread(can_crawl, page_url):
                    if progress_cb:
                        await progress_cb(page_url, "blocked", 0, "", counter["count"])
                    return

        async with sem:
            try:
                # Re-check max_pages inside the semaphore (counter may have advanced)
                if max_pages > 0 and counter["count"] >= max_pages:
                    return

                if progress_cb:
                    await progress_cb(page_url, "crawling", 0, "", counter["count"])

                raw_content, discovered_links = await fetch_fn(page_url)

                # llms.txt manifest → parse links and queue them all
                if is_llms_txt(page_url, raw_content):
                    manifest_links = parse_llms_txt(raw_content, page_url)
                    if progress_cb:
                        await progress_cb(page_url, "parsed_llms_txt",
                                          len(manifest_links), "", counter["count"])
                    for lnk in manifest_links:
                        lurl = _strip_fragment(lnk["url"])
                        if max_pages > 0 and counter["count"] >= max_pages:
                            break
                        async with visited_lock:
                            if lurl in visited_urls:
                                continue
                            visited_urls.add(lurl)
                        queue.put_nowait((lurl, 0))  # depth=0: don't follow further links
                    return

                # Regular page: extract text → chunk+dedup → hand off to embed worker
                # js_render path: raw_content is already clean markdown from crawl4ai
                text = raw_content if js_render else extract_text(raw_content, page_url)
                if not text.strip():
                    if progress_cb:
                        await progress_cb(page_url, "skipped", 0, "empty content", counter["count"])
                    return

                # _prepare_chunks: CPU-light (no model calls) — chunk, dedup, reserve IDs
                records = await asyncio.to_thread(_prepare_chunks, instance, text, page_url, _crawl_rc, force_reindex)
                # Always mark URL as visited so future crawls don't re-fetch duplicate-content pages
                await _mark_url_indexed(page_url)
                if not records:
                    if progress_cb:
                        await progress_cb(page_url, "skipped", 0, "duplicate content", counter["count"])
                    return

                counter["count"] += 1
                # Embed worker picks this up and calls progress_cb("indexed", ...) after flush
                await embed_queue.put((page_url, records))

                # Enqueue child links if depth allows
                if page_depth > 0:
                    for href in discovered_links:
                        if max_pages > 0 and counter["count"] >= max_pages:
                            break
                        href = _strip_fragment(href)
                        parsed_href = urlparse(href)
                        if local_only and parsed_href.netloc != root_netloc:
                            continue
                        if path_prefix_only and not parsed_href.path.startswith(root_path_prefix):
                            continue
                        async with visited_lock:
                            if href in visited_urls:
                                continue
                            visited_urls.add(href)
                        queue.put_nowait((href, page_depth - 1))

            except Exception as e:
                append_log({
                    "ts": datetime.now(timezone.utc).isoformat(),
                    "instance": instance, "source": page_url,
                    "chunks": 0, "status": "error", "error": str(e),
                })
                if progress_cb:
                    await progress_cb(page_url, "error", 0, str(e), counter["count"])

    # ── Cross-page embed queue — decouples fetch workers from model calls ───
    # process_page puts (page_url, records) here after chunking/dedup.
    # embed_worker drains it in batches so encode() is called once per batch
    # instead of once per page — amortises the ~100 ms per-call model overhead.
    embed_queue: asyncio.Queue = asyncio.Queue()

    # ── BFS driver: N concurrent workers pulling from the shared queue ─────
    async def run_bfs(fetch_fn):
        EMBED_BATCH = 64      # max chunks per encode() call
        EMBED_FLUSH_MS = 0.10 # seconds to wait before flushing a partial batch

        async def embed_worker():
            """Drain embed_queue in cross-page batches; call add_chunks once per batch."""
            batch_records: list[dict] = []
            batch_meta: list[tuple[str, int]] = []   # (page_url, n_chunks)

            async def flush():
                if not batch_records:
                    return
                records_snap = list(batch_records)
                meta_snap    = list(batch_meta)
                batch_records.clear()
                batch_meta.clear()
                try:
                    await asyncio.to_thread(add_chunks, instance, records_snap, _crawl_rc)
                except Exception as e:
                    log.warning("embed_worker: batch embed/store failed (%s) — %d chunks lost", e, len(records_snap))
                    for page_url, n in meta_snap:
                        append_log({
                            "ts": datetime.now(timezone.utc).isoformat(),
                            "instance": instance, "source": page_url,
                            "chunks": 0, "status": "error", "error": str(e),
                        })
                        if progress_cb:
                            await progress_cb(page_url, "error", 0, str(e), counter["count"])
                    return
                for page_url, n in meta_snap:
                    append_log({
                        "ts": datetime.now(timezone.utc).isoformat(),
                        "instance": instance, "source": page_url,
                        "chunks": n, "status": "ok",
                    })
                    if progress_cb:
                        await progress_cb(page_url, "indexed", n, "", counter["count"])

            while True:
                try:
                    item = await asyncio.wait_for(embed_queue.get(), timeout=EMBED_FLUSH_MS)
                except asyncio.TimeoutError:
                    await flush()
                    continue
                except asyncio.CancelledError:
                    await flush()
                    raise

                if item is None:   # sentinel — crawl finished
                    await flush()
                    break

                page_url, records = item
                batch_records.extend(records)
                batch_meta.append((page_url, len(records)))
                if len(batch_records) >= EMBED_BATCH:
                    await flush()

        async def worker():
            while True:
                page_url, page_depth = await queue.get()
                try:
                    await process_page(page_url, page_depth, fetch_fn)
                except Exception as e:
                    log.warning("worker: unhandled exception for %s: %s", page_url, e)
                finally:
                    queue.task_done()

        num_workers = min(max(1, concurrency), 50)
        embed_task   = asyncio.create_task(embed_worker())
        worker_tasks = [asyncio.create_task(worker()) for _ in range(num_workers)]

        try:
            await queue.join()   # blocks until every queued item has been task_done()
        finally:
            for t in worker_tasks:
                t.cancel()
            await asyncio.gather(*worker_tasks, return_exceptions=True)

        # Signal embed worker to flush remaining chunks and exit
        await embed_queue.put(None)
        await embed_task

        # Flush all newly indexed URLs to Redis in a single pipeline
        if _newly_indexed:
            try:
                pipe = _crawl_rc.pipeline(transaction=False)
                for u in _newly_indexed:
                    pipe.sadd(_indexed_urls_key, u)
                await asyncio.to_thread(pipe.execute)   # only the round-trip blocks
            except Exception:
                pass

    # ── Helpers shared by all browser-based fetch paths ────────────────────
    def _links_from_result(r) -> list[str]:
        """Extract absolute HTTP links from a crawl4ai result object."""
        return (
            [l["href"] for l in r.links.get("internal", []) if l.get("href", "").startswith("http")] +
            [l["href"] for l in r.links.get("external", []) if l.get("href", "").startswith("http")]
        )

    async def fetch_httpx(u: str) -> tuple[str, list[str]]:
        html  = await fetch_url(u)
        links = _extract_html_links(html, u)
        return html, links

    # ── Resolve which modes are actually available ──────────────────────────
    use_js    = js_render
    use_smart = smart_mode and not js_render   # smart = httpx-first + JS fallback

    if (use_js or use_smart) and not HAS_CRAWL4AI:
        if use_js:
            log.warning(
                "js_render=True but crawl4ai not installed — "
                "falling back to httpx. Run: pip install crawl4ai && playwright install chromium"
            )
        use_js    = False
        use_smart = False

    # ── Browser config — shared by both full-JS and smart-fallback paths ───
    # One browser process is reused across ALL worker coroutines for the
    # lifetime of this crawl.  Extra args block image rendering and suppress
    # background network activity, cutting per-tab RAM/CPU by ~60%.
    _browser_cfg = _C4AIBrowserConfig(
        headless=True,
        verbose=False,
        extra_args=_BROWSER_EXTRA_ARGS,
    ) if HAS_CRAWL4AI else None

    # domcontentloaded fires as soon as HTML is parsed — no need to wait for
    # analytics/tracking calls that never resolve on networkidle.
    _run_cfg = _C4AIRunConfig(
        wait_until="domcontentloaded",   # was "networkidle" — 3-10× faster per page
        page_timeout=15_000,             # was 30 000 ms
        word_count_threshold=10,
        exclude_all_images=True,         # skip image downloads entirely
        exclude_external_images=True,
    ) if HAS_CRAWL4AI else None

    # Separate semaphore caps Playwright tabs independently of httpx workers.
    # Browser tabs are ~150 MB each; default cap of 3 keeps RAM predictable.
    _js_sem = asyncio.Semaphore(max(1, js_concurrency))

    # ── Dispatch ────────────────────────────────────────────────────────────
    if use_js:
        # Full JS mode: every page goes through Playwright.
        async with _C4AIWebCrawler(config=_browser_cfg) as _c4ai:
            async def fetch_js(u: str) -> tuple[str, list[str]]:
                assert_public_url(u)          # block internal targets before the browser fetch
                async with _js_sem:
                    r = await _c4ai.arun(u, config=_run_cfg)
                _assert_c4ai_result_public(r, u)   # block redirect-to-internal
                text = (r.markdown.fit_markdown
                        if (r.markdown and r.markdown.fit_markdown)
                        else (r.html or ""))
                return text, _links_from_result(r)

            await run_bfs(fetch_js)

    elif use_smart:
        # Smart mode: try httpx first (fast, zero browser overhead).
        # Only pages whose extracted text is below min_words are retried
        # with Playwright — typically 5-15% of pages in a real docs crawl.
        async with _C4AIWebCrawler(config=_browser_cfg) as _c4ai:
            async def fetch_smart(u: str) -> tuple[str, list[str]]:
                # SSRF guard for BOTH the httpx and browser paths. Raised here,
                # OUTSIDE the try below, so an internal-target rejection is never
                # swallowed into the browser fallback (which would defeat it).
                assert_public_url(u)
                # ── Fast path: pooled httpx ────────────────────────────────
                httpx_html:  str       = ""
                httpx_links: list[str] = []
                try:
                    httpx_html, httpx_links = await fetch_httpx(u)
                    text = extract_text(httpx_html, u)
                    if len(text.split()) >= min_words:
                        return httpx_html, httpx_links   # sufficient content, done
                    if not httpx_html:
                        # Empty response (4xx) — JS won't help; skip browser entirely
                        return httpx_html, httpx_links
                    log.debug("smart-crawl: thin content (%d words) on %s — retrying with JS",
                              len(text.split()), u)
                except Exception as exc:
                    log.debug("smart-crawl: httpx failed for %s (%s) — retrying with JS", u, exc)

                # ── JS fallback: Playwright (rate-limited by _js_sem) ──────
                try:
                    async with _js_sem:
                        r = await _c4ai.arun(u, config=_run_cfg)
                    _assert_c4ai_result_public(r, u)   # block redirect-to-internal
                    text = (r.markdown.fit_markdown
                            if (r.markdown and r.markdown.fit_markdown)
                            else (r.html or ""))
                    return text, _links_from_result(r)
                except Exception as js_exc:
                    # Browser context was None or crashed (e.g. after a cancelled crawl).
                    # Return whatever httpx gave us rather than failing the page entirely.
                    log.debug("smart-crawl: JS fallback failed for %s (%s) — using httpx result", u, js_exc)
                    return httpx_html, httpx_links

            await run_bfs(fetch_smart)

    else:
        # Pure httpx — no browser, maximum speed for pre-rendered HTML sites.
        await run_bfs(fetch_httpx)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# LLM PROVIDER — OLLAMA
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def ollama_base() -> str:
    """Return the base URL for the configured Ollama server."""
    cfg = _config.get("ollama", {})
    host = cfg.get("host", "http://localhost").rstrip("/")
    port = cfg.get("port", 11434)
    return f"{host}:{port}"

# Cache the model list for 30 s to avoid hammering Ollama on concurrent calls.
_ollama_models_cache: list[dict] = []
_ollama_models_ts: float = 0.0
_OLLAMA_MODELS_TTL = 30.0


async def ollama_models() -> list[dict]:
    """
    Fetch available models from the Ollama server.
    Returns a list of {name, size, vision, details} dicts.
    Vision detection checks model family tags and common naming conventions.
    Result is cached for _OLLAMA_MODELS_TTL seconds to avoid hammering Ollama
    when /api/status/ollama and /api/ollama/models are called concurrently.
    """
    global _ollama_models_cache, _ollama_models_ts
    if time.time() - _ollama_models_ts < _OLLAMA_MODELS_TTL and _ollama_models_cache:
        return _ollama_models_cache
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            res = await client.get(f"{ollama_base()}/api/tags")
            data = res.json()
            models = []
            for m in data.get("models", []):
                name = m["name"]
                details = m.get("details", {})
                families = details.get("families", []) or []
                vision = any(f in ["clip", "llava"] for f in families) or any(
                    v in name.lower()
                    for v in ["llava", "bakllava", "moondream", "vision", "minicpm", "gemma3", "qwen-vl"]
                )
                models.append({"name": name, "size": m.get("size", 0), "vision": vision, "details": details})
            _ollama_models_cache = models
            _ollama_models_ts = time.time()
            return models
    except Exception as e:
        log.error(f"Ollama models error: {e}")
        return []


_IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

def _tool_result_to_markdown(tool_calls: list) -> str:
    """
    Convert Ollama tool_call entries into renderable markdown.
    If a tool result looks like an image (data URI, local path, or HTTP URL),
    emit markdown that the frontend can detect and render.
    """
    parts = []
    for tc in tool_calls:
        fn   = tc.get("function", {})
        name = fn.get("name", "tool")
        args = fn.get("arguments", {})

        # Check if any argument value looks like an image
        for val in (args.values() if isinstance(args, dict) else []):
            val = str(val).strip()
            # Data URI
            if val.startswith("data:image/"):
                parts.append(f"\n![{name} result]({val})\n")
                continue
            # Local file path with image extension
            p = Path(val)
            if p.suffix.lower() in _IMG_EXTS and p.is_file():
                parts.append(f"\n![{name} result](/api/files/image?path={val})\n")
                continue
            # HTTP URL pointing to an image
            if val.startswith(("http://", "https://")) and Path(val).suffix.lower() in _IMG_EXTS:
                parts.append(f"\n![{name} result]({val})\n")
                continue
        else:
            # No image found — emit the raw tool call as a code block
            parts.append(f"\n```json\n[tool: {name}] {json.dumps(args, ensure_ascii=False)}\n```\n")
    return "".join(parts)


def _to_ollama_messages(messages: list) -> list:
    """
    Convert OpenAI-style messages to Ollama format.

    OpenAI multi-modal content:
        [{"type": "text", "text": "..."}, {"type": "image_url", "image_url": {"url": "data:image/...;base64,..."}}]

    Ollama expects:
        {"role": "user", "content": "...", "images": ["<raw_base64>"]}

    Non-multi-modal messages are passed through unchanged.
    """
    result = []
    for m in messages:
        content = m["content"]
        if isinstance(content, list):
            text_parts: list[str] = []
            b64_images: list[str] = []
            for part in content:
                if part.get("type") == "text":
                    text_parts.append(part["text"])
                elif part.get("type") == "image_url":
                    url = part["image_url"]["url"]
                    if "base64," in url:
                        b64_images.append(url.split("base64,", 1)[1])
            msg: dict = {"role": m["role"], "content": " ".join(text_parts)}
            if b64_images:
                msg["images"] = b64_images
        else:
            msg = {"role": m["role"], "content": content}
        result.append(msg)
    return result


async def ollama_stream(messages: list, model: str, images: list[str] | None = None):
    """
    Stream tokens from Ollama /api/chat endpoint.
    Yields (token: str, done: bool) tuples.

    Handles both plain text responses and tool_call responses:
    - Plain content tokens are yielded directly.
    - tool_calls entries are converted to renderable markdown (images or code blocks).
    """
    payload = {"model": model, "messages": _to_ollama_messages(messages), "stream": True}
    url = f"{ollama_base()}/api/chat"
    async with httpx.AsyncClient(timeout=120) as client:
        async with client.stream("POST", url, json=payload) as resp:
            async for line in resp.aiter_lines():
                if not line.strip():
                    continue
                try:
                    d    = json.loads(line)
                    msg  = d.get("message", {})
                    done = d.get("done", False)

                    # Standard text content
                    token = msg.get("content", "")

                    # Tool calls — convert to markdown so the frontend can render images
                    tool_calls = msg.get("tool_calls", [])
                    if tool_calls:
                        token += _tool_result_to_markdown(tool_calls)

                    if token:
                        yield token, done
                    elif done:
                        yield "", True
                    if done:
                        break
                except Exception:
                    pass

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# LLM PROVIDER — ANTHROPIC CLAUDE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Static model list (Claude API doesn't have a /models endpoint that lists them)
CLAUDE_MODELS = [
    {"id": "claude-opus-4-6",           "name": "Claude Opus 4.6",   "context": 200000},
    {"id": "claude-sonnet-4-6",         "name": "Claude Sonnet 4.6", "context": 200000},
    {"id": "claude-haiku-4-5-20251001", "name": "Claude Haiku 4.5",  "context": 200000},
]


def _to_claude_content(content):
    """
    Convert OpenAI-style message content to Claude API format.

    OpenAI image:  {"type": "image_url", "image_url": {"url": "data:image/jpeg;base64,..."}}
    Claude image:  {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": "..."}}

    Plain strings pass through unchanged.
    """
    if not isinstance(content, list):
        return content
    result = []
    for part in content:
        if part.get("type") == "text":
            result.append({"type": "text", "text": part["text"]})
        elif part.get("type") == "image_url":
            url = part["image_url"]["url"]
            if "base64," in url:
                header, data = url.split("base64,", 1)
                media_type = header.rstrip(";").split(":")[-1]  # e.g. "image/jpeg"
                result.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": media_type, "data": data},
                })
    return result


async def claude_stream(messages: list, model: str):
    """
    Stream tokens from Anthropic Claude using the native anthropic SDK.
    Yields (token: str, done: bool) tuples.
    Falls back to a clear error if the SDK is not installed.
    """
    api_key  = _config.get("claude", {}).get("api_key", "")
    base_url = _config.get("claude", {}).get("base_url", "https://api.anthropic.com").rstrip("/")

    if not api_key:
        yield "Error: Claude API key not configured.", True
        return
    if not _ANTHROPIC_AVAILABLE:
        yield "Error: anthropic package not installed. Run: pip install anthropic", True
        return

    system_msg, claude_msgs = "", []
    for m in messages:
        if m["role"] == "system":
            system_msg = m["content"] if isinstance(m["content"], str) else str(m["content"])
        else:
            claude_msgs.append({"role": m["role"], "content": _to_claude_content(m["content"])})

    try:
        client = _cached_client("anthropic", api_key, base_url)
        kwargs: dict = {"model": model, "messages": claude_msgs, "max_tokens": 4096}
        if system_msg:
            kwargs["system"] = system_msg
        async with client.messages.stream(**kwargs) as stream:
            async for text in stream.text_stream:
                yield text, False
        yield "", True
    except _anthropic.APIStatusError as e:
        yield f"Error: {e.message}", True
    except Exception as e:
        yield f"Claude error: {e}", True

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# LLM PROVIDER — OPENAI
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Models we expose in the UI.  Fetched dynamically from /v1/models when possible
# but kept as a static fallback so the UI always has something to show.
OPENAI_MODELS_STATIC = [
    {"id": "gpt-4o",            "name": "GPT-4o",               "context": 128000},
    {"id": "gpt-4o-mini",       "name": "GPT-4o mini",          "context": 128000},
    {"id": "gpt-4.1",           "name": "GPT-4.1",              "context": 1047576},
    {"id": "gpt-4.1-mini",      "name": "GPT-4.1 mini",         "context": 1047576},
    {"id": "gpt-4.1-nano",      "name": "GPT-4.1 nano",         "context": 1047576},
    {"id": "o1",                "name": "o1",                   "context": 200000},
    {"id": "o3",                "name": "o3",                   "context": 200000},
    {"id": "o3-mini",           "name": "o3-mini",              "context": 200000},
    {"id": "o4-mini",           "name": "o4-mini",              "context": 200000},
]


async def openai_models() -> list[dict]:
    """
    Fetch the list of available models from the OpenAI /v1/models endpoint.
    Uses the native openai SDK. Falls back to the static list on any error.
    Only returns chat-capable models (id starts with 'gpt-' or 'o').
    """
    api_key  = _config.get("openai", {}).get("api_key", "")
    base_url = _config.get("openai", {}).get("base_url", "https://api.openai.com").rstrip("/")
    if not api_key or not _OPENAI_SDK_AVAILABLE:
        return OPENAI_MODELS_STATIC

    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        page = await client.models.list()
        models = []
        for m in sorted(page.data, key=lambda x: x.id):
            mid = m.id
            if mid.startswith(("gpt-4o", "gpt-4.1", "o1", "o3", "o4")):
                name = next((x["name"] for x in OPENAI_MODELS_STATIC if x["id"] == mid), mid)
                ctx  = next((x["context"] for x in OPENAI_MODELS_STATIC if x["id"] == mid), 128000)
                models.append({"id": mid, "name": name, "context": ctx})
        return models if models else OPENAI_MODELS_STATIC
    except Exception as e:
        log.error(f"OpenAI models error: {e}")
        return OPENAI_MODELS_STATIC


async def openai_stream(messages: list, model: str):
    """
    Stream tokens from the OpenAI /v1/chat/completions endpoint using the native openai SDK.
    Yields (token: str, done: bool) tuples — same interface as claude_stream.

    Compatible with any OpenAI-compatible API by changing the base_url in settings.
    """
    api_key  = _config.get("openai", {}).get("api_key", "")
    base_url = _config.get("openai", {}).get("base_url", "https://api.openai.com").rstrip("/")

    if not api_key:
        yield "Error: OpenAI API key not configured.", True
        return
    if not _OPENAI_SDK_AVAILABLE:
        yield "Error: openai package not installed. Run: pip install openai", True
        return

    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        stream = await client.chat.completions.create(
            model=model, messages=messages, stream=True,
        )
        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content, False
            if chunk.choices[0].finish_reason:
                break
        yield "", True
    except Exception as e:
        yield f"OpenAI error: {e}", True


# ── Qwen / DashScope ──────────────────────────────────────────────────────────
# DashScope exposes an OpenAI-compatible /v1/chat/completions endpoint.
# Free-tier models: qwen-plus, qwen-turbo, qwen-long (generous monthly quotas).
# Premium: qwen-max, qwen-max-longcontext.
# Get a free API key at: https://qwen.ai

QWEN_MODELS_STATIC = [
    {"id": "qwen-plus",              "name": "Qwen Plus (free tier)",       "context": 131072},
    {"id": "qwen-turbo",             "name": "Qwen Turbo (free tier)",      "context": 131072},
    {"id": "qwen-long",              "name": "Qwen Long (free tier)",       "context": 10000000},
    {"id": "qwen-max",               "name": "Qwen Max",                    "context": 131072},
    {"id": "qwen-max-longcontext",   "name": "Qwen Max (long ctx)",         "context": 1000000},
    {"id": "qwen2.5-72b-instruct",   "name": "Qwen 2.5 72B Instruct",      "context": 131072},
    {"id": "qwen2.5-7b-instruct",    "name": "Qwen 2.5 7B Instruct",       "context": 131072},
    {"id": "qwen2.5-coder-32b-instruct", "name": "Qwen 2.5 Coder 32B",     "context": 131072},
]


async def qwen_stream(messages: list, model: str):
    """
    Stream tokens from Alibaba DashScope using the openai SDK with DashScope's base URL.
    Yields (token: str, done: bool) — same interface as openai_stream / claude_stream.

    Free-tier models: qwen-plus, qwen-turbo, qwen-long.
    API key: get a free key at qwen.ai (DASHSCOPE_API_KEY env var supported).
    """
    api_key  = _config.get("qwen", {}).get("api_key", "")
    base_url = _config.get("qwen", {}).get("base_url",
                           "https://dashscope.aliyuncs.com/compatible-mode/v1").rstrip("/")

    if not api_key:
        yield "Error: Qwen API key not configured. Get a free key at qwen.ai", True
        return
    if not _OPENAI_SDK_AVAILABLE:
        yield "Error: openai package not installed. Run: pip install openai", True
        return

    try:
        # DashScope base_url already includes /v1 path — pass it directly to the SDK
        client = _cached_client("openai", api_key, base_url)
        stream = await client.chat.completions.create(
            model=model, messages=messages, stream=True,
        )
        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content, False
            if chunk.choices[0].finish_reason:
                break
        yield "", True
    except Exception as e:
        yield f"Qwen error: {e}", True

# ── Mistral ─────────────────────────────────────────────────────────────────────
# OpenAI-compatible endpoint (api.mistral.ai/v1). EU-hosted. The free "Experiment"
# tier covers every model with generous limits and needs no credit card (phone
# verification only) — https://console.mistral.ai.
MISTRAL_MODELS_STATIC = [
    {"id": "mistral-small-latest",   "name": "Mistral Small (free tier)",   "context": 32000},
    {"id": "open-mistral-nemo",      "name": "Mistral Nemo (free tier)",    "context": 128000},
    {"id": "mistral-large-latest",   "name": "Mistral Large",               "context": 128000},
    {"id": "codestral-latest",       "name": "Codestral (code)",            "context": 256000},
    {"id": "ministral-8b-latest",    "name": "Ministral 8B",                "context": 128000},
    {"id": "ministral-3b-latest",    "name": "Ministral 3B",                "context": 128000},
    {"id": "pixtral-12b-2409",       "name": "Pixtral 12B (vision)",        "context": 128000},
]

async def mistral_stream(messages: list, model: str):
    """
    Stream tokens from Mistral La Plateforme via the openai SDK (Mistral is
    OpenAI-compatible). Yields (token: str, done: bool) — same interface as the others.

    Free "Experiment" tier: all models, ~1B tokens/month, no credit card.
    API key: https://console.mistral.ai (MISTRAL_API_KEY env var supported).
    """
    api_key  = _config.get("mistral", {}).get("api_key", "")
    base_url = _config.get("mistral", {}).get("base_url",
                           "https://api.mistral.ai/v1").rstrip("/")

    if not api_key:
        yield "Error: Mistral API key not configured. Get a free key at console.mistral.ai", True
        return
    if not _OPENAI_SDK_AVAILABLE:
        yield "Error: openai package not installed. Run: pip install openai", True
        return

    try:
        # Mistral's base_url already includes /v1 — pass it directly to the SDK.
        client = _cached_client("openai", api_key, base_url)
        stream = await client.chat.completions.create(
            model=model, messages=messages, stream=True,
        )
        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content, False
            if chunk.choices[0].finish_reason:
                break
        yield "", True
    except Exception as e:
        yield f"Mistral error: {e}", True

# ── Groq ──────────────────────────────────────────────────────────────────────
# OpenAI-compatible endpoint; very fast inference; generous free tier.
# Get a free API key at console.groq.com (no credit card required).

GROQ_MODELS_STATIC = [
    {"id": "llama-3.3-70b-versatile",  "name": "Llama 3.3 70B (free)",        "context": 128000},
    {"id": "llama-3.1-8b-instant",     "name": "Llama 3.1 8B Instant (free)", "context": 131072},
    {"id": "llama3-70b-8192",          "name": "Llama 3 70B (free)",          "context": 8192},
    {"id": "llama3-8b-8192",           "name": "Llama 3 8B (free)",           "context": 8192},
    {"id": "mixtral-8x7b-32768",       "name": "Mixtral 8×7B (free)",         "context": 32768},
    {"id": "gemma2-9b-it",             "name": "Gemma 2 9B (free)",           "context": 8192},
]


async def groq_models() -> list[dict]:
    """Fetch available models from Groq using the openai SDK (Groq is OpenAI-compatible)."""
    api_key  = _config.get("groq", {}).get("api_key", "")
    base_url = _config.get("groq", {}).get("base_url", "https://api.groq.com/openai").rstrip("/")
    if not api_key or not _OPENAI_SDK_AVAILABLE:
        return GROQ_MODELS_STATIC
    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        page = await client.models.list()
        models = []
        for m in sorted(page.data, key=lambda x: x.id):
            mid = m.id
            if not mid:
                continue
            name = next((x["name"] for x in GROQ_MODELS_STATIC if x["id"] == mid), mid)
            ctx  = next((x["context"] for x in GROQ_MODELS_STATIC if x["id"] == mid),
                        getattr(m, "context_window", None) or 8192)
            models.append({"id": mid, "name": name, "context": ctx})
        return models if models else GROQ_MODELS_STATIC
    except Exception as e:
        log.error(f"Groq models error: {e}")
        return GROQ_MODELS_STATIC


async def groq_stream(messages: list, model: str):
    """
    Stream tokens from Groq using the openai SDK (Groq is fully OpenAI-compatible).
    Yields (token: str, done: bool) — same interface as openai_stream.
    """
    api_key  = _config.get("groq", {}).get("api_key", "")
    base_url = _config.get("groq", {}).get("base_url", "https://api.groq.com/openai").rstrip("/")

    if not api_key:
        yield "Error: Groq API key not configured. Get a free key at console.groq.com", True
        return
    if not _OPENAI_SDK_AVAILABLE:
        yield "Error: openai package not installed. Run: pip install openai", True
        return

    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        stream = await client.chat.completions.create(
            model=model, messages=messages, stream=True,
        )
        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content, False
            if chunk.choices[0].finish_reason:
                break
        yield "", True
    except Exception as e:
        err = str(e)
        if "429" in err or "rate_limit" in err.lower():
            yield ("Error: Groq rate limit reached (free tier has per-minute token limits). "
                   "Wait a moment and try again, or switch to a smaller model like llama-3.1-8b-instant."), True
        else:
            yield f"Groq error: {e}", True


# ── Google Gemini ──────────────────────────────────────────────────────────────
# Uses the native google-genai SDK (pip install google-genai).
# Get a free API key at aistudio.google.com · Set GEMINI_API_KEY env var.

GEMINI_MODELS_STATIC = [
    {"id": "gemini-3-flash-preview",        "name": "Gemini 3 Flash Preview",   "context": 1048576},
    {"id": "gemini-2.5-pro-preview-03-25",  "name": "Gemini 2.5 Pro Preview",   "context": 1048576},
    {"id": "gemini-2.5-flash-preview-04-17","name": "Gemini 2.5 Flash Preview", "context": 1048576},
    {"id": "gemini-2.0-flash",              "name": "Gemini 2.0 Flash",         "context": 1048576},
    {"id": "gemini-2.0-flash-lite",         "name": "Gemini 2.0 Flash Lite",    "context": 1048576},
    {"id": "gemini-1.5-flash",              "name": "Gemini 1.5 Flash",         "context": 1048576},
    {"id": "gemini-1.5-flash-8b",           "name": "Gemini 1.5 Flash 8B",      "context": 1048576},
    {"id": "gemini-1.5-pro",                "name": "Gemini 1.5 Pro",           "context": 2097152},
]


def _to_gemini_contents(messages: list) -> tuple[list, str | None]:
    """
    Convert OpenAI-format messages to Gemini native content objects.
    Returns (contents, system_instruction_text).
    Roles: OpenAI "assistant" → Gemini "model". System messages → system_instruction.
    """
    if not _GENAI_AVAILABLE:
        return [], None
    system_parts: list[str] = []
    contents = []
    for m in messages:
        role    = m.get("role", "user")
        content = m.get("content", "")
        if role == "system":
            txt = " ".join(p.get("text", "") for p in content if p.get("type") == "text") \
                  if isinstance(content, list) else content
            system_parts.append(txt)
            continue
        gemini_role = "model" if role == "assistant" else "user"
        if isinstance(content, list):
            parts = []
            for part in content:
                if part.get("type") == "text":
                    parts.append(_genai_types.Part.from_text(text=part["text"]))
                elif part.get("type") == "image_url":
                    url = part["image_url"]["url"]
                    if "base64," in url:
                        header, data = url.split("base64,", 1)
                        mime = header.rstrip(";").split(":")[-1]
                        parts.append(_genai_types.Part.from_bytes(
                            data=base64.b64decode(data), mime_type=mime))
        else:
            parts = [_genai_types.Part.from_text(text=content)]
        contents.append(_genai_types.Content(role=gemini_role, parts=parts))
    system_text = "\n\n".join(system_parts) if system_parts else None
    return contents, system_text


def _gemini_err_msg(exc: Exception) -> str:
    """Return a human-readable error message from a Gemini SDK exception."""
    s = str(exc)
    if "limit: 0" in s or "free_tier" in s.lower():
        return ("Gemini free-tier quota is zero for this project. "
                "Enable billing at console.cloud.google.com or create a new project at aistudio.google.com.")
    if "RESOURCE_EXHAUSTED" in s or "quota" in s.lower() or "429" in s:
        return "Gemini quota exhausted. Wait for daily reset (midnight Pacific) or enable billing."
    if "API_KEY_INVALID" in s or "401" in s or "403" in s:
        return "Gemini API key is invalid or revoked. Enter a new key in Settings → Gemini."
    if "NOT_FOUND" in s or "404" in s:
        return "Gemini model not found. Open Settings → Gemini, click Refresh Models and pick another."
    return f"Gemini error: {exc}"


async def gemini_models() -> list[dict]:
    """Fetch available Gemini models via the native SDK. Falls back to static list."""
    api_key = _config.get("gemini", {}).get("api_key", "")
    if not api_key or not _GENAI_AVAILABLE:
        return GEMINI_MODELS_STATIC
    try:
        client = _google_genai.Client(api_key=api_key)
        raw = await asyncio.get_event_loop().run_in_executor(
            None, lambda: list(client.models.list()))
        result = []
        for m in sorted(raw, key=lambda x: getattr(x, "name", "")):
            mid = getattr(m, "name", "")
            if mid.startswith("models/"):
                mid = mid[7:]
            if not mid.startswith("gemini"):
                continue
            name = next((x["name"] for x in GEMINI_MODELS_STATIC if x["id"] == mid), mid)
            ctx  = next((x["context"] for x in GEMINI_MODELS_STATIC if x["id"] == mid), 1048576)
            result.append({"id": mid, "name": name, "context": ctx})
        return result if result else GEMINI_MODELS_STATIC
    except Exception as e:
        log.error(f"Gemini models error: {e}")
        return GEMINI_MODELS_STATIC


async def gemini_stream(messages: list, model: str):
    """
    Stream tokens from Google Gemini using the native google-genai SDK.
    Yields (token: str, done: bool) — same interface as openai_stream / claude_stream.
    """
    api_key = _config.get("gemini", {}).get("api_key", "")
    if not api_key:
        yield "Error: Gemini API key not configured. Get a free key at aistudio.google.com", True
        return
    if not _GENAI_AVAILABLE:
        yield "Error: google-genai package not installed. Run: pip install google-genai", True
        return

    try:
        client   = _google_genai.Client(api_key=api_key)
        contents, system_text = _to_gemini_contents(messages)
        cfg = _genai_types.GenerateContentConfig(
            system_instruction=system_text if system_text else None,
        )
        log.info(f"Gemini request: model={model!r} turns={len(contents)}")
        async for chunk in await client.aio.models.generate_content_stream(
            model=model, contents=contents, config=cfg,
        ):
            text = getattr(chunk, "text", None)
            if text:
                yield text, False
        yield "", True
    except Exception as e:
        log.warning(f"Gemini stream error: {e}")
        yield _gemini_err_msg(e), True


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# HYDE — HYPOTHETICAL DOCUMENT EMBEDDINGS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

async def hyde_generate(query: str, provider: str, model: str) -> str:
    """
    Generate a short hypothetical document that would answer the query.

    HyDE (Hypothetical Document Embeddings) improves RAG retrieval by bridging
    the vocabulary gap between short queries and long documents:
      1. Ask the LLM to write a brief answer to the query.
      2. Embed the hypothetical answer (not the raw query).
      3. Use that embedding for KNN vector search.

    The hypothesis is never shown to the user — it is only used as the search
    vector.  The original query is still used for BM25 keyword search in hybrid mode.
    """
    prompt = [{"role": "user", "content": (
        "Write a concise 2-3 sentence factual passage that directly answers "
        "this question. Only write the passage, no preamble or explanation. "
        f"Question: {query}"
    )}]
    hypothesis = ""
    try:
        if provider == "claude":
            async for tok, done in claude_stream(prompt, model):
                hypothesis += tok
                if done: break
        elif provider == "openai":
            async for tok, done in openai_stream(prompt, model):
                hypothesis += tok
                if done: break
        elif provider == "qwen":
            async for tok, done in qwen_stream(prompt, model):
                hypothesis += tok
                if done: break
        elif provider == "mistral":
            async for tok, done in mistral_stream(prompt, model):
                hypothesis += tok
                if done: break
        elif provider == "groq":
            async for tok, done in groq_stream(prompt, model):
                hypothesis += tok
                if done: break
        elif provider == "gemini":
            async for tok, done in gemini_stream(prompt, model):
                hypothesis += tok
                if done: break
        else:
            async for tok, done in ollama_stream(prompt, model):
                hypothesis += tok
                if done: break
    except Exception as e:
        log.warning(f"HyDE generation failed: {e}")
    return hypothesis.strip()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# RAG INSTANCE MANAGEMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def list_rag_instances() -> list[dict]:
    """
    Return metadata for all RAG instances across all configured Redis endpoints.

    Each instance can live on a different endpoint (stored in its rag_meta).
    We scan every endpoint so the UI shows a unified list.
    """
    all_instances: dict[str, dict] = {}   # "endpoint:name" -> info

    # Build the list of endpoints to scan: default + any extras in config
    endpoints_to_scan = [("default", r())]
    for ep in _config.get("redis_endpoints", []):
        ep_name = ep.get("name", "")
        if ep_name and ep_name != "default":
            try:
                endpoints_to_scan.append((ep_name, r_for(ep_name)))
            except Exception:
                pass

    for ep_name, rc in endpoints_to_scan:
        # Skip endpoints already known to be down — avoids hanging on socket_timeout
        if _endpoint_health.get(ep_name, None) is False:
            continue
        try:
            # Count chunks per instance — scan_iter avoids blocking Redis
            for k in rc.scan_iter("rag:*:chunk:*", count=500):
                parts = k.decode().split(":")
                if len(parts) >= 3:
                    inst = parts[1]
                    key  = f"{ep_name}:{inst}"
                    all_instances.setdefault(key, {"count": 0, "ep": ep_name, "name": inst})
                    all_instances[key]["count"] += 1
            # Also pick up instances that exist only as metadata (0 chunks)
            for mk in rc.scan_iter("rag_meta:*", count=200):
                inst = mk.decode().replace("rag_meta:", "")
                key  = f"{ep_name}:{inst}"
                all_instances.setdefault(key, {"count": 0, "ep": ep_name, "name": inst})
        except Exception:
            pass

    # Pipeline metadata fetches grouped by endpoint to avoid N round-trips
    ep_instances: dict[str, list[tuple[str, str]]] = {}  # ep_name → [(key, inst)]
    for key, info in all_instances.items():
        ep_instances.setdefault(info["ep"], []).append((key, info["name"]))

    meta_cache: dict[str, dict] = {}
    for ep_name, pairs in ep_instances.items():
        try:
            rc = r_for(ep_name)
            pipe = rc.pipeline(transaction=False)
            for _, inst in pairs:
                pipe.get(f"rag_meta:{inst}")
            raws = pipe.execute()
            for (key, _), raw in zip(pairs, raws):
                meta_cache[key] = json.loads(raw) if raw else {}
        except Exception:
            for key, _ in pairs:
                meta_cache[key] = {}

    result = []
    for key, info in all_instances.items():
        ep_name, inst = info["ep"], info["name"]
        meta = meta_cache.get(key, {})
        resolved_ep = meta.get("redis_endpoint", ep_name)
        # Unknown endpoints (not yet probed) are treated as reachable
        reachable = _endpoint_health.get(resolved_ep, True)
        result.append({
            "name":            inst,
            "chunks":          info["count"],
            "color":           meta.get("color",   "#6366f1"),
            "tags":            meta.get("tags",    []),
            "created":         meta.get("created", ""),
            "enabled":         meta.get("enabled", True),
            "redis_endpoint":  resolved_ep,
            "reachable":       reachable,
        })
    return result


def reset_rag(instance: str, rc: redis.Redis | None = None):
    """
    Delete all chunk keys and the FT index for a RAG instance.

    Uses FT.DROPINDEX with the DD (Delete Documents) flag which atomically
    drops the index AND removes all indexed HASH keys in a single command —
    far more efficient than the previous KEYS scan + bulk DELETE approach.
    The chunk counter key is cleaned up separately.
    """
    rc = rc or r()
    prefix = rag_prefix(instance)
    _index_ensured.discard(instance)   # force re-creation on next ingest
    try:
        rc.execute_command("FT.DROPINDEX", f"{prefix}:idx", "DD")
    except Exception:
        pass  # index may not exist yet — that's fine
    # Belt-and-suspenders: explicitly delete any remaining chunk HASH keys.
    # FT.DROPINDEX DD may silently fail (e.g. index never existed), leaving
    # chunk keys behind so the instance keeps reappearing in list scans.
    batch: list = []
    for k in rc.scan_iter(f"{prefix}:chunk:*", count=500):
        batch.append(k)
        if len(batch) >= 500:
            rc.delete(*batch)
            batch = []
    if batch:
        rc.delete(*batch)
    # Remove counter, chunk hash dedup set, and URL skip list in one call
    rc.delete(f"rag:{instance}:chunk_counter", f"rag:{instance}:chunk_hashes", f"rag:{instance}:indexed_urls")
    append_log({
        "ts": datetime.now(timezone.utc).isoformat(),
        "instance": instance,
        "source": "RESET",
        "chunks": 0,
        "status": "reset",
    })


# ── rag_meta read-through cache ──────────────────────────────────────────────
# rag_meta:{instance} = {"enabled": bool, "redis_endpoint": str, ...}. It is read on
# every chat turn — to pick the owning endpoint and check the enabled flag — but written
# only when an instance is created, toggled or deleted. Reading it fresh each turn cost
# up to two Redis round trips per instance (rc_for_instance re-read the same key), which
# is pure latency against a remote Redis. This short-TTL in-process cache turns the hot
# path into a dict lookup; writes invalidate explicitly, so the TTL only bounds staleness
# in the (unexpected) event an invalidation is ever missed.
_RAG_META_TTL = 3.0                       # seconds
_rag_meta_cache: dict[str, tuple[float, tuple[dict | None, str]]] = {}
# The cache is read from the event loop AND from to_thread / FastAPI-threadpool workers,
# and writes to rag_meta invalidate it. A generation counter guarded by a lock closes the
# resolve-then-store race: if an invalidation lands while a resolve is in flight, the stale
# value is not cached (we re-resolve on the next read instead).
_rag_meta_lock = threading.Lock()
_rag_meta_gen = 0

def _resolve_rag_meta(instance: str) -> tuple[dict | None, str]:
    """Find an instance's rag_meta across endpoints. Returns (meta_or_None, endpoint_name).
    One GET on the default endpoint in the common case; extra endpoints are consulted only
    when the instance is not on the default. Mirrors the old rc_for_instance search order."""
    try:
        meta_raw = r().get(f"rag_meta:{instance}")
        if meta_raw:
            meta = json.loads(meta_raw)
            return meta, meta.get("redis_endpoint", "default")
    except Exception:
        pass
    for ep in _config.get("redis_endpoints", []):
        ep_name = ep.get("name", "")
        if ep_name and ep_name != "default":
            try:
                meta_raw = r_for(ep_name).get(f"rag_meta:{instance}")
                if meta_raw:
                    return json.loads(meta_raw), ep_name
            except Exception:
                pass
    return None, "default"

def _rag_meta_cached(instance: str) -> tuple[dict | None, str]:
    """Cached (meta, endpoint) for an instance. Cache hit performs no Redis I/O."""
    ent = _rag_meta_cache.get(instance)
    if ent and (time.time() - ent[0]) < _RAG_META_TTL:
        return ent[1]
    with _rag_meta_lock:
        gen_before = _rag_meta_gen
    val = _resolve_rag_meta(instance)                 # Redis I/O outside the lock
    with _rag_meta_lock:
        if _rag_meta_gen == gen_before:               # no invalidation raced us
            _rag_meta_cache[instance] = (time.time(), val)
    return val

async def _rag_meta_cached_async(instance: str) -> tuple[dict | None, str]:
    """Same as _rag_meta_cached, but a cache miss resolves off the event loop."""
    ent = _rag_meta_cache.get(instance)
    if ent and (time.time() - ent[0]) < _RAG_META_TTL:
        return ent[1]
    return await asyncio.to_thread(_rag_meta_cached, instance)

def invalidate_rag_meta(instance: str | None = None):
    """Drop cached rag_meta after a write so the next read reflects it immediately."""
    global _rag_meta_gen
    with _rag_meta_lock:
        _rag_meta_gen += 1                             # cancels any in-flight resolve's store
        if instance is None:
            _rag_meta_cache.clear()
        else:
            _rag_meta_cache.pop(instance, None)

def rc_for_instance(instance: str) -> redis.Redis:
    """
    Return the Redis client that owns a specific RAG instance.
    Endpoint is resolved via the cached rag_meta (default endpoint first, then extras),
    falling back to the default client if the instance is not found anywhere.
    """
    _meta, ep = _rag_meta_cached(instance)
    return r_for(ep)

def _rc_for(instance: str, endpoint: str | None = None) -> redis.Redis:
    """
    Return the Redis client for an instance.
    When ``endpoint`` is supplied explicitly (e.g. from a query parameter) it
    is used directly, bypassing the metadata lookup.  This lets callers target
    the correct server even when two instances share the same name on different
    endpoints.
    """
    if endpoint:
        return r_for(endpoint)
    return rc_for_instance(instance)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# WEBSOCKET MANAGER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class ConnManager:
    """Tracks active WebSocket connections by session ID."""

    def __init__(self):
        self.active: dict[str, WebSocket] = {}

    async def connect(self, ws: WebSocket, sid: str):
        await ws.accept()
        self.active[sid] = ws

    def disconnect(self, sid: str):
        self.active.pop(sid, None)

    async def send(self, sid: str, data: dict):
        ws = self.active.get(sid)
        if ws:
            await ws.send_json(data)


mgr = ConnManager()

async def _recrawl_loop():
    """
    Background task: periodically re-crawl all scheduled web sources.
    Runs every 60 seconds internally; actual crawls only trigger when
    ``now - last_crawled >= interval_minutes * 60``.
    """
    while True:
        await asyncio.sleep(60)
        if not _config.get("recrawl", {}).get("enabled", False):
            continue
        interval_secs = int(_config.get("recrawl", {}).get("interval_minutes", 60)) * 60
        now = time.time()
        scheduled = _config.get("scheduled_sources", [])
        if not scheduled:
            continue
        changed = False
        for src in scheduled:
            last = float(src.get("last_crawled", 0))
            if now - last < interval_secs:
                continue
            url      = src.get("url", "").strip()
            instance = src.get("instance", "default")
            depth    = int(src.get("depth", 0))
            if not url:
                continue
            log.info(f"Recrawl scheduler: crawling {url} → instance '{instance}'")
            try:
                rc = rc_for_instance(instance)
                await crawl_url(instance, url, depth, rc=rc)
                src["last_crawled"] = now
                changed = True
            except Exception as e:
                log.warning(f"Recrawl failed for {url}: {e}")
        if changed:
            save_config(_config)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STARTUP
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _primary_lan_ip() -> str | None:
    """Best-effort primary LAN IPv4 (the interface used for outbound traffic).
    Uses a UDP socket's chosen source address — no packet is actually sent."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        return ip if not ip.startswith("127.") else None
    except Exception:
        return None
    finally:
        s.close()


@app.on_event("startup")
async def startup():
    """
    Called once when Uvicorn starts.
    - Loads config from disk (merged with defaults)
    - Reads API keys from environment (they override config-file keys)
    - Connects to Redis and warms up the embedding model
    - Prints diagnostic info to stdout
    """
    global _config, _env_key, _openai_env_key, _qwen_env_key, _mistral_env_key, _groq_env_key, _gemini_env_key

    _config = load_config()
    load_logs()

    log.info("=" * 60)
    log.info(f"RediRecall v{__version__} — startup diagnostics")
    log.info(f"  Config:   {CONFIG_PATH} ({'found' if CONFIG_PATH.exists() else 'NOT FOUND — using defaults'})")
    log.info(f"  Provider: {_config.get('provider', 'ollama')}")

    # ── Access URLs ────────────────────────────────────────────────────────
    # Show every URL the UI can be reached at, so the user can copy one from the
    # console. Host/port mirror what cli() binds (REDIRECALL_HOST / REDIRECALL_PORT);
    # a raw `uvicorn --host/--port` bypasses these env vars, so the URLs below assume
    # the standard `redirecall` entrypoint.
    _bind_host = os.environ.get("REDIRECALL_HOST", "127.0.0.1")
    _port      = os.environ.get("REDIRECALL_PORT", "8420")
    _in_docker = os.path.exists("/.dockerenv")
    _urls: list[str] = []
    if _bind_host in ("0.0.0.0", "::", ""):
        # Bound to all interfaces — reachable on loopback AND the LAN.
        _urls.append(f"http://localhost:{_port}")
        # Inside a container the detected IP is the container's, not the host's, and the
        # published-port mapping is unknown here — so don't print a misleading LAN URL.
        _lan_ip = None if _in_docker else _primary_lan_ip()
        if _lan_ip:
            _urls.append(f"http://{_lan_ip}:{_port}  (LAN — no built-in auth; trusted networks only)")
        elif _in_docker:
            _urls.append(f"(in Docker: reach it at http://<host>:<published-port>)")
    else:
        _urls.append(f"http://{_bind_host}:{_port}")
        if _bind_host == "127.0.0.1":
            _urls.append(f"http://localhost:{_port}")
    log.info(f"  Access:   {_urls[0]}")
    for _u in _urls[1:]:
        log.info(f"            {_u}")

    # ── Anthropic Claude API key ───────────────────────────────────────────
    _env_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    stored_claude = _config.get("claude", {}).get("api_key", "")
    if _env_key:
        _config.setdefault("claude", {})["api_key"] = _env_key
        masked = _env_key[:8] + "…" + _env_key[-4:] if len(_env_key) > 12 else "***"
        log.info(f"  ANTHROPIC_API_KEY: env ({masked}) — takes precedence over config.json")
    elif stored_claude:
        log.info("  ANTHROPIC_API_KEY: from config.json")
    else:
        log.info("  ANTHROPIC_API_KEY: not set — Claude unavailable")

    # ── OpenAI API key ─────────────────────────────────────────────────────
    _openai_env_key = os.environ.get("OPENAI_API_KEY", "").strip()
    stored_openai = _config.get("openai", {}).get("api_key", "")
    if _openai_env_key:
        _config.setdefault("openai", {})["api_key"] = _openai_env_key
        masked = _openai_env_key[:8] + "…" + _openai_env_key[-4:] if len(_openai_env_key) > 12 else "***"
        log.info(f"  OPENAI_API_KEY:    env ({masked}) — takes precedence over config.json")
    elif stored_openai:
        log.info("  OPENAI_API_KEY:    from config.json")
    else:
        log.info("  OPENAI_API_KEY:    not set — OpenAI unavailable")

    # ── Qwen / DashScope API key ───────────────────────────────────────────
    _qwen_env_key = os.environ.get("DASHSCOPE_API_KEY", "").strip()
    stored_qwen = _config.get("qwen", {}).get("api_key", "")
    if _qwen_env_key:
        _config.setdefault("qwen", {})["api_key"] = _qwen_env_key
        masked = _qwen_env_key[:8] + "…" + _qwen_env_key[-4:] if len(_qwen_env_key) > 12 else "***"
        log.info(f"  DASHSCOPE_API_KEY: env ({masked}) — takes precedence over config.json")
    elif stored_qwen:
        log.info("  DASHSCOPE_API_KEY: from config.json")
    else:
        log.info("  DASHSCOPE_API_KEY: not set — Qwen unavailable")

    # ── Mistral API key ────────────────────────────────────────────────────
    _mistral_env_key = os.environ.get("MISTRAL_API_KEY", "").strip()
    stored_mistral = _config.get("mistral", {}).get("api_key", "")
    if _mistral_env_key:
        _config.setdefault("mistral", {})["api_key"] = _mistral_env_key
        masked = _mistral_env_key[:8] + "…" + _mistral_env_key[-4:] if len(_mistral_env_key) > 12 else "***"
        log.info(f"  MISTRAL_API_KEY:   env ({masked}) — takes precedence over config.json")
    elif stored_mistral:
        log.info("  MISTRAL_API_KEY:   from config.json")
    else:
        log.info("  MISTRAL_API_KEY:   not set — Mistral unavailable")

    # ── Groq API key ───────────────────────────────────────────────────────
    _groq_env_key = os.environ.get("GROQ_API_KEY", "").strip()
    stored_groq = _config.get("groq", {}).get("api_key", "")
    if _groq_env_key:
        _config.setdefault("groq", {})["api_key"] = _groq_env_key
        masked = _groq_env_key[:8] + "…" + _groq_env_key[-4:] if len(_groq_env_key) > 12 else "***"
        log.info(f"  GROQ_API_KEY:      env ({masked}) — takes precedence over config.json")
    elif stored_groq:
        log.info("  GROQ_API_KEY:      from config.json")
    else:
        log.info("  GROQ_API_KEY:      not set — Groq unavailable")

    # ── Google Gemini API key ──────────────────────────────────────────────
    _gemini_env_key = os.environ.get("GEMINI_API_KEY", "").strip()
    stored_gemini = _config.get("gemini", {}).get("api_key", "")
    if _gemini_env_key:
        _config.setdefault("gemini", {})["api_key"] = _gemini_env_key
        masked = _gemini_env_key[:8] + "…" + _gemini_env_key[-4:] if len(_gemini_env_key) > 12 else "***"
        log.info(f"  GEMINI_API_KEY:    env ({masked}) — takes precedence over config.json")
    elif stored_gemini:
        log.info("  GEMINI_API_KEY:    from config.json")
    else:
        log.info("  GEMINI_API_KEY:    not set — Gemini unavailable")

    # Redis host/port env overrides — used by the Docker image to point at the
    # `redis` compose service without editing config.json.
    _redis_host_env = os.environ.get("REDIRECALL_REDIS_HOST", "").strip()
    _redis_port_env = os.environ.get("REDIRECALL_REDIS_PORT", "").strip()
    if _redis_host_env:
        _config.setdefault("redis", {})["host"] = _redis_host_env
        log.info(f"  REDIS host:        env ({_redis_host_env})")
    if _redis_port_env:
        _config.setdefault("redis", {})["port"] = int(_redis_port_env)
        log.info(f"  REDIS port:        env ({_redis_port_env})")

    rc = _config.get("redis", {})
    log.info(f"  Ollama:   {_config.get('ollama', {}).get('host')}:{_config.get('ollama', {}).get('port')}")
    log.info(f"  Redis:    {rc.get('host')}:{rc.get('port')}  db={rc.get('db', 0)}")

    extra = _config.get("redis_endpoints", [])
    if extra:
        log.info(f"  Extra Redis endpoints: {[e.get('name') for e in extra]}")
    log.info("=" * 60)

    # ── All slow work runs in background so startup returns immediately ────
    # SentenceTransformer model loading and Redis SCAN over all endpoints
    # can both take 5-60 s.  Doing them here would block Uvicorn from
    # accepting connections, causing a blank browser until they finish.
    async def _bg_init():
        # 1. Probe all endpoints — fast PING, marks unreachable ones offline
        await asyncio.to_thread(refresh_endpoint_health)
        log.info(f"  Endpoint health: { {k: ('OK' if v else 'OFFLINE') for k, v in _endpoint_health.items()} }")

        # 2. Warm up embedding model (5-15 s, CPU-bound)
        try:
            await asyncio.to_thread(get_embed_model)
            log.info(f"  Embedding model ready: {_embed_model_name}")
        except Exception as e:
            log.warning(f"  Embedding model warmup failed: {e}")

        # 2b. Warm the semantic cache and (if enabled) the cross-encoder reranker.
        # Both used to initialise lazily inside handle_chat, ON the event loop, so the
        # first question after a restart stalled seconds while the whole server froze.
        try:
            await asyncio.to_thread(_get_semantic_cache)
            log.info("  Semantic cache ready")
        except Exception as e:
            log.warning(f"  Semantic cache warmup failed: {e}")
        finally:
            # Warm done (success or not): from here cache_lookup/store use the cache
            # normally instead of skipping to avoid the inline build.
            global _semantic_cache_ready
            _semantic_cache_ready = True
        if _config.get("reranker", {}).get("enabled", False):
            try:
                await asyncio.to_thread(get_reranker)
                log.info("  Reranker ready")
            except Exception as e:
                log.warning(f"  Reranker warmup failed: {e}")

        # 3. Ensure FT indexes only on reachable endpoints
        try:
            instances = await asyncio.to_thread(list_rag_instances)
            for inst in instances:
                inst_name = inst.get("name", "")
                ep_name   = inst.get("redis_endpoint", "default")
                if inst_name and _endpoint_health.get(ep_name, True):
                    try:
                        rc_inst = rc_for_instance(inst_name)
                        _index_ensured.discard(inst_name)
                        await asyncio.to_thread(ensure_rag_index, inst_name, rc_inst)
                    except Exception as idx_err:
                        log.warning(f"  Could not ensure index for '{inst_name}': {idx_err}")
        except Exception:
            pass

    asyncio.create_task(_bg_init())

    # ── Start background recrawl scheduler ────────────────────────────────
    global _recrawl_task
    _recrawl_task = asyncio.create_task(_recrawl_loop())
    log.info("  Recrawl scheduler: started")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — CONFIG
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/config")
def api_get_config():
    """
    Return the full runtime config.
    Includes boolean flags indicating whether API keys came from env vars
    so the UI can show the correct hint without ever seeing the actual keys.
    """
    return {
        **_redact_secrets(_config),
        # Read-only: lets the UI offer "Reset to default" for the base instruction,
        # so an improved shipped default can be adopted after it has been saved once.
        "base_instruction_default": DEFAULT_BASE_INSTRUCTION,
        "claude_key_from_env":  bool(_env_key),
        "openai_key_from_env":  bool(_openai_env_key),
        "qwen_key_from_env":    bool(_qwen_env_key),
        "mistral_key_from_env": bool(_mistral_env_key),
        "groq_key_from_env":    bool(_groq_env_key),
        "gemini_key_from_env":  bool(_gemini_env_key),
    }


@app.post("/api/config")
async def api_save_config(payload: dict):
    """
    Save updated config.
    After merging the payload, env-sourced API keys are always restored so a
    settings save never loses a key that was loaded from the environment.
    """
    global _config, _embed_model
    # Restore any secrets the UI posted back as sentinels before applying,
    # so a settings save never wipes a stored key/password it never received.
    _unredact_secrets(payload, _config)
    _config.update(payload)

    # Restore env-sourced keys — they must never be overwritten by a config save
    if _env_key:
        _config.setdefault("claude", {})["api_key"] = _env_key
    if _openai_env_key:
        _config.setdefault("openai", {})["api_key"] = _openai_env_key
    if _qwen_env_key:
        _config.setdefault("qwen", {})["api_key"] = _qwen_env_key
    if _mistral_env_key:
        _config.setdefault("mistral", {})["api_key"] = _mistral_env_key
    if _groq_env_key:
        _config.setdefault("groq", {})["api_key"] = _groq_env_key
    if _gemini_env_key:
        _config.setdefault("gemini", {})["api_key"] = _gemini_env_key

    save_config(_config)      # strips env keys before writing to disk
    invalidate_redis_clients()
    invalidate_provider_clients()   # keys/base URLs may have changed
    _embed_model = None       # force model reload on next request

    log.info(
        f"Config saved — provider={_config.get('provider')} "
        f"claude={'env' if _env_key else ('set' if _config.get('claude',{}).get('api_key') else 'empty')} "
        f"openai={'env' if _openai_env_key else ('set' if _config.get('openai',{}).get('api_key') else 'empty')}"
    )

    # Warm up connections in background — don't block the event loop
    async def _rewarm():
        try:
            await asyncio.to_thread(get_redis)
            await asyncio.to_thread(get_embed_model)
        except Exception:
            pass
    asyncio.create_task(_rewarm())

    return {"ok": True}


@app.get("/api/config/export")
def api_export_config():
    """Export current settings as a JSON download, with all secrets redacted.

    Secrets are replaced by a sentinel rather than the raw file so an export can
    be shared/backed up without leaking keys; re-importing keeps existing keys.
    """
    return JSONResponse(
        content=_redact_secrets(_config),
        headers={"Content-Disposition": 'attachment; filename="redirecall_config.json"'},
    )


@app.post("/api/config/import")
async def api_import_config(file: UploadFile = File(...)):
    global _config
    content = await file.read()
    cfg = json.loads(content)
    # Sentinel secrets in an exported file mean "keep what's already stored".
    _unredact_secrets(cfg, _config)
    _config = {**DEFAULT_CONFIG, **cfg}
    save_config(_config)
    return {"ok": True}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — STATUS / HEALTH CHECKS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/health")
def api_health():
    """Lightweight liveness probe (no I/O) for the Docker HEALTHCHECK and run.sh.

    Always returns 200 while the process is up. The ``services`` booleans report
    which optional capabilities were importable at startup so run.sh can show
    them in its status banner. Redis reachability lives on /api/status/redis
    (kept out of here so this probe never blocks on a slow connection).
    """
    return {
        "status":   "ok",
        "app":      "RediRecall",
        "version":  __version__,
        # AGPL-3.0 §13: users interacting with this instance over a network must be
        # offered the Corresponding Source. The UI links this too; keeping it here
        # makes the offer available to any client, not just the browser UI.
        "license":  "AGPL-3.0-or-later",
        "source":   SOURCE_URL,
        "provider": _config.get("provider", ""),
        "services": {
            "pdf":         HAS_PYMUPDF,
            "docx":        HAS_PYTHON_DOCX,
            "xlsx":        HAS_OPENPYXL,
            "web_extract": HAS_TRAFILATURA,
            "html_parse":  HAS_BS4,
            "reranker":    HAS_CROSSENCODER,
            "js_crawl":    HAS_CRAWL4AI,
            "anthropic":   _ANTHROPIC_AVAILABLE,
            "openai_sdk":  _OPENAI_SDK_AVAILABLE,
            "gemini_sdk":  _GENAI_AVAILABLE,
        },
    }


def probe_search(rc: redis.Redis, ep_name: str = "default") -> bool:
    """
    Check whether a Redis connection has the Search module (RediSearch) loaded.

    The result is cached in ``_search_available`` so we only probe once per
    endpoint per server lifetime — subsequent calls return instantly.

    Returns True if FT commands are available, False otherwise.
    """
    global _search_available
    cached = _search_available.get(ep_name)
    if cached is not None:
        return cached
    try:
        # FT.INFO on a non-existent index raises "Unknown index name" (not
        # "unknown command") when Search IS available — so any response other
        # than an "unknown command" error means Search is present.
        rc.execute_command("FT.INFO", "__probe__")
        _search_available[ep_name] = True
    except Exception as e:
        err = str(e).lower()
        if "unknown command" in err or "unknown subcommand" in err:
            _search_available[ep_name] = False
        else:
            # "Unknown index name" or similar — Search is available, index just doesn't exist
            _search_available[ep_name] = True
    return _search_available[ep_name]


@app.get("/api/status/redis")
def api_redis_status():
    """Ping the primary Redis, refresh endpoint health, and return server info."""
    # Refresh health for all endpoints on every status poll (called every 30 s by the UI)
    refresh_endpoint_health()
    try:
        rc = get_redis()
        info = rc.info()
        return {
            "ok": True,
            "version":           info.get("redis_version"),
            "memory_used":       info.get("used_memory_human"),
            "memory_peak":       info.get("used_memory_peak_human"),
            "connected_clients": info.get("connected_clients"),
            "uptime_days":       info.get("uptime_in_days"),
            "mode":              info.get("redis_mode", "standalone"),
            "search_available":  probe_search(rc, "default"),
            "endpoint_health":   dict(_endpoint_health),
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/redis/{endpoint_name}")
def api_redis_endpoint_status(endpoint_name: str):
    """Ping a named Redis endpoint and return Search module availability."""
    try:
        rc = r_for(endpoint_name)
        info = rc.info()
        return {
            "ok":               True,
            "version":          info.get("redis_version"),
            "memory_used":      info.get("used_memory_human"),
            "search_available": probe_search(rc, endpoint_name),
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/ollama")
async def api_ollama_status():
    """Ping the Ollama server (reuses the cached model list to avoid a redundant /api/tags call)."""
    try:
        await ollama_models()
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/claude")
async def api_claude_status(key: str | None = None):
    """
    Verify the Claude API key. Accepts an optional ?key= param so the UI
    can test an unsaved key without requiring a save-first workflow.
    """
    # The Settings form pre-fills an already-saved key's field with
    # _SECRET_SENTINEL (the redacted placeholder from GET /api/config), and the
    # Test button sends whatever's in the field. Without this check, testing an
    # already-saved key would send the literal sentinel string as the API key
    # instead of falling back to the real one below.
    if key == _SECRET_SENTINEL:
        key = None
    api_key  = key or _config.get("claude", {}).get("api_key", "")
    base_url = _config.get("claude", {}).get("base_url", "https://api.anthropic.com").rstrip("/")
    if not api_key:
        return {"ok": False, "error": "No API key configured"}
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            res = await client.post(
                f"{base_url}/v1/messages/count_tokens",
                json={"model": "claude-haiku-4-5-20251001", "messages": [{"role": "user", "content": "hi"}]},
                headers={"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            )
            if res.status_code == 200:
                return {"ok": True}
            if res.status_code == 401:
                return {"ok": False, "error": res.json().get("error", {}).get("message", "Unauthorized")}
            # Fallback
            res2 = await client.get(
                f"{base_url}/v1/models",
                headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"},
            )
            if res2.status_code == 200:
                return {"ok": True}
            return {"ok": False, "error": res2.json().get("error", {}).get("message", f"HTTP {res2.status_code}")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/openai")
async def api_openai_status(key: str | None = None):
    """Verify the OpenAI API key using the native SDK. Accepts optional ?key= for testing unsaved keys."""
    if key == _SECRET_SENTINEL:
        key = None
    api_key  = key or _config.get("openai", {}).get("api_key", "")
    base_url = _config.get("openai", {}).get("base_url", "https://api.openai.com").rstrip("/")
    if not api_key:
        return {"ok": False, "error": "No API key configured"}
    if not _OPENAI_SDK_AVAILABLE:
        return {"ok": False, "error": "openai package not installed. Run: pip install openai"}
    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        await client.models.list()
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/qwen")
async def api_qwen_status(key: str | None = None):
    """Verify the Qwen API key. Accepts optional ?key= for testing unsaved keys."""
    if key == _SECRET_SENTINEL:
        key = None
    api_key  = key or _config.get("qwen", {}).get("api_key", "")
    base_url = _config.get("qwen", {}).get("base_url",
                           "https://dashscope.aliyuncs.com/compatible-mode/v1").rstrip("/")
    if not api_key:
        return {"ok": False, "error": "No API key configured. Get a free key at qwen.ai"}
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            res = await client.get(
                f"{base_url}/models",
                headers={"Authorization": f"Bearer {api_key}"},
            )
            if res.status_code == 200:
                return {"ok": True}
            body = res.json()
            return {"ok": False, "error": body.get("message", f"HTTP {res.status_code}")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/mistral")
async def api_mistral_status(key: str | None = None, probe: bool = False):
    """Verify the Mistral API key. Accepts optional ?key= for testing unsaved keys.
    The cheap check is GET /models; some free-tier keys are denied that specific
    endpoint (401) while chat completions work fine, so a bare /models 401 is not
    proof the key itself is invalid. The periodic background poll (checkCloudStatus,
    every 5 min) relies on the cheap check alone and accepts that rare imprecision.
    ?probe=1 — sent only by the user-initiated Settings "Test" button — confirms a
    /models failure with one minimal real chat-completion call before reporting the
    key as invalid."""
    if key == _SECRET_SENTINEL:
        key = None
    api_key  = key or _config.get("mistral", {}).get("api_key", "")
    base_url = _config.get("mistral", {}).get("base_url", "https://api.mistral.ai/v1").rstrip("/")
    if not api_key:
        return {"ok": False, "error": "No API key configured. Get a free key at console.mistral.ai"}
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            res = await client.get(
                f"{base_url}/models",
                headers={"Authorization": f"Bearer {api_key}"},
            )
            if res.status_code == 200:
                return {"ok": True}
            try:
                body = res.json()
                msg = body.get("message") or body.get("error") or f"HTTP {res.status_code}"
            except Exception:
                msg = f"HTTP {res.status_code}"
            if not probe or not _OPENAI_SDK_AVAILABLE:
                return {"ok": False, "error": msg}
            try:
                model = _config.get("mistral", {}).get("model") or MISTRAL_MODELS_STATIC[0]["id"]
                client2 = _cached_client("openai", api_key, base_url)
                await client2.chat.completions.create(
                    model=model, messages=[{"role": "user", "content": "hi"}], max_tokens=1,
                )
                return {"ok": True}
            except Exception as e2:
                return {"ok": False, "error": f"Mistral error: {e2}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/groq")
async def api_groq_status(key: str | None = None):
    """Verify the Groq API key using the openai SDK. Accepts optional ?key= for testing unsaved keys."""
    if key == _SECRET_SENTINEL:
        key = None
    api_key  = key or _config.get("groq", {}).get("api_key", "")
    base_url = _config.get("groq", {}).get("base_url", "https://api.groq.com/openai").rstrip("/")
    if not api_key:
        return {"ok": False, "error": "No API key configured. Get a free key at console.groq.com"}
    if not _OPENAI_SDK_AVAILABLE:
        return {"ok": False, "error": "openai package not installed. Run: pip install openai"}
    try:
        client = _cached_client("openai", api_key, f"{base_url}/v1")
        await client.models.list()
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/status/gemini")
async def api_gemini_status(key: str | None = None, probe: bool = False):
    """Verify the Gemini API key using the native SDK. Accepts optional ?key= for
    unsaved keys. The cheap check is models.list(); some keys are denied that
    specific method (403 PERMISSION_DENIED) while generateContent still works
    fine, so a bare list() failure is not proof the key itself is invalid. The
    periodic background poll (checkCloudStatus, every 5 min) relies on the cheap
    check alone and accepts that rare imprecision. ?probe=1 — sent only by the
    user-initiated Settings "Test" button — confirms a list() failure with one
    minimal real generateContent call before reporting the key as invalid."""
    if key == _SECRET_SENTINEL:
        key = None
    api_key = key or _config.get("gemini", {}).get("api_key", "")
    if not api_key:
        return {"ok": False, "error": "No API key configured. Get a free key at aistudio.google.com"}
    if not _GENAI_AVAILABLE:
        return {"ok": False, "error": "google-genai not installed. Run: pip install google-genai"}
    client = _google_genai.Client(api_key=api_key)
    try:
        await asyncio.get_event_loop().run_in_executor(
            None, lambda: next(iter(client.models.list()), None))
        return {"ok": True}
    except Exception as e:
        if not probe:
            return {"ok": False, "error": _gemini_err_msg(e)}
        try:
            model = _config.get("gemini", {}).get("model") or GEMINI_MODELS_STATIC[0]["id"]
            await asyncio.get_event_loop().run_in_executor(
                None, lambda: client.models.generate_content(model=model, contents="hi"))
            return {"ok": True}
        except Exception as e2:
            return {"ok": False, "error": _gemini_err_msg(e2)}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — MODELS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/ollama/models")
async def api_ollama_models():
    return await ollama_models()

@app.get("/api/claude/models")
def api_claude_models():
    return CLAUDE_MODELS

@app.get("/api/openai/models")
async def api_openai_models():
    return await openai_models()

@app.get("/api/qwen/models")
def api_qwen_models():
    return QWEN_MODELS_STATIC

@app.get("/api/mistral/models")
async def api_mistral_models():
    """Fetch live models from Mistral's /v1/models; fall back to the static list."""
    api_key  = _config.get("mistral", {}).get("api_key", "")
    base_url = _config.get("mistral", {}).get("base_url", "https://api.mistral.ai/v1").rstrip("/")
    if not api_key:
        return MISTRAL_MODELS_STATIC
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            res = await client.get(f"{base_url}/models",
                                   headers={"Authorization": f"Bearer {api_key}"})
            if res.status_code == 200:
                data = res.json().get("data", [])
                # Chat-capable models only; keep the ids, sort for stable display.
                ids = sorted({m.get("id") for m in data if m.get("id")})
                if ids:
                    return [{"id": i, "name": i, "context": 0} for i in ids]
    except Exception:
        pass
    return MISTRAL_MODELS_STATIC

@app.get("/api/groq/models")
async def api_groq_models():
    return await groq_models()

@app.get("/api/gemini/models")
async def api_gemini_models():
    return await gemini_models()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — LOCAL FILE IMAGE PROXY
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

import mimetypes
import tempfile

_ALLOWED_IMAGE_DIRS: list[Path] = [
    Path(tempfile.gettempdir()),   # /tmp — tool-generated images usually land here
    Path.cwd(),                    # project working directory
]

@app.get("/api/files/image")
async def api_serve_image(path: str):
    """
    Proxy for locally-generated image files produced by Ollama tools.
    Only serves files that:
      - Exist on disk
      - Have an image extension (png, jpg, gif, webp, bmp)
      - Are located inside an allowed directory (tmp or cwd) to prevent
        arbitrary filesystem access.
    """
    try:
        p = Path(path).resolve()
    except Exception:
        raise HTTPException(400, "Invalid path")

    if p.suffix.lower() not in _IMG_EXTS:
        raise HTTPException(400, "Not an image file")

    if not p.is_file():
        raise HTTPException(404, "File not found")

    # Security: only serve from allowed directories
    allowed = any(
        str(p).startswith(str(d.resolve()))
        for d in _ALLOWED_IMAGE_DIRS
    )
    if not allowed:
        raise HTTPException(403, "Path not in an allowed directory")

    mime, _ = mimetypes.guess_type(str(p))
    mime = mime or "application/octet-stream"
    return FileResponse(str(p), media_type=mime)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — REDIS ENDPOINTS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/redis/endpoints")
def api_list_endpoints():
    """Return all configured Redis endpoints (primary + extras), passwords redacted."""
    def _strip(ep: dict) -> dict:
        e = dict(ep)
        if e.get("password"):
            e["password"] = _SECRET_SENTINEL
        return e
    primary = _strip({**_config.get("redis", {}), "name": "default", "primary": True})
    extras  = [_strip(e) for e in _config.get("redis_endpoints", [])]
    return [primary] + extras


@app.post("/api/redis/endpoints")
async def api_add_endpoint(payload: dict):
    """
    Add a new named Redis endpoint.
    The name must be unique and must not be 'default'.
    """
    name = payload.get("name", "").strip()
    if not name or name == "default":
        raise HTTPException(400, "Endpoint name must be a non-empty string other than 'default'")
    endpoints = _config.get("redis_endpoints", [])
    # A sentinel password means "keep the existing one" (the UI never saw it).
    password = payload.get("password", "")
    if password == _SECRET_SENTINEL:
        existing = next((e for e in endpoints if e.get("name") == name), {})
        password = existing.get("password", "")
    # Upsert: replace if name already exists
    endpoints = [e for e in endpoints if e.get("name") != name]
    endpoints.append({
        "name":     name,
        "host":     payload.get("host", "localhost"),
        "port":     int(payload.get("port", 6379)),
        "db":       int(payload.get("db", 0)),
        "password": password,
        "ssl":      bool(payload.get("ssl", False)),
    })
    _config["redis_endpoints"] = endpoints
    save_config(_config)
    # Invalidate cached client for this endpoint
    _redis_clients.pop(name, None)
    return {"ok": True}


@app.delete("/api/redis/endpoints/{name}")
def api_delete_endpoint(name: str):
    """Remove a named Redis endpoint (cannot delete 'default')."""
    if name == "default":
        raise HTTPException(400, "Cannot delete the default endpoint")
    endpoints = [e for e in _config.get("redis_endpoints", []) if e.get("name") != name]
    _config["redis_endpoints"] = endpoints
    save_config(_config)
    _redis_clients.pop(name, None)
    return {"ok": True}


@app.get("/api/redis/endpoints/{name}/discover")
def api_discover_endpoint(name: str):
    """
    Scan a named Redis endpoint for existing RAG instances.
    Returns a list of {name, chunks, has_meta} dicts so the UI can prompt
    the user to re-register instances found on that server.
    """
    try:
        rc = r_for(name)
        rc.ping()  # fail fast if connection is broken
    except Exception as e:
        raise HTTPException(400, f"Cannot connect to endpoint '{name}': {e}")

    discovered: dict[str, dict] = {}
    try:
        # SCAN can return the same key more than once under concurrent writes, so count
        # DISTINCT chunk keys per instance rather than incrementing on every hit.
        seen_chunks: dict[str, set] = {}
        for k in rc.scan_iter("rag:*:chunk:*", count=500):   # SCAN, not KEYS (O(N) blocks Redis)
            parts = k.decode().split(":")
            if len(parts) >= 3:
                inst = parts[1]
                seen_chunks.setdefault(inst, set()).add(k)
        for inst, keys in seen_chunks.items():
            discovered.setdefault(inst, {"name": inst, "chunks": 0, "has_meta": False})
            discovered[inst]["chunks"] = len(keys)
        for mk in rc.scan_iter("rag_meta:*", count=500):     # SCAN, not KEYS
            inst = mk.decode().replace("rag_meta:", "")
            discovered.setdefault(inst, {"name": inst, "chunks": 0, "has_meta": False})
            discovered[inst]["has_meta"] = True
    except Exception as e:
        raise HTTPException(500, f"Discovery scan failed: {e}")

    return sorted(discovered.values(), key=lambda x: x["name"])


@app.post("/api/redis/endpoints/{name}/register")
def api_register_discovered(name: str, payload: dict):
    """
    Register a list of discovered RAG instances on a named endpoint.
    Ensures each instance has a rag_meta key that references this endpoint
    and that its search index exists.
    """
    instances = payload.get("instances", [])
    if not instances:
        return {"ok": True, "registered": 0}
    try:
        rc = r_for(name)
    except Exception as e:
        raise HTTPException(400, f"Cannot connect to endpoint '{name}': {e}")

    registered = 0
    for inst_name in instances:
        try:
            meta_raw = rc.get(f"rag_meta:{inst_name}")
            meta = json.loads(meta_raw) if meta_raw else {}
            meta["redis_endpoint"] = name
            meta.setdefault("created", datetime.now(timezone.utc).isoformat())
            meta.setdefault("color", "#6366f1")
            meta.setdefault("enabled", True)
            rc.set(f"rag_meta:{inst_name}", json.dumps(meta))
            invalidate_rag_meta(inst_name)
            ensure_rag_index(inst_name, rc)
            registered += 1
        except Exception as e:
            log.warning(f"Failed to register '{inst_name}' on endpoint '{name}': {e}")

    return {"ok": True, "registered": registered}


@app.post("/api/redis/test")
async def api_test_redis_adhoc(payload: dict):
    """
    Test a Redis connection using parameters supplied in the request body.
    Does NOT use saved config — validates form-field values before saving.
    Body: {host, port, db, password, ssl}
    """
    try:
        host = payload.get("host", "localhost")
        port = int(payload.get("port", 6379))
        password = payload.get("password", "")
        # The UI never receives real passwords (they're redacted to a sentinel).
        # If it echoes the sentinel back, resolve it against the stored config
        # for the matching host:port so "Test" works without re-typing the key.
        if password == _SECRET_SENTINEL:
            candidates = [_config.get("redis", {})] + list(_config.get("redis_endpoints", []))
            match = next(
                (e for e in candidates
                 if e.get("host") == host and int(e.get("port", 6379)) == port),
                {},
            )
            password = match.get("password", "")
        rc = redis.Redis(
            host=host,
            port=port,
            db=int(payload.get("db", 0)),
            password=password or None,
            ssl=bool(payload.get("ssl", False)),
            socket_connect_timeout=3,
            socket_timeout=3,
            decode_responses=False,
        )
        info = await asyncio.to_thread(rc.info)   # keep the sync client off the event loop
        # Use a unique key so the probe result doesn't pollute the saved-endpoint cache
        probe_key = f"__adhoc_{payload.get('host')}:{payload.get('port')}"
        search_ok = await asyncio.to_thread(probe_search, rc, probe_key)
        # Don't persist this ad-hoc probe result permanently
        _search_available.pop(probe_key, None)
        return {
            "ok":               True,
            "version":          info.get("redis_version"),
            "memory_used":      info.get("used_memory_human"),
            "connected_clients":info.get("connected_clients"),
            "mode":             info.get("redis_mode", "standalone"),
            "search_available": search_ok,
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/api/redis/endpoints/{name}/test")
async def api_test_endpoint(name: str):
    """Test connectivity to a named Redis endpoint (uses saved config)."""
    try:
        rc = r_for(name)
        info = await asyncio.to_thread(rc.info)
        search_ok = await asyncio.to_thread(probe_search, rc, name)
        return {
            "ok":               True,
            "version":          info.get("redis_version"),
            "memory":           info.get("used_memory_human"),
            "search_available": search_ok,
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — CACHE ANALYTICS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/cache/stats")
def api_cache_stats():
    """
    Return a summary of entries in the SemanticCache.
    redisvl stores the prompt in a field named 'prompt' (previously 'query').
    """
    try:
        rc   = r()
        keys = list(rc.scan_iter(f"{CACHE_PREFIX}*", count=200))
        entries = []
        pipe = rc.pipeline(transaction=False)
        for k in keys:
            pipe.hget(k, "prompt")
            pipe.ttl(k)
        results = pipe.execute()
        for i in range(0, len(results), 2):
            prompt_raw = results[i]
            ttl        = results[i + 1]
            prompt     = prompt_raw.decode() if isinstance(prompt_raw, bytes) else (prompt_raw or "")
            entries.append({"query": prompt, "ttl": ttl})
        return {"count": len(keys), "entries": entries}
    except Exception as e:
        return {"count": 0, "entries": [], "error": str(e)}


@app.delete("/api/cache")
def api_cache_clear():
    """Clear all semantic cache entries. Uses SemanticCache.clear() when available."""
    global _semantic_cache
    cache = _get_semantic_cache()
    if cache is not None:
        try:
            cache.clear()
            return {"deleted": "all"}
        except Exception:
            pass  # fall through to manual key deletion
    rc = r()
    deleted = 0
    batch: list = []
    for k in rc.scan_iter(f"{CACHE_PREFIX}*", count=500):   # SCAN, not KEYS (O(N) blocks Redis)
        batch.append(k)
        if len(batch) >= 500:
            rc.delete(*batch); deleted += len(batch); batch = []
    if batch:
        rc.delete(*batch); deleted += len(batch)
    _semantic_cache = None  # reset so it re-indexes on next use
    return {"deleted": deleted}


@app.delete("/api/cache/entry")
def api_delete_cache_entry(entry_id: str):
    """
    Delete a single cache entry by its entry_id.
    redisvl SemanticCache stores entries as HASH keys named "{name}:{entry_id}".
    """
    if not entry_id:
        return {"ok": False, "error": "entry_id required"}
    try:
        # SemanticCache name is CACHE_PREFIX stripped of the trailing ":"
        key = f"{CACHE_PREFIX.rstrip(':')}:{entry_id}"
        deleted = r().delete(key)
        return {"ok": bool(deleted)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/rag/stats")
def api_rag_stats():
    """
    Return per-instance query statistics accumulated since the server started.

    Metrics:
      queries            — total number of times this instance was searched
      hits               — queries that returned ≥1 chunk above the threshold
      hit_rate           — hits / queries  (0.0–1.0)
      avg_top_score      — mean cosine similarity of the top-1 chunk on hit queries.
                           Serves as an accuracy proxy: higher = better semantic match.
      avg_best_raw_score — mean cosine similarity of the top-1 KNN result *before*
                           threshold filtering, across all queries.  If this is high
                           but hit_rate is low, your similarity_threshold is too strict.
      avg_chunks         — mean number of chunks returned per query (including misses)
    """
    result = []
    for inst, s in _rag_stats.items():
        q = s["queries"]
        h = s["hits"]
        result.append({
            "name":               inst,
            "queries":            q,
            "hits":               h,
            "misses":             q - h,
            "hit_rate":           round(h / q, 4)                          if q else 0.0,
            "avg_top_score":      round(s["score_sum"] / h, 4)             if h else 0.0,
            "avg_best_raw_score": round(s.get("raw_score_sum", 0.0) / q, 4) if q else 0.0,
            "avg_chunks":         round(s["chunks_total"] / q, 2)          if q else 0.0,
        })
    result.sort(key=lambda x: x["queries"], reverse=True)
    return result


@app.delete("/api/rag/stats")
def api_rag_stats_reset():
    """Reset all per-instance RAG statistics."""
    _rag_stats.clear()
    return {"ok": True}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — RAG INSTANCES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/rag/instances")
def api_rag_instances():
    return list_rag_instances()


@app.post("/api/rag/instances")
async def api_create_instance(payload: dict):
    """
    Create a new RAG instance with optional metadata.
    Accepts redis_endpoint to store this instance on a specific Redis server.
    """
    name     = payload.get("name", f"rag_{uuid.uuid4().hex[:6]}")
    ep_name  = payload.get("redis_endpoint", "default")
    rc       = r_for(ep_name)
    meta = {
        "color":           payload.get("color", "#6366f1"),
        "tags":            payload.get("tags", []),
        "created":         datetime.now(timezone.utc).isoformat(),
        "redis_endpoint":  ep_name,
    }
    await asyncio.to_thread(rc.set, f"rag_meta:{name}", json.dumps(meta))
    invalidate_rag_meta(name)
    await asyncio.to_thread(ensure_rag_index, name, rc)
    return {"name": name, **meta}


@app.delete("/api/rag/instances/{instance}")
def api_delete_instance(instance: str, endpoint: str | None = None):
    rc = _rc_for(instance, endpoint)
    reset_rag(instance, rc)
    rc.delete(f"rag_meta:{instance}")
    invalidate_rag_meta(instance)
    return {"ok": True}


@app.post("/api/rag/{instance}/toggle")
def api_toggle_rag(instance: str, payload: dict, endpoint: str | None = None):
    """Enable or disable a RAG instance without deleting its data."""
    rc       = _rc_for(instance, endpoint)
    meta_raw = rc.get(f"rag_meta:{instance}")
    meta     = json.loads(meta_raw) if meta_raw else {}
    meta["enabled"] = bool(payload.get("enabled", True))
    rc.set(f"rag_meta:{instance}", json.dumps(meta))
    invalidate_rag_meta(instance)
    return {"ok": True, "enabled": meta["enabled"]}


@app.post("/api/rag/{instance}/reset")
def api_reset_rag(instance: str, endpoint: str | None = None):
    reset_rag(instance, _rc_for(instance, endpoint))
    return {"ok": True}


@app.get("/api/rag/{instance}/chunks")
def api_rag_chunks(instance: str, limit: int = 50, endpoint: str | None = None):
    """Return a sample of stored chunks for inspection."""
    rc     = _rc_for(instance, endpoint)
    prefix = rag_prefix(instance)
    # SCAN and stop at `limit` — never pull the whole keyspace into memory the way KEYS did.
    keys = []
    for k in rc.scan_iter(f"{prefix}:chunk:*", count=500):
        keys.append(k)
        if len(keys) >= limit:
            break
    chunks = []
    for k in keys:
        d = rc.hgetall(k)
        chunks.append({
            "id":     d.get(b"chunk_id", b"0").decode(),
            "text":   d.get(b"text", b"").decode()[:200],
            "source": d.get(b"source", b"").decode(),
        })
    return chunks

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — INGESTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/rag/{instance}/ingest/files")
async def api_ingest_files(instance: str, files: list[UploadFile] = File(...), endpoint: str | None = None):
    rc = _rc_for(instance, endpoint)
    results = []
    for f in files:
        dest, safe_name = safe_upload_dest(f.filename)
        dest.write_bytes(await f.read())
        result = await ingest_file(instance, dest, safe_name, rc)
        results.append(result)
        # Remove the uploaded file now that it has been indexed — the content
        # lives in Redis; keeping the file on disk serves no further purpose.
        try:
            dest.unlink(missing_ok=True)
        except Exception as e:
            log.warning(f"Could not delete upload '{dest}': {e}")
    return results


@app.post("/api/rag/{instance}/ingest/files/stream")
async def api_ingest_files_stream(instance: str, files: list[UploadFile] = File(...), endpoint: str | None = None):
    """
    SSE-streaming file ingestion.  Processes files one-by-one and emits a
    progress event after each file so the UI can show a live progress bar.

    Events: {file, status, chunks, error, index, total}
    Final:  {done: true, total: N}
    """
    rc = _rc_for(instance, endpoint)

    # Buffer all uploads to disk first (so we can stream progress after)
    saved: list[tuple[Path, str]] = []
    for f in files:
        dest, safe_name = safe_upload_dest(f.filename)   # strip path + verify containment
        dest.write_bytes(await f.read())
        saved.append((dest, safe_name))

    async def generate():
        for idx, (path, name) in enumerate(saved):
            try:
                result = await ingest_file(instance, path, name, rc)
                yield f"data: {json.dumps({'file': name, 'status': result.get('status','ok'), 'chunks': result.get('chunks', 0), 'index': idx, 'total': len(saved)})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'file': name, 'status': 'error', 'error': str(e), 'index': idx, 'total': len(saved)})}\n\n"
            finally:
                # Always remove the uploaded file after indexing — success or
                # failure — to prevent unbounded growth of the uploads directory.
                try:
                    path.unlink(missing_ok=True)
                except Exception as e:
                    log.warning(f"Could not delete upload '{path}': {e}")
        yield f"data: {json.dumps({'done': True, 'total': len(saved)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/rag/{instance}/optimize")
def api_optimize_rag(instance: str, endpoint: str | None = None):
    """
    Deduplicate chunks in a RAG instance.

    Two chunks are considered exact duplicates when their text is identical
    after case-folding and whitespace normalisation.  The first occurrence
    is kept; all subsequent duplicates are deleted.

    Returns: {removed, remaining, total_before}
    """
    rc = _rc_for(instance, endpoint)
    prefix = rag_prefix(instance)

    seen: dict[str, str] = {}   # hash → first key that owns it
    to_delete: list = []
    total_before = 0

    # Pipeline hget("text") in batches to avoid N round-trips
    _BATCH = 200
    batch_keys: list = []

    def _process_batch(bkeys: list) -> None:
        pipe = rc.pipeline(transaction=False)
        for bk in bkeys:
            pipe.hget(bk, "text")
        for bk, text_raw in zip(bkeys, pipe.execute()):
            if not text_raw:
                continue
            text = text_raw.decode() if isinstance(text_raw, bytes) else text_raw
            normalised = " ".join(text.lower().split())
            h = hashlib.sha256(normalised.encode()).hexdigest()
            key_s = bk.decode() if isinstance(bk, bytes) else bk
            if h in seen:
                to_delete.append(bk)
            else:
                seen[h] = key_s

    for k in rc.scan_iter(f"{prefix}:chunk:*", count=500):
        total_before += 1
        batch_keys.append(k)
        if len(batch_keys) >= _BATCH:
            _process_batch(batch_keys)
            batch_keys = []
    if batch_keys:
        _process_batch(batch_keys)

    if to_delete:
        # Delete in batches to avoid oversized commands
        for i in range(0, len(to_delete), 500):
            rc.delete(*to_delete[i:i + 500])

    return {
        "removed":      len(to_delete),
        "remaining":    total_before - len(to_delete),
        "total_before": total_before,
    }


@app.get("/api/crawl/active")
async def api_crawl_active():
    """Return the state of all currently running (and recently finished) crawls."""
    return list(_active_crawls.values())


@app.post("/api/crawl/cancel")
async def api_crawl_cancel(payload: dict):
    """Cancel a running crawl by seed URL."""
    url  = payload.get("url", "")
    task = _crawl_tasks.get(url)
    if task and not task.done():
        task.cancel()
        # Don't await — return immediately; background task cleans up on its own
    if url in _active_crawls:
        _active_crawls[url]["done"] = True
    return {"ok": True}


@app.post("/api/rag/{instance}/ingest/url")
async def api_ingest_url(instance: str, payload: dict, endpoint: str | None = None):
    """Non-streaming URL ingest (waits for full completion before returning)."""
    url              = payload.get("url", "")
    try:
        assert_public_url(url)          # SSRF guard on the seed URL
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"URL rejected: {e}")
    depth            = int(payload.get("depth", 0))
    respect_robots   = bool(payload.get("respect_robots", True))
    local_only       = bool(payload.get("local_only", True))
    path_prefix_only = bool(payload.get("path_prefix_only", False))
    force_reindex    = bool(payload.get("force_reindex", False))
    max_pages        = int(payload.get("max_pages", 0))
    _crawl_cfg       = _config.get("crawl", {})
    concurrency      = int(payload.get("concurrency",    _crawl_cfg.get("concurrency", 10)))
    js_render        = bool(payload.get("js_render",     _crawl_cfg.get("js_render", False)))
    js_concurrency   = int(payload.get("js_concurrency", _crawl_cfg.get("js_concurrency", 3)))
    smart_mode       = bool(payload.get("smart_mode",    _crawl_cfg.get("smart_mode", True)))
    min_words        = int(payload.get("min_words",      _crawl_cfg.get("min_words", 100)))
    rc               = _rc_for(instance, endpoint)
    results: list    = []

    _active_crawls[url] = {
        "url": url, "instance": instance,
        "pages_done": 0, "chunks": 0, "errors": 0, "blocked": 0, "skipped": 0,
        "start_ts": datetime.now(timezone.utc).isoformat(), "done": False,
    }

    async def cb(u, status, n=0, err="", count=0):
        results.append({"url": u, "status": status, "chunks": n, "error": err, "pages_done": count})
        state = _active_crawls.get(url)
        if state:
            state["pages_done"] = count
            if status == "indexed":   state["chunks"]  += n
            elif status == "error":   state["errors"]  += 1
            elif status == "blocked": state["blocked"] += 1
            elif status == "skipped": state["skipped"] += 1

    task = asyncio.create_task(
        crawl_url(instance, url, depth, progress_cb=cb,
                  respect_robots=respect_robots, local_only=local_only,
                  path_prefix_only=path_prefix_only,
                  max_pages=max_pages, concurrency=concurrency, js_render=js_render,
                  js_concurrency=js_concurrency, smart_mode=smart_mode,
                  min_words=min_words, rc=rc, force_reindex=force_reindex)
    )
    _crawl_tasks[url] = task
    try:
        await task
    finally:
        if url in _active_crawls:
            _active_crawls[url]["done"] = True
    return results


@app.get("/api/rag/{instance}/ingest/url/stream")
async def api_ingest_url_stream(
    instance: str,
    url: str,
    depth: int = 0,
    respect_robots: bool = True,
    local_only: bool = True,
    path_prefix_only: bool = False,
    force_reindex: bool = False,
    max_pages: int = 0,
    concurrency: int = 10,
    js_render: bool = False,
    js_concurrency: int = 3,
    smart_mode: bool = True,
    min_words: int = 100,
    endpoint: str | None = None,
):
    """
    SSE endpoint that streams crawl progress in real-time.
    Each event is a JSON object: {url, status, chunks, error, pages_done}.
    A final {done: true} event signals completion.
    When the client disconnects the server-side crawl task is cancelled.
    """
    queue: asyncio.Queue = asyncio.Queue()
    rc = _rc_for(instance, endpoint)

    try:
        assert_public_url(url)          # SSRF guard on the seed URL
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"URL rejected: {e}")

    _active_crawls[url] = {
        "url": url, "instance": instance,
        "pages_done": 0, "chunks": 0, "errors": 0, "blocked": 0, "skipped": 0,
        "start_ts": datetime.now(timezone.utc).isoformat(), "done": False,
        # Store params so the UI can reconnect with the exact same settings
        "params": {
            "depth": depth, "respect_robots": respect_robots,
            "local_only": local_only, "path_prefix_only": path_prefix_only,
            "force_reindex": force_reindex,
            "max_pages": max_pages, "concurrency": concurrency,
            "js_render": js_render, "js_concurrency": js_concurrency,
            "smart_mode": smart_mode, "min_words": min_words,
        },
    }

    async def cb(u, status, n=0, err="", count=0):
        state = _active_crawls.get(url)
        if state:
            state["pages_done"] = count
            if status == "indexed":   state["chunks"]  += n
            elif status == "error":   state["errors"]  += 1
            elif status == "blocked": state["blocked"] += 1
            elif status == "skipped": state["skipped"] += 1
        await queue.put({"url": u, "status": status, "chunks": n, "error": err, "pages_done": count})

    async def run():
        try:
            await crawl_url(
                instance, url, depth, progress_cb=cb,
                respect_robots=respect_robots, local_only=local_only,
                path_prefix_only=path_prefix_only,
                force_reindex=force_reindex,
                max_pages=max_pages, concurrency=concurrency, js_render=js_render,
                js_concurrency=js_concurrency, smart_mode=smart_mode,
                min_words=min_words, rc=rc,
            )
        finally:
            if url in _active_crawls:
                _active_crawls[url]["done"] = True
            await queue.put(None)   # sentinel: signals the generator to stop

    task = asyncio.create_task(run())
    _crawl_tasks[url] = task

    async def generate():
        try:
            while True:
                item = await queue.get()
                if item is None:
                    yield f"data: {json.dumps({'done': True})}\n\n"
                    break
                yield f"data: {json.dumps(item)}\n\n"
        except GeneratorExit:
            pass
        finally:
            # Client disconnected — cancel the crawl task so it doesn't keep
            # running and consuming resources with no one listening.
            if not task.done():
                task.cancel()
                try:
                    await task
                except (asyncio.CancelledError, Exception):
                    pass
            if url in _active_crawls:
                _active_crawls[url]["done"] = True

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/rag/logs")
def api_ingest_logs():
    return _ingestion_logs[-200:]

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — RAG EXPORT / IMPORT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_EXPORT_BATCH  = 200          # hgetall commands per pipeline call
_EXPORT_BUFSIZE = 256 * 1024  # bytes to accumulate before yielding an NDJSON chunk


def _iter_chunks_pipelined(rc: redis.Redis, prefix: str):
    """
    Yield chunk dicts from Redis efficiently:
    - scan_iter instead of KEYS  → non-blocking cursor scan, safe on large DBs
    - pipelined HGETALL in batches of _EXPORT_BATCH → N/200 round-trips instead of N
    - base64-encode embedding inline (unavoidable, but done once per chunk)
    """
    batch: list = []
    for key in rc.scan_iter(f"{prefix}:chunk:*", count=500):
        batch.append(key)
        if len(batch) >= _EXPORT_BATCH:
            pipe = rc.pipeline(transaction=False)
            for k in batch:
                pipe.hgetall(k)
            for d in pipe.execute():
                if d:
                    emb_raw = d.get(b"embedding", b"")
                    yield {
                        "id":            d.get(b"chunk_id", b"").decode(),
                        "text":          d.get(b"text",     b"").decode(),
                        "source":        d.get(b"source",   b"").decode(),
                        "embedding_b64": base64.b64encode(emb_raw).decode() if emb_raw else "",
                    }
            batch = []
    # flush remainder
    if batch:
        pipe = rc.pipeline(transaction=False)
        for k in batch:
            pipe.hgetall(k)
        for d in pipe.execute():
            if d:
                emb_raw = d.get(b"embedding", b"")
                yield {
                    "id":            d.get(b"chunk_id", b"").decode(),
                    "text":          d.get(b"text",     b"").decode(),
                    "source":        d.get(b"source",   b"").decode(),
                    "embedding_b64": base64.b64encode(emb_raw).decode() if emb_raw else "",
                }


@app.get("/api/rag/{instance}/export")
def api_export_rag(instance: str, endpoint: str | None = None):
    """
    Export a RAG instance as a ZIP file (ZIP_STORED — no compression).

    Embeddings are float32 random bytes that compress < 1%; using DEFLATE
    wastes CPU for almost no size gain.  ZIP_STORED skips compression and
    lets the data flow straight to the client, cutting export time by 50-80%.
    """
    rc     = _rc_for(instance, endpoint)
    prefix = rag_prefix(instance)
    chunks = list(_iter_chunks_pipelined(rc, prefix))
    meta_raw = rc.get(f"rag_meta:{instance}")
    meta     = json.loads(meta_raw) if meta_raw else {}
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as zf:
        zf.writestr("meta.json",   json.dumps(meta))
        zf.writestr("chunks.json", json.dumps(chunks))
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={instance}_rag.zip"},
    )


@app.get("/api/rag/{instance}/export/stream")
def api_export_rag_stream(instance: str, endpoint: str | None = None):
    """
    Stream-export a RAG instance as NDJSON (one JSON object per line).
    Each line has a ``_t`` discriminator: "meta" | "chunk" | "done".

    Optimisations vs the naïve approach:
    - scan_iter + pipelined HGETALL  → far fewer Redis round-trips
    - output buffered to _EXPORT_BUFSIZE before yielding → fewer HTTP frames
    """
    rc     = _rc_for(instance, endpoint)
    prefix = rag_prefix(instance)

    def generate():
        meta_raw = rc.get(f"rag_meta:{instance}")
        meta = json.loads(meta_raw) if meta_raw else {}
        yield json.dumps({"_t": "meta", **meta}) + "\n"

        count    = 0
        out_buf: list[str] = []
        out_size = 0

        for chunk in _iter_chunks_pipelined(rc, prefix):
            line = json.dumps({"_t": "chunk", **chunk}) + "\n"
            out_buf.append(line)
            out_size += len(line)
            count += 1
            if out_size >= _EXPORT_BUFSIZE:
                yield "".join(out_buf)
                out_buf  = []
                out_size = 0

        if out_buf:
            yield "".join(out_buf)

        yield json.dumps({"_t": "done", "total": count}) + "\n"

    fname = f"{instance}_rag.jsonl"
    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={
            "Content-Disposition": f"attachment; filename={fname}",
            "Cache-Control": "no-cache",
        },
    )


@app.post("/api/rag/{instance}/import")
async def api_import_rag(instance: str, file: UploadFile = File(...), endpoint: str | None = None):
    """Import a RAG ZIP or NDJSON (jsonl) export. Re-uses stored embeddings when present."""
    content = await file.read()
    fname   = (file.filename or "").lower()

    # ── NDJSON / jsonl format ──────────────────────────────────────────────────
    if fname.endswith(".jsonl") or fname.endswith(".ndjson"):
        meta: dict      = {}
        chunks_raw: list[dict] = []
        for raw_line in content.decode().splitlines():
            raw_line = raw_line.strip()
            if not raw_line:
                continue
            try:
                obj = json.loads(raw_line)
            except json.JSONDecodeError:
                continue
            t = obj.pop("_t", None)
            if t == "meta":
                meta = obj
            elif t == "chunk":
                chunks_raw.append(obj)
            # "done" line is informational only
    # ── ZIP format (legacy) ───────────────────────────────────────────────────
    else:
        buf = io.BytesIO(content)
        with zipfile.ZipFile(buf) as zf:
            meta       = json.loads(zf.read("meta.json"))
            chunks_raw = json.loads(zf.read("chunks.json"))

    # Always write to the instance's CURRENT assigned endpoint, not the one
    # recorded in the ZIP (which may be from a different server entirely).
    rc = _rc_for(instance, endpoint)

    prefix = rag_prefix(instance)

    def _do_import():
        # The whole import — read existing meta, write meta, build the index, embed and
        # pipeline the chunks, sync the counter — is synchronous Redis/CPU work. Run it in
        # one thread so a large import never blocks the event loop (nor other sessions).
        existing_meta_raw = rc.get(f"rag_meta:{instance}")
        if existing_meta_raw:
            existing_meta = json.loads(existing_meta_raw)
            meta["redis_endpoint"] = existing_meta.get("redis_endpoint", "default")
            meta.setdefault("color",   existing_meta.get("color",   "#6366f1"))
            meta.setdefault("enabled", existing_meta.get("enabled", True))
        rc.set(f"rag_meta:{instance}", json.dumps(meta))

        ensure_rag_index(instance, rc)
        pipe = rc.pipeline(transaction=False)
        for ch in chunks_raw:
            emb_b64   = ch.get("embedding_b64", "")
            emb_bytes = (
                base64.b64decode(emb_b64)
                if emb_b64
                else embed(ch["text"]).astype(np.float32).tobytes()
            )
            pipe.hset(f"{prefix}:chunk:{ch['id']}", mapping={
                "text":      ch["text"].encode(),
                "source":    ch.get("source", "").encode(),
                "chunk_id":  str(ch["id"]),
                "embedding": emb_bytes,
            })
        pipe.execute()

        # Sync the chunk counter so future additions don't collide with imported IDs.
        max_id = -1
        for ch in chunks_raw:
            try:
                max_id = max(max_id, int(ch["id"]))
            except (ValueError, TypeError):
                pass
        if max_id >= 0:
            rc.set(f"rag:{instance}:chunk_counter", str(max_id + 1))

    await asyncio.to_thread(_do_import)
    invalidate_rag_meta(instance)   # in-process; the write above has completed

    return {"ok": True, "chunks": len(chunks_raw)}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — REDIS MEMORY MONITOR
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/redis/memory")
def api_redis_memory():
    try:
        info = r().info("memory")
        return {
            "used":       info.get("used_memory", 0),
            "used_human": info.get("used_memory_human", "?"),
            "peak":       info.get("used_memory_peak", 0),
            "peak_human": info.get("used_memory_peak_human", "?"),
            "max":        info.get("maxmemory", 0),
            "max_human":  info.get("maxmemory_human", "nomax"),
        }
    except Exception as e:
        return {"error": str(e)}


def _collect_redis_server_stats(ep_name: str, rc: redis.Redis, ep_cfg: dict) -> dict:
    """
    Collect comprehensive stats for a single Redis endpoint.
    Returns a dict suitable for the analytics UI.
    """
    base: dict = {
        "name":    ep_name,
        "host":    ep_cfg.get("host", "localhost"),
        "port":    int(ep_cfg.get("port", 6379)),
        "db":      int(ep_cfg.get("db", 0)),
        "ok":      False,
    }
    try:
        info = rc.info("all")
    except Exception as e:
        base["error"] = str(e)
        return base

    # ── basic ──────────────────────────────────────────────────────────────────
    base["ok"]      = True
    base["version"] = info.get("redis_version", "?")
    base["mode"]    = info.get("redis_mode", "standalone")
    base["role"]    = info.get("role", "master")
    base["os"]      = info.get("os", "")

    # ── memory ─────────────────────────────────────────────────────────────────
    base["mem_used"]        = info.get("used_memory_human", "?")
    base["mem_used_bytes"]  = info.get("used_memory", 0)
    base["mem_peak"]        = info.get("used_memory_peak_human", "?")
    base["mem_rss"]         = info.get("used_memory_rss_human", "?")
    base["mem_max_bytes"]   = info.get("maxmemory", 0)
    base["mem_max"]         = info.get("maxmemory_human", "0") if info.get("maxmemory") else "unlimited"
    base["mem_pct"]         = (
        round(info.get("used_memory", 0) / info.get("maxmemory") * 100)
        if info.get("maxmemory") else 0
    )
    base["mem_fragmentation_ratio"] = info.get("mem_fragmentation_ratio", 1.0)

    # ── clients & throughput ───────────────────────────────────────────────────
    base["connected_clients"]        = info.get("connected_clients", 0)
    base["blocked_clients"]          = info.get("blocked_clients", 0)
    base["uptime_days"]              = info.get("uptime_in_days", 0)
    base["uptime_seconds"]           = info.get("uptime_in_seconds", 0)
    base["ops_per_sec"]              = info.get("instantaneous_ops_per_sec", 0)
    base["total_commands_processed"] = info.get("total_commands_processed", 0)
    base["total_connections"]        = info.get("total_connections_received", 0)
    base["net_input_bytes"]          = info.get("total_net_input_bytes", 0)
    base["net_output_bytes"]         = info.get("total_net_output_bytes", 0)

    # ── keyspace hit rate ──────────────────────────────────────────────────────
    hits   = info.get("keyspace_hits",   0)
    misses = info.get("keyspace_misses", 0)
    base["keyspace_hits"]   = hits
    base["keyspace_misses"] = misses
    base["keyspace_hit_rate"] = round(hits / (hits + misses) * 100, 1) if (hits + misses) else None

    # ── evictions / expirations ────────────────────────────────────────────────
    base["evicted_keys"] = info.get("evicted_keys", 0)
    base["expired_keys"] = info.get("expired_keys", 0)

    # ── persistence ────────────────────────────────────────────────────────────
    base["rdb_enabled"]            = info.get("rdb_last_bgsave_status") is not None
    base["rdb_last_save_status"]   = info.get("rdb_last_bgsave_status", "?")
    base["rdb_changes_since_save"] = info.get("rdb_changes_since_last_save", 0)
    base["aof_enabled"]            = bool(info.get("aof_enabled", 0))
    base["aof_rewrite_running"]    = bool(info.get("aof_rewrite_in_progress", 0))

    # ── keyspace (databases) ───────────────────────────────────────────────────
    keyspace = []
    total_keys = 0
    for k, v in info.items():
        if k.startswith("db") and isinstance(v, dict):
            keys = v.get("keys", 0)
            total_keys += keys
            keyspace.append({
                "db":      int(k[2:]),
                "keys":    keys,
                "expires": v.get("expires", 0),
                "avg_ttl": v.get("avg_ttl", 0),
            })
    base["keyspace"]   = sorted(keyspace, key=lambda x: x["db"])
    base["total_keys"] = total_keys

    # ── replication ────────────────────────────────────────────────────────────
    base["repl_slaves"]     = info.get("connected_slaves", 0)
    base["repl_master"]     = None
    if info.get("role") == "slave":
        base["repl_master"] = f"{info.get('master_host','?')}:{info.get('master_port','?')}"
    base["repl_backlog"]    = info.get("repl_backlog_active", 0)

    # ── cluster ────────────────────────────────────────────────────────────────
    cluster_enabled = bool(info.get("cluster_enabled", 0))
    base["cluster_enabled"] = cluster_enabled
    if cluster_enabled:
        try:
            ci = rc.execute_command("CLUSTER INFO")
            # CLUSTER INFO returns a bulk string of "key:value\r\n" lines
            if isinstance(ci, (bytes, str)):
                ci_str = ci.decode() if isinstance(ci, bytes) else ci
                ci_map = dict(
                    line.split(":", 1)
                    for line in ci_str.strip().splitlines()
                    if ":" in line
                )
            else:
                ci_map = ci  # some clients parse it to dict already
            base["cluster_state"]       = ci_map.get("cluster_state", "?")
            base["cluster_slots_ok"]    = int(ci_map.get("cluster_slots_ok", 0))
            base["cluster_slots_fail"]  = int(ci_map.get("cluster_slots_fail", 0))
            base["cluster_known_nodes"] = int(ci_map.get("cluster_known_nodes", 0))
            base["cluster_size"]        = int(ci_map.get("cluster_size", 0))
        except Exception:
            base["cluster_state"] = "unknown"
    else:
        base["cluster_state"] = None

    # ── search module ──────────────────────────────────────────────────────────
    base["search_available"] = probe_search(rc, ep_name)

    # ── RAG instances on this endpoint ─────────────────────────────────────────
    all_insts = list_rag_instances()
    base["rag_instances"] = [
        {"name": i["name"], "chunks": i["chunks"], "enabled": i.get("enabled", True)}
        for i in all_insts
        if (i.get("redis_endpoint") or "default") == ep_name
    ]

    return base


@app.get("/api/redis/all-stats")
def api_redis_all_stats():
    """
    Return comprehensive stats for all configured Redis endpoints
    (default + any extras).  Used by the analytics panel.
    """
    primary_cfg = {**_config.get("redis", {}), "name": "default"}
    endpoints_to_check = [("default", r(), primary_cfg)]
    for ep in _config.get("redis_endpoints", []):
        ep_name = ep.get("name", "")
        if ep_name and ep_name != "default":
            try:
                endpoints_to_check.append((ep_name, r_for(ep_name), ep))
            except Exception:
                endpoints_to_check.append((ep_name, None, ep))

    results = []
    for ep_name, rc, cfg in endpoints_to_check:
        if rc is None:
            results.append({"name": ep_name, "ok": False, "error": "Could not connect"})
        else:
            results.append(_collect_redis_server_stats(ep_name, rc, cfg))
    return results

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — FEEDBACK
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/feedback")
async def api_feedback(payload: dict):
    entry = {**payload, "ts": datetime.now(timezone.utc).isoformat()}
    _feedback.append(entry)
    with open(FEEDBACK_PATH, "w") as f:
        json.dump(_feedback, f)
    return {"ok": True}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — SESSIONS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/sessions")
def api_sessions():
    # In-memory sessions (active connections)
    result = []
    for sid, msgs in _sessions.items():
        preview = ""
        for m in reversed(msgs):
            if m.get("role") == "assistant":
                preview = m.get("content", "")[:60]
                break
        result.append({"id": sid, "messages": len(msgs), "preview": preview})
    # Merge in Redis-persisted sessions not currently loaded in memory
    result.extend(list_sessions_from_redis())
    return result

@app.get("/api/sessions/{sid}")
def api_session(sid: str):
    if sid in _sessions:
        return _sessions[sid]
    # Try Redis
    msgs = load_session(sid)
    if msgs:
        _sessions[sid] = msgs   # cache in memory
    return msgs

@app.delete("/api/sessions/{sid}")
def api_delete_session(sid: str):
    _sessions.pop(sid, None)
    delete_session_from_redis(sid)
    return {"ok": True}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — PROMPT TEMPLATES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/templates")
def api_templates():
    return _config.get("prompt_templates", [])

@app.post("/api/templates")
async def api_save_templates(payload: list):
    _config["prompt_templates"] = payload
    save_config(_config)
    return {"ok": True}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FILE UPLOAD — CHAT CONTEXT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/chat/upload-file")
async def api_chat_upload_file(file: UploadFile = File(...)):
    """Extract plain text from an uploaded document (TXT, MD, CSV, PDF, DOCX, XLSX).

    The extracted text is returned to the browser and included in the next chat
    message as inline context — similar to pasting the document into the prompt,
    but with server-side format conversion.  Files are NOT stored on disk or
    indexed into Redis; they exist only for the duration of the browser session.
    """
    data = await file.read()
    if len(data) > _CHAT_FILE_MAX_BYTES:
        raise HTTPException(413, f"File too large (max {_CHAT_FILE_MAX_BYTES // (1024*1024)} MB)")
    try:
        text = await asyncio.to_thread(extract_file_text, file.filename or "", data)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(422, f"Could not parse file: {e}")
    truncated = len(text) > _CHAT_FILE_MAX_CHARS
    if truncated:
        text = text[:_CHAT_FILE_MAX_CHARS]
    return {
        "filename": file.filename,
        "chars":    len(text),
        "truncated": truncated,
        "text":     text,
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# WEBSOCKET — CHAT HANDLER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.websocket("/ws/chat/{sid}")
async def ws_chat(ws: WebSocket, sid: str):
    """One WebSocket per session; each message is a full chat request.

    The streaming task runs as an asyncio background task so the receive loop
    can concurrently process an {"type":"abort"} message and cancel it.
    """
    await mgr.connect(ws, sid)
    if sid not in _sessions:
        _sessions[sid] = await asyncio.to_thread(load_session, sid)
    try:
        while True:
            raw = await ws.receive_text()
            msg = json.loads(raw)

            # ── Abort: cancel the active streaming task if any ──────────────
            if msg.get("type") == "abort":
                task = _chat_tasks.pop(sid, None)
                if task and not task.done():
                    task.cancel()
                    try:
                        await task
                    except (asyncio.CancelledError, Exception):
                        pass
                # Notify client that the stream is done (aborted)
                try:
                    await ws.send_json({"type": "stream_end", "aborted": True, "latency": {}})
                except Exception:
                    pass
                continue

            # ── New chat turn: start as a background task ────────────────────
            # If a task is somehow still running, let it finish (shouldn't happen
            # in normal usage since send-btn is disabled during streaming).
            task = _chat_tasks.get(sid)
            if task and not task.done():
                continue

            t = asyncio.create_task(handle_chat(ws, sid, msg))
            _chat_tasks[sid] = t

    except WebSocketDisconnect:
        # Cancel any in-flight task when the client disconnects
        task = _chat_tasks.pop(sid, None)
        if task and not task.done():
            task.cancel()
        mgr.disconnect(sid)


async def handle_chat(ws: WebSocket, sid: str, msg: dict):
    """
    Process one chat turn:
    1. Semantic cache lookup  — return immediately on hit
    2. RAG retrieval          — find relevant chunks from the knowledge base
    3. Build message history  — last 10 turns + system prompt with context
    4. Stream LLM response    — Ollama / Claude / OpenAI token-by-token
    5. Auto-title generation  — on first turn, generate a short session title
    6. Cache store            — save the response for future similar queries
    """
    query         = msg.get("content", "")
    provider      = msg.get("provider", _config.get("provider", "ollama"))
    model         = msg.get("model", "")
    images        = msg.get("images", [])       # list of base64 data URIs
    # base_instruction (always) + the selected template's system prompt (additive)
    system_prompt = compose_system_prompt(msg.get("system_prompt"))
    source_filter = msg.get("source_filter", "")   # optional substring filter on chunk source
    bypass_cache  = msg.get("bypass_cache", False)  # skip cache lookup (re-run fresh)
    file_context  = msg.get("file_context", [])    # list of {name, text} uploaded documents

    # ── RAG instance selection ───────────────────────────────────────────────
    # Accept either:
    #   rag_instances: ["inst1", "inst2"]  → parallel multi-instance query
    #   rag_instance:  "inst1"             → single-instance query (legacy/default)
    rag_instances_raw = msg.get("rag_instances")   # multi-instance list
    rag_inst          = msg.get("rag_instance", _config.get("active_rag", "default"))
    # Normalise to a list for uniform handling below
    if rag_instances_raw and isinstance(rag_instances_raw, list):
        rag_instances = [i for i in rag_instances_raw if i]
    else:
        rag_instances = [rag_inst] if rag_inst else []
    parallel_mode = len(rag_instances) > 1
    t0            = time.time()

    # Fall back to default model for the active provider if none specified
    if not model:
        if provider == "claude":
            model = _config.get("claude", {}).get("model", "claude-sonnet-4-6")
        elif provider == "openai":
            model = _config.get("openai", {}).get("model", "gpt-4o")
        elif provider == "qwen":
            model = _config.get("qwen", {}).get("model", "qwen-plus")
        elif provider == "mistral":
            model = _config.get("mistral", {}).get("model", "mistral-small-latest")
        elif provider == "groq":
            model = _config.get("groq", {}).get("model", "llama-3.3-70b-versatile")
        elif provider == "gemini":
            model = _config.get("gemini", {}).get("model", "gemini-3-flash-preview")
        else:
            model = _config.get("ollama", {}).get("model", "")

    if not model:
        await ws.send_json({"type": "error", "content": "No model selected."})
        return

    # ── 1. Semantic cache check ─────────────────────────────────────────────
    # Skip cache entirely for vision requests: same text + different image ≠ same answer.
    await ws.send_json({"type": "status", "phase": "cache"})
    cache_threshold = _config.get("cache", {}).get("similarity_threshold", 0.92)
    t_cache_start   = time.time()
    # Off the event loop: while these run, no other session's tokens can flush and
    # the WS receive loop cannot read an {'type':'abort'} frame (Stop goes dead).
    hit             = (await asyncio.to_thread(cache_lookup, query, cache_threshold)
                       if not images and not bypass_cache and not wants_visual(query) else None)
    t_cache         = round(time.time() - t_cache_start, 3)

    if hit:
        await ws.send_json({
            "type":     "cache_hit",
            "content":  hit["response"],
            "score":    hit["score"],
            "entry_id": hit.get("entry_id", ""),
            "latency":  {"cache": t_cache, "total": round(time.time() - t0, 3)},
        })
        if hit.get("chunks"):
            await ws.send_json({"type": "rag_context", "chunks": hit["chunks"], "latency": {"cache": t_cache, "rag": 0}})
        return

    # ── 2. RAG retrieval ────────────────────────────────────────────────────
    await ws.send_json({"type": "status", "phase": "rag"})
    rag_cfg       = _config.get("rag", {})
    rag_threshold = rag_cfg.get("similarity_threshold", 0.75)
    top_k         = rag_cfg.get("top_k", 5)
    hybrid_search = rag_cfg.get("hybrid_search", True)
    # ── HyDE: generate a hypothetical answer and use its embedding for search.
    # Timed on its own — it is an LLM generation, not a Redis/vector cost — so the
    # latency badge attributes it to HyDE instead of silently inflating "rag".
    t_hyde = 0.0
    search_vec: "np.ndarray | None" = None
    if rag_instances and _config.get("hyde", {}).get("enabled", False):
        await ws.send_json({"type": "status", "phase": "hyde"})
        t_hyde_start = time.time()
        hypothesis = await hyde_generate(query, provider, model)
        if hypothesis:
            search_vec = (await asyncio.to_thread(embed, hypothesis)).astype(np.float32)
            log.info(f"HyDE hypothesis ({len(hypothesis)} chars) embedded for RAG search")
        t_hyde = round(time.time() - t_hyde_start, 3)

    t_rag_start   = time.time()   # retrieval + rerank only; HyDE is timed above

    if parallel_mode:
        # Multi-instance parallel query — search all requested instances simultaneously.
        # search_rag_parallel filters disabled instances internally.
        chunks = await search_rag_parallel(rag_instances, query, top_k, rag_threshold, hybrid_search, search_vec)
    elif rag_instances:
        # Single-instance query (normal mode)
        rag_inst = rag_instances[0]
        meta, _ep = await _rag_meta_cached_async(rag_inst)   # primes the cache off-loop
        rag_enabled = (meta or {}).get("enabled", True)
        chunks = (
            await asyncio.to_thread(search_rag, rag_inst, query, top_k, rag_threshold,
                                    rc_for_instance(rag_inst), hybrid_search, search_vec)
            if rag_enabled else []
        )
    else:
        chunks = []

    # ── Source filter: restrict chunks to a specific source prefix/substring
    if source_filter and chunks:
        chunks = [c for c in chunks if source_filter.lower() in c.get("source", "").lower()]

    # ── Cross-encoder reranking (runs after fast retrieval, before LLM)
    if chunks:
        top_n = _config.get("reranker", {}).get("top_n", top_k)
        chunks = await asyncio.to_thread(rerank_chunks, query, chunks, top_n)

    t_rag = round(time.time() - t_rag_start, 3)

    # Inject uploaded file text into the system prompt (before RAG context)
    if file_context:
        file_parts = [f"[File: {f['name']}]\n{f['text']}" for f in file_context if f.get("text")]
        if file_parts:
            system_prompt += "\n\nThe user has attached the following document(s) — use them to answer:\n\n" \
                             + "\n\n---\n\n".join(file_parts)

    # Inject retrieved RAG context into the system prompt
    if chunks:
        context = "\n\n".join(f"[Source: {c['source']}]\n{c['text']}" for c in chunks)
        system_prompt += f"\n\nUse the following context to answer:\n{context}"

    # ── 3. Build message list ───────────────────────────────────────────────
    history  = _sessions[sid][-10:]   # last 10 turns for context window management
    messages = [{"role": "system", "content": system_prompt}]
    for h in history:
        messages.append({"role": h["role"], "content": h["content"]})

    # Vision: wrap user content as a multi-modal list when images are attached
    user_content: Any = query
    if images:
        user_content = (
            [{"type": "text", "text": query}]
            + [{"type": "image_url", "image_url": {"url": img}} for img in images]
        )
    messages.append({"role": "user", "content": user_content})

    # ── 4. Stream LLM response ──────────────────────────────────────────────
    # stream_start MUST fire first so the frontend creates the message element
    # and sets currentAiMsgId before rag_context arrives.
    await ws.send_json({"type": "stream_start"})
    await ws.send_json({"type": "rag_context", "chunks": chunks, "latency": {"cache": t_cache, "hyde": t_hyde, "rag": t_rag}})

    full_response = ""
    stream_error  = False
    t_llm_start   = time.time()

    try:
        # Route to the correct provider
        if provider == "claude":
            stream_gen = claude_stream(messages, model)
        elif provider == "openai":
            stream_gen = openai_stream(messages, model)
        elif provider == "qwen":
            stream_gen = qwen_stream(messages, model)
        elif provider == "mistral":
            stream_gen = mistral_stream(messages, model)
        elif provider == "groq":
            stream_gen = groq_stream(messages, model)
        elif provider == "gemini":
            stream_gen = gemini_stream(messages, model)
        else:
            stream_gen = ollama_stream(messages, model, images or None)

        async for token, done in stream_gen:
            full_response += token
            await ws.send_json({"type": "token", "content": token, "done": done})
            if done:
                # Error tokens must not be cached
                if token and (token.startswith("Error:") or "error:" in token.lower()[:20]):
                    stream_error = True
                break

    except asyncio.CancelledError:
        # Client sent abort — clean exit, do not send stream_end (ws_chat handles it)
        raise
    except Exception as e:
        await ws.send_json({"type": "error", "content": str(e)})
        # Report the real phase timings we already have (cache/hyde/rag ran before the
        # LLM error); llm is however far the failed stream got.
        await ws.send_json({"type": "stream_end", "latency": {
            "cache": t_cache, "hyde": t_hyde, "rag": t_rag,
            "llm": round(time.time() - t_llm_start, 3),
            "total": round(time.time() - t0, 3)}, "title": None})
        return

    t_llm  = round(time.time() - t_llm_start, 3)
    total  = round(time.time() - t0, 3)

    # ── Store turn in session (memory + Redis) ──────────────────────────────
    _sessions[sid].append({"role": "user",      "content": query})
    _sessions[sid].append({"role": "assistant", "content": full_response})
    await asyncio.to_thread(save_session, sid, _sessions[sid])

    # ── 5. Release the client FIRST ─────────────────────────────────────────
    # stream_end unlocks the composer. Auto-titling is a second, full LLM call
    # (measured at ~423 ms on local gemma4, seconds on a large model) and the cache
    # store touches Redis — neither is something the user should wait behind while
    # staring at a finished answer. Send stream_end now; the title arrives later as
    # its own event.
    await ws.send_json({
        "type":    "stream_end",
        "latency": {"cache": t_cache, "hyde": t_hyde, "rag": t_rag, "llm": t_llm, "total": total},
        "title":   None,
    })
    _chat_tasks.pop(sid, None)

    # ── 6. Cache store — off the critical path ──────────────────────────────
    if not stream_error and not images and not wants_visual(query):
        await asyncio.to_thread(cache_store, query, full_response, chunks)

    # ── 7. Auto-title (first turn only), delivered as a follow-up event ─────
    title_msg = None
    if len(_sessions[sid]) == 2:
        try:
            t_payload = [{"role": "user", "content": (
                "Reply with ONLY a short title of 2-5 words for this query. "
                "No punctuation, no explanation, no numbering, no quotes. "
                f"Just the title words. Query: {query}"
            )}]
            title_chunks = ""

            # Use the same provider that answered the question
            if provider == "claude":
                async for tok, done in claude_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            elif provider == "openai":
                async for tok, done in openai_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            elif provider == "qwen":
                async for tok, done in qwen_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            elif provider == "mistral":
                async for tok, done in mistral_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            elif provider == "groq":
                async for tok, done in groq_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            elif provider == "gemini":
                async for tok, done in gemini_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break
            else:
                async for tok, done in ollama_stream(t_payload, model):
                    title_chunks += tok
                    if done:
                        break

            # Clean the response: take first non-empty line, strip list markers
            raw_title  = title_chunks.strip()
            first_line = next((l.strip() for l in raw_title.splitlines() if l.strip()), raw_title)
            first_line = re.sub(r'^[\d]+[.)]\s*|^[-*•]\s*', '', first_line).strip()
            first_line = first_line.strip('"\'').rstrip('.:,;')
            title_msg  = first_line[:60] if first_line else None

        except Exception:
            pass   # title generation is non-critical

    if title_msg:
        try:
            await ws.send_json({"type": "session_title", "title": title_msg})
        except Exception:
            pass   # client may have navigated away; the title is non-critical

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — BATCH CHAT (non-streaming REST)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/chat")
async def api_chat(payload: dict):
    """
    Non-streaming chat endpoint — returns the full LLM response as JSON.

    Same semantics as the WebSocket handler but suitable for scripting,
    batch pipelines, or any client that does not support WebSockets.

    Request body fields (all optional except ``content``):
      content        — the user message
      session_id     — reuse an existing session (creates new if omitted)
      provider       — "ollama" | "claude" | "openai" | "qwen"
      model          — model name (uses config default if omitted)
      system_prompt  — overrides the default system prompt
      rag_instance   — single RAG instance to query
      rag_instances  — list of RAG instances for parallel multi-instance query
      source_filter  — substring filter applied to chunk sources
      use_cache      — bool (default true); skip semantic cache when false

    Response:
      session_id, response, chunks (list of RAG chunks used)
    """
    query         = payload.get("content", "")
    provider      = payload.get("provider", _config.get("provider", "ollama"))
    model         = payload.get("model", "")
    # base_instruction (always) + the selected template's system prompt (additive)
    system_prompt = compose_system_prompt(payload.get("system_prompt"))
    source_filter = payload.get("source_filter", "")
    images        = payload.get("images", [])       # list of base64 data URIs
    file_context  = payload.get("file_context", []) # list of {name, text} uploaded documents
    use_cache     = bool(payload.get("use_cache", True))
    sid = payload.get("session_id") or f"rest_{uuid.uuid4().hex[:8]}"

    # Ensure session exists in memory
    if sid not in _sessions:
        _sessions[sid] = await asyncio.to_thread(load_session, sid)

    # Model fallback
    if not model:
        if provider == "claude":
            model = _config.get("claude", {}).get("model", "claude-sonnet-4-6")
        elif provider == "openai":
            model = _config.get("openai", {}).get("model", "gpt-4o")
        elif provider == "qwen":
            model = _config.get("qwen", {}).get("model", "qwen-plus")
        elif provider == "mistral":
            model = _config.get("mistral", {}).get("model", "mistral-small-latest")
        elif provider == "groq":
            model = _config.get("groq", {}).get("model", "llama-3.3-70b-versatile")
        elif provider == "gemini":
            model = _config.get("gemini", {}).get("model", "gemini-3-flash-preview")
        else:
            model = _config.get("ollama", {}).get("model", "")
    if not model:
        raise HTTPException(400, "No model selected — configure a model in settings")

    # Semantic cache — skip for vision requests and chart requests (see wants_visual)
    if use_cache and not images and not wants_visual(query):
        cache_threshold = _config.get("cache", {}).get("similarity_threshold", 0.92)
        hit = await asyncio.to_thread(cache_lookup, query, cache_threshold)
        if hit:
            return {"session_id": sid, "response": hit["response"], "chunks": hit.get("chunks", []), "cache_hit": True, "cache_score": hit["score"]}

    # RAG instance resolution
    rag_instances_raw = payload.get("rag_instances")
    rag_inst          = payload.get("rag_instance", _config.get("active_rag", "default"))
    if rag_instances_raw and isinstance(rag_instances_raw, list):
        rag_instances = [i for i in rag_instances_raw if i]
    else:
        rag_instances = [rag_inst] if rag_inst else []
    parallel_mode = len(rag_instances) > 1

    rag_cfg       = _config.get("rag", {})
    rag_threshold = rag_cfg.get("similarity_threshold", 0.75)
    top_k         = rag_cfg.get("top_k", 5)
    hybrid_search = rag_cfg.get("hybrid_search", True)

    # HyDE
    search_vec: "np.ndarray | None" = None
    if rag_instances and _config.get("hyde", {}).get("enabled", False):
        hypothesis = await hyde_generate(query, provider, model)
        if hypothesis:
            search_vec = (await asyncio.to_thread(embed, hypothesis)).astype(np.float32)

    # RAG retrieval
    if parallel_mode:
        chunks = await search_rag_parallel(rag_instances, query, top_k, rag_threshold, hybrid_search, search_vec)
    elif rag_instances:
        inst = rag_instances[0]
        meta, _ep = await _rag_meta_cached_async(inst)       # primes the cache off-loop
        rag_enabled = (meta or {}).get("enabled", True)
        chunks = (
            await asyncio.to_thread(search_rag, inst, query, top_k, rag_threshold,
                                    rc_for_instance(inst), hybrid_search, search_vec)
            if rag_enabled else []
        )
    else:
        chunks = []

    if source_filter and chunks:
        chunks = [c for c in chunks if source_filter.lower() in c.get("source", "").lower()]

    top_n = _config.get("reranker", {}).get("top_n", top_k)
    chunks = await asyncio.to_thread(rerank_chunks, query, chunks, top_n)

    if file_context:
        file_parts = [f"[File: {f['name']}]\n{f['text']}" for f in file_context if f.get("text")]
        if file_parts:
            system_prompt += "\n\nThe user has attached the following document(s) — use them to answer:\n\n" \
                             + "\n\n---\n\n".join(file_parts)

    if chunks:
        context = "\n\n".join(f"[Source: {c['source']}]\n{c['text']}" for c in chunks)
        system_prompt += f"\n\nUse the following context to answer:\n{context}"

    history  = _sessions[sid][-10:]
    messages = [{"role": "system", "content": system_prompt}]
    for h in history:
        messages.append({"role": h["role"], "content": h["content"]})
    messages.append({"role": "user", "content": query})

    full_response = ""
    stream_error  = False
    try:
        if provider == "claude":
            async for tok, done in claude_stream(messages, model):
                full_response += tok
                if done: break
        elif provider == "openai":
            async for tok, done in openai_stream(messages, model):
                full_response += tok
                if done: break
        elif provider == "qwen":
            async for tok, done in qwen_stream(messages, model):
                full_response += tok
                if done: break
        elif provider == "mistral":
            async for tok, done in mistral_stream(messages, model):
                full_response += tok
                if done: break
        elif provider == "groq":
            async for tok, done in groq_stream(messages, model):
                full_response += tok
                if done: break
        elif provider == "gemini":
            async for tok, done in gemini_stream(messages, model):
                full_response += tok
                if done: break
        else:
            async for tok, done in ollama_stream(messages, model):
                full_response += tok
                if done: break
        if full_response.startswith("Error:") or "error:" in full_response.lower()[:20]:
            stream_error = True
    except Exception as e:
        raise HTTPException(500, str(e))

    _sessions[sid].append({"role": "user",      "content": query})
    _sessions[sid].append({"role": "assistant", "content": full_response})
    await asyncio.to_thread(save_session, sid, _sessions[sid])

    if not stream_error and not images and not wants_visual(query):
        await asyncio.to_thread(cache_store, query, full_response, chunks)

    return {"session_id": sid, "response": full_response, "chunks": chunks, "cache_hit": False}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — RAG SOURCE LISTING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/rag/{instance}/sources")
def api_rag_sources(instance: str, endpoint: str | None = None):
    """Return all unique source identifiers stored in a RAG instance."""
    rc     = _rc_for(instance, endpoint)
    prefix = rag_prefix(instance)
    sources: set[str] = set()
    batch: list = []
    for key in rc.scan_iter(f"{prefix}:chunk:*", count=500):
        batch.append(key)
        if len(batch) >= _EXPORT_BATCH:
            pipe = rc.pipeline(transaction=False)
            for k in batch:
                pipe.hget(k, "source")
            for raw in pipe.execute():
                if raw:
                    src = raw.decode() if isinstance(raw, bytes) else raw
                    if src:
                        sources.add(src)
            batch = []
    if batch:
        pipe = rc.pipeline(transaction=False)
        for k in batch:
            pipe.hget(k, "source")
        for raw in pipe.execute():
            if raw:
                src = raw.decode() if isinstance(raw, bytes) else raw
                if src:
                    sources.add(src)
    return sorted(sources)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES — SCHEDULED RECRAWL MANAGEMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/recrawl/sources")
def api_list_scheduled_sources():
    """List all scheduled re-crawl sources."""
    return _config.get("scheduled_sources", [])


@app.post("/api/recrawl/sources")
async def api_add_scheduled_source(payload: dict):
    """
    Add or update a URL in the re-crawl schedule.

    Body: {url, instance, depth}
    The scheduler will re-crawl this URL into the given RAG instance
    whenever recrawl.enabled is true and the configured interval has elapsed.
    """
    url = payload.get("url", "").strip()
    if not url:
        raise HTTPException(400, "url is required")
    sources = [s for s in _config.get("scheduled_sources", []) if s.get("url") != url]
    sources.append({
        "url":          url,
        "instance":     payload.get("instance", "default"),
        "depth":        int(payload.get("depth", 0)),
        "last_crawled": 0,
    })
    _config["scheduled_sources"] = sources
    save_config(_config)
    return {"ok": True}


@app.delete("/api/recrawl/sources")
async def api_delete_scheduled_source(url: str):
    """Remove a URL from the re-crawl schedule."""
    _config["scheduled_sources"] = [
        s for s in _config.get("scheduled_sources", []) if s.get("url") != url
    ]
    save_config(_config)
    return {"ok": True}


@app.post("/api/recrawl/trigger")
async def api_trigger_recrawl():
    """Immediately trigger a re-crawl of all scheduled sources (ignores interval)."""
    scheduled = _config.get("scheduled_sources", [])
    if not scheduled:
        return {"ok": True, "triggered": 0}
    triggered = 0
    now = time.time()
    for src in scheduled:
        url      = src.get("url", "").strip()
        instance = src.get("instance", "default")
        depth    = int(src.get("depth", 0))
        if not url:
            continue
        asyncio.create_task(crawl_url(instance, url, depth, rc=rc_for_instance(instance)))
        src["last_crawled"] = now
        triggered += 1
    save_config(_config)
    return {"ok": True, "triggered": triggered}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SERVE FRONTEND
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Serve index.html from beside this module (not the current working directory),
# so the app works regardless of where it is launched from.
_INDEX_HTML = Path(__file__).resolve().parent / "index.html"
if _INDEX_HTML.exists():
    @app.get("/")
    def serve_index():
        # no-cache so a rebuilt/edited UI is always picked up on reload (the
        # single-file frontend changes far more often than it's worth caching).
        return FileResponse(str(_INDEX_HTML), headers={"Cache-Control": "no-cache, must-revalidate"})


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CLI ENTRY POINT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def cli() -> None:
    """Run the RediRecall server (the ``redirecall`` console script).

    Host/port come from REDIRECALL_HOST / REDIRECALL_PORT (default
    127.0.0.1:8420). Bound to loopback by default — there is no built-in auth
    yet, so front it with a reverse proxy before exposing it to a network.
    """
    import uvicorn
    host = os.environ.get("REDIRECALL_HOST", "127.0.0.1")
    port = int(os.environ.get("REDIRECALL_PORT", "8420"))
    uvicorn.run(app, host=host, port=port)


if __name__ == "__main__":
    cli()
